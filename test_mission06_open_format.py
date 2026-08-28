"""Mission 06 — Open Format System: persistence, reference, custom, guided."""

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from open_format.config import GUIDED_FIELDS
from open_format.guided import create_guided_session, update_guided_session
from open_format.profile_store import (
    delete_user_profile,
    list_user_profiles,
    reload_user_profiles,
    save_user_profile,
    update_user_profile,
)
from open_format.reference_builder import analyze_reference_docx
from profiles import registry as profile_registry
from reformat_engine.profile_coverage import validate_profile_coverage
from reformat_engine.service import reformat_document


def _make_paragraph(doc, text, font=None, size=None, bold=None, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    if font is not None:
        run.font.name = font
        run._element.rPr.rFonts.set(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", font
        )
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    return p


# ---------------------------------------------------------------- persistence


def test_persistence_save_reload_update_delete(tmp_path, monkeypatch):
    monkeypatch.setenv("FORGE_HOME", str(tmp_path / "forge_home"))

    result = save_user_profile("my_custom", "我的格式", "desc", "custom", "generic_document", {"title": {"font": "黑体", "size_pt": 20}})
    assert result["status"] == "ok"
    assert profile_registry.is_user_profile("my_custom")
    assert profile_registry.resolve_profile("my_custom").title["font"] == "黑体"

    # simulate a fresh process / registry reload
    profile_registry.clear_user_profiles()
    assert not profile_registry.is_user_profile("my_custom")
    corrupted = reload_user_profiles()
    assert corrupted == []
    assert profile_registry.is_user_profile("my_custom")
    assert profile_registry.resolve_profile("my_custom").title["size_pt"] == 20

    update = update_user_profile("my_custom", {"rules": {"body": {"font": "仿宋", "size_pt": 14}}})
    assert update["status"] == "ok"
    assert profile_registry.resolve_profile("my_custom").body["font"] == "仿宋"

    assert delete_user_profile("my_custom")["status"] == "ok"
    assert not profile_registry.is_user_profile("my_custom")


def test_builtin_protection(tmp_path, monkeypatch):
    monkeypatch.setenv("FORGE_HOME", str(tmp_path / "forge_home"))
    result = save_user_profile("official_standard", "x", "", "custom", "generic_document", {})
    assert result["error"] == "PROFILE_ID_CONFLICT"
    result = delete_user_profile("official_standard")
    assert result["error"] == "BUILTIN_PROFILE_PROTECTED"


def test_inherited_parent_in_use(tmp_path, monkeypatch):
    monkeypatch.setenv("FORGE_HOME", str(tmp_path / "forge_home"))
    save_user_profile("parent_profile", "父", "", "custom", "generic_document", {})
    save_user_profile("child_profile", "子", "", "custom", "parent_profile", {})
    result = delete_user_profile("parent_profile")
    assert result["error"] == "PROFILE_IN_USE"
    # child can delete first
    assert delete_user_profile("child_profile")["status"] == "ok"
    assert delete_user_profile("parent_profile")["status"] == "ok"


def test_corrupted_user_json_reports_warning(tmp_path, monkeypatch):
    monkeypatch.setenv("FORGE_HOME", str(tmp_path / "forge_home"))
    profiles_dir = tmp_path / "forge_home" / "profiles"
    profiles_dir.mkdir(parents=True)
    (profiles_dir / "bad.json").write_text("{not valid json", encoding="utf-8")
    corrupted = reload_user_profiles()
    assert len(corrupted) == 1
    assert profile_registry.list_profiles()  # built-ins still present


def test_custom_profile_coverage_complete(tmp_path, monkeypatch):
    monkeypatch.setenv("FORGE_HOME", str(tmp_path / "forge_home"))
    save_user_profile(
        "coverage_ok",
        "覆盖完整",
        "",
        "custom",
        "generic_document",
        {"title": {"font": "黑体"}, "body": {"font": "宋体"}, "heading_1": {"font": "黑体"}, "page_number": {"enabled": True}},
    )
    coverage = validate_profile_coverage("coverage_ok")
    assert coverage.complete is True
    assert coverage.missing_slots == []


# ---------------------------------------------------------------- reference


def _make_reference_docx(path, body_font="宋体", body_size=12, body_count=10, split_font=None):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)
    _make_paragraph(doc, "参考标题", font="方正小标宋简体", size=22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _make_paragraph(doc, "一、参考一级标题", font="黑体", size=16, bold=False)
    half = body_count // 2
    for i in range(body_count):
        text = f"参考正文段落内容第{i}段，这里是用于参考学习的正文格式样本。"
        if split_font and i >= half:
            _make_paragraph(doc, text, font=split_font, size=body_size + 2)
        else:
            _make_paragraph(doc, text, font=body_font, size=body_size)
    doc.save(path)


def test_reference_learning_e2e(tmp_path, monkeypatch):
    monkeypatch.setenv("FORGE_HOME", str(tmp_path / "forge_home"))
    reference = tmp_path / "reference.docx"
    _make_reference_docx(reference, body_font="仿宋_GB2312", body_size=16)
    analysis = analyze_reference_docx(reference, base_profile_id="generic_document", profile_id="ref_learned", name="学到的格式")
    assert analysis["status"] == "ok"
    draft = analysis["draft"]
    assert draft["rules"]["body"]["font"] == "仿宋_GB2312"
    assert draft["rules"]["title"]["font"] == "方正小标宋简体"
    assert abs(draft["rules"]["page"]["top_cm"] - 3.7) < 0.05

    # finalize/save reference profile
    saved = save_user_profile(
        profile_id="ref_learned",
        name="学到的格式",
        description="reference learning",
        source="reference",
        inherits="generic_document",
        rules=draft["rules"],
    )
    assert saved["status"] == "ok"

    # use it on a completely different document
    source = tmp_path / "source.docx"
    doc = Document()
    p = doc.add_paragraph("完全不同内容的标题")
    p.runs[0].font.bold = True
    p.runs[0].font.size = Pt(22)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("完全不同内容的正文段落。")
    doc.save(source)
    out = tmp_path / "out.docx"
    result = reformat_document(source, output_path=out, saved_profile_id="ref_learned")
    assert result["status"] == "ok"
    from document_ir import read_docx

    out_ir = read_docx(out)
    texts = " ".join(getattr(b, "text", "") for b in out_ir.blocks)
    assert "完全不同内容" in texts
    assert "参考正文段落" not in texts  # reference content did not leak


def test_reference_conflict_needs_review(tmp_path, monkeypatch):
    monkeypatch.setenv("FORGE_HOME", str(tmp_path / "forge_home"))
    reference = tmp_path / "reference.docx"
    _make_reference_docx(reference, body_font="宋体", body_size=12, body_count=20, split_font="仿宋")
    analysis = analyze_reference_docx(reference)
    assert analysis["status"] == "needs_review"
    slots = [c["slot"] for c in analysis["conflicts"]]
    props = [c["property"] for c in analysis["conflicts"]]
    assert "body" in slots
    assert "font" in props


# ---------------------------------------------------------------- guided


def test_guided_e2e(tmp_path, monkeypatch):
    monkeypatch.setenv("FORGE_HOME", str(tmp_path / "forge_home"))
    created = create_guided_session(profile_id="guided_profile", name="引导格式", intent="正式公文")
    assert created["status"] == "needs_guidance"
    assert len(created["questions"]) <= 5
    session_id = created["session_id"]

    # answer: customize body, inherit everything else
    answers = []
    for field in GUIDED_FIELDS:
        if field == "body":
            answers.append({"field": field, "value": {"font": "仿宋", "size_pt": 16}})
        else:
            answers.append({"field": field, "inherit": True})
    updated = update_guided_session(session_id, answers)
    assert updated["status"] == "saved"
    assert profile_registry.is_user_profile("guided_profile")
    assert profile_registry.resolve_profile("guided_profile").body["font"] == "仿宋"


# ---------------------------------------------------------------- hard acceptance


def test_new_profile_without_core_code_change(tmp_path, monkeypatch):
    monkeypatch.setenv("FORGE_HOME", str(tmp_path / "forge_home"))
    save_user_profile(
        "my_new_profile",
        "我的新格式",
        "",
        "custom",
        "generic_document",
        {"title": {"font": "黑体", "size_pt": 20}, "body": {"font": "仿宋", "size_pt": 14}, "page_number": {"enabled": True}},
    )
    src = tmp_path / "src.docx"
    doc = Document()
    p = doc.add_paragraph("新格式标题")
    p.runs[0].font.bold = True
    p.runs[0].font.size = Pt(20)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("新格式正文")
    doc.save(src)
    out = tmp_path / "out.docx"
    result = reformat_document(src, output_path=out, saved_profile_id="my_new_profile")
    assert result["status"] == "ok"
    from document_ir import read_docx

    out_ir = read_docx(out)
    from semantics.annotator import annotate_document

    annotate_document(out_ir)
    title = next(b for b in out_ir.blocks if b.text == "新格式标题")
    assert title.style.font_name == "黑体"

    # assemble with the same saved profile
    from assembly_engine.service import assemble_documents

    a = tmp_path / "a.docx"
    b = tmp_path / "b.docx"
    for path, text in ((a, "汇编A"), (b, "汇编B")):
        d = Document()
        d.add_paragraph(text)
        d.save(path)
    assembled = tmp_path / "assembled.docx"
    asm = assemble_documents([str(a), str(b)], str(assembled), saved_profile_id="my_new_profile")
    assert asm["status"] == "ok"
    assert assembled.exists()

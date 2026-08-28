"""Mission 04-A Reformat Planning Layer 测试。"""

from dataclasses import asdict

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from document_ir import OpaqueBlock, ParagraphBlock, TableBlock, read_docx
from format_model import FormatProfile, FormatSource
from reformat_engine.models import Operation
from reformat_engine.planner import build_plan
from reformat_engine.profile_coverage import validate_profile_coverage
from semantics.annotator import annotate_document


def _annotated_ir(paragraphs, with_table=False, with_opaque=False):
    doc = Document()
    for text, style in paragraphs:
        p = doc.add_paragraph(text)
        if style.get("center"):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if style.get("bold"):
            for run in p.runs:
                run.font.bold = True
        if style.get("size"):
            for run in p.runs:
                run.font.size = Pt(style["size"])
    if with_table:
        table = doc.add_table(rows=1, cols=1)
        table.rows[0].cells[0].text = "数据"
    doc.save("/tmp/_planner.docx")
    if with_opaque:
        from zipfile import ZIP_DEFLATED, ZipFile

        with ZipFile("/tmp/_planner.docx") as zin:
            parts = {info.filename: zin.read(info.filename) for info in zin.infolist()}
        xml = parts["word/document.xml"].decode("utf-8")
        xml = xml.replace("</w:body>", "<w:customThing><w:r><w:t>隐藏</w:t></w:r></w:customThing></w:body>", 1)
        parts["word/document.xml"] = xml.encode("utf-8")
        with ZipFile("/tmp/_planner.docx", "w", ZIP_DEFLATED) as zout:
            for name, data in parts.items():
                zout.writestr(name, data)
    ir = read_docx("/tmp/_planner.docx")
    annotate_document(ir)
    return ir


def test_1_title_slot():
    ir = _annotated_ir([("幼儿园秋季运动会活动方案", {"center": True, "size": 22, "bold": True})])
    plan = build_plan(ir, "official_standard")
    assert plan.operations[0].style_slot == "title"
    assert plan.operations[0].action == "apply_profile_style"


def test_2_heading_1_slot():
    ir = _annotated_ir([("一、活动目标", {})])
    plan = build_plan(ir, "official_standard")
    assert plan.operations[0].style_slot == "heading_1"


def test_3_heading_2_slot():
    ir = _annotated_ir([("（一）活动准备", {})])
    plan = build_plan(ir, "official_standard")
    assert plan.operations[0].style_slot == "heading_2"


def test_4_heading_3_slot():
    ir = _annotated_ir([("1. 场地准备", {})])
    plan = build_plan(ir, "official_standard")
    assert plan.operations[0].style_slot == "heading_3"


def test_5_body_slot():
    ir = _annotated_ir([("这是一段比较长的普通正文段落，用来测试正文角色。", {})])
    plan = build_plan(ir, "official_standard")
    assert plan.operations[0].style_slot == "body"


def test_6_organization_slot():
    ir = _annotated_ir([("XX幼儿园", {}), ("活动方案", {"center": True, "size": 22, "bold": True})])
    plan = build_plan(ir, "official_standard")
    assert plan.operations[0].style_slot == "organization"


def test_7_author_slot():
    ir = _annotated_ir([("作者：张三", {})])
    plan = build_plan(ir, "official_standard")
    assert plan.operations[0].style_slot == "author"


def test_8_date_slot():
    ir = _annotated_ir([("2026年8月28日", {})])
    plan = build_plan(ir, "official_standard")
    assert plan.operations[0].style_slot == "date"


def test_9_signature_slot():
    ir = _annotated_ir([("XX幼儿园", {}), ("2026年8月28日", {})])
    plan = build_plan(ir, "official_standard")
    assert plan.operations[0].style_slot == "signature"


def test_10_caption_slot():
    ir = _annotated_ir([("图1 活动现场", {})])
    plan = build_plan(ir, "official_standard")
    assert plan.operations[0].style_slot == "caption"


def test_11_empty_preserve():
    ir = _annotated_ir([("", {})])
    plan = build_plan(ir, "official_standard")
    assert plan.operations[0].action == "preserve_structure"


def test_12_unknown_review_required():
    ir = _annotated_ir([("一、目标", {})])
    # Planner 相信语义层：手动模拟 unknown
    ir.blocks[0].semantic_role = "unknown"
    plan = build_plan(ir, "official_standard")
    assert plan.operations[0].action == "review_required"


def test_13_table_slot():
    ir = _annotated_ir([("正文", {})], with_table=True)
    plan = build_plan(ir, "official_standard")
    table_ops = [op for op in plan.operations if op.block_type == "table"]
    assert len(table_ops) == 1
    assert table_ops[0].style_slot == "table"


def test_14_opaque_unsupported():
    ir = _annotated_ir([("正文", {})], with_opaque=True)
    plan = build_plan(ir, "official_standard")
    opaque_ops = [op for op in plan.operations if op.block_type == "opaque"]
    assert len(opaque_ops) == 1
    assert opaque_ops[0].action == "unsupported"
    assert plan.ready is False


def test_15_full_coverage():
    from profiles import registry

    profile = registry.resolve_profile("generic_document")
    coverage = validate_profile_coverage(profile)
    assert coverage.complete is True
    assert coverage.missing_slots == []


def test_16_missing_slots():
    profile = FormatProfile(profile_id="partial", source=FormatSource())
    coverage = validate_profile_coverage(profile)
    assert coverage.complete is False
    assert "title" in coverage.missing_slots


def test_17_inherits_generic_complete():
    from profiles import registry

    registry.register_profile(
        FormatProfile(
            profile_id="_test_child_cover",
            source=FormatSource(),
            inherits="generic_document",
            title={"font": "宋体", "size_pt": 22},
        )
    )
    coverage = validate_profile_coverage("_test_child_cover")
    assert coverage.complete is True


def test_18_unknown_profile_not_found():
    ir = _annotated_ir([("正文", {})])
    plan = build_plan(ir, "no_such_profile")
    assert plan.ready is False
    assert "PROFILE_NOT_FOUND" in plan.blockers


def test_19_fingerprints_preserved():
    ir = _annotated_ir([("XX幼儿园", {}), ("活动方案", {"center": True, "size": 22, "bold": True}), ("作者：张三", {}), ("一、目标", {}), ("正文内容", {}), ("XX幼儿园", {}), ("2026年8月28日", {})])
    before = asdict(ir.content_fingerprint)
    build_plan(ir, "official_standard")
    after = asdict(ir.content_fingerprint)
    assert before == after


def test_20_deterministic():
    ir1 = _annotated_ir([("XX幼儿园", {}), ("活动方案", {"center": True, "size": 22, "bold": True}), ("一、目标", {}), ("正文", {})])
    ir2 = _annotated_ir([("XX幼儿园", {}), ("活动方案", {"center": True, "size": 22, "bold": True}), ("一、目标", {}), ("正文", {})])
    plan1 = build_plan(ir1, "official_standard")
    plan2 = build_plan(ir2, "official_standard")
    assert plan1.target_profile_id == plan2.target_profile_id
    assert [(o.block_id, o.style_slot, o.action) for o in plan1.operations] == [
        (o.block_id, o.style_slot, o.action) for o in plan2.operations
    ]

"""Mission 04-C — explicit override / ambiguous / invalid profile routing tests."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from document_ir import read_docx
from reformat_engine.service import reformat_document
from semantics.annotator import annotate_document


def _make_activity_plan(path):
    doc = Document()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("科技之春活动方案")
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    doc.add_paragraph("一、活动主题")
    doc.add_paragraph("二、活动目标")
    doc.add_paragraph("三、活动流程")
    doc.add_paragraph("四、活动准备")
    doc.save(path)


def _make_weak_document(path):
    doc = Document()
    doc.add_paragraph("今天天气不错。")
    doc.add_paragraph("大家讨论了一些日常安排。")
    doc.add_paragraph("会议在友好的气氛中结束。")
    doc.save(path)


def test_explicit_user_override_hint_beats_classification(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    _make_activity_plan(src)

    result = reformat_document(src, output_path=out, explicit_format_hint="正式公文")

    assert result["status"] == "ok"
    assert result["classification"]["intent"] == "activity_plan"
    assert result["resolution"]["profile_id"] == "official_standard"
    assert result["resolution"]["decision_basis"] == "explicit_user_choice"
    assert out.is_file()
    # 输出按 official_standard：标题字体为方正小标宋简体
    out_ir = read_docx(out)
    annotate_document(out_ir)
    title = next(b for b in out_ir.blocks if b.text == "科技之春活动方案")
    assert title.semantic_role == "title"
    assert title.style.font_name == "方正小标宋简体"


def test_ambiguous_returns_needs_guidance_no_output(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    _make_weak_document(src)

    result = reformat_document(src, output_path=out)

    assert result["status"] == "needs_guidance"
    assert result["output"] is None
    assert not out.exists()
    assert result["resolution"]["status"] == "needs_guidance"


def test_invalid_explicit_profile_not_found_no_output(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    _make_activity_plan(src)

    result = reformat_document(src, output_path=out, explicit_profile_id="no_such_profile")

    assert result["status"] == "error"
    assert result["resolution"]["error"] == "PROFILE_NOT_FOUND"
    assert result["output"] is None
    assert not out.exists()


def test_output_path_none_generates_forge_name(tmp_path):
    src = tmp_path / "src.docx"
    _make_activity_plan(src)
    expected = tmp_path / "src_FORGE.docx"

    result = reformat_document(src, explicit_format_hint="正式公文")

    assert result["status"] == "ok"
    assert result["output"] == str(expected)
    assert expected.is_file()

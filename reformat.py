"""Server-side document reformatting: extract content from an existing docx and
re-render it with a registered template, without routing the full text through
the AI.  This keeps large documents intact during reformatting.

Supported target types:
- 论文 / 演讲稿 / 发言稿 / 长文
- 传统公文 / 计划 / 总结 / 方案 / 汇报 / 报告 / 请示
- 传统公文活动方案 / 活动方案 / 活动总结
- 行政周报 / 周报
- 培训通知 / 培训活动记录 / 培训活动影像
- 活动影像
- 培训通知记录 (三合一完整版)
"""

from __future__ import annotations

import re
import shutil
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn

NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
}

CHAPTER_RE = re.compile(r"^第[一二三四五六七八九十百]+[章节篇]\s*\S*|^[一二三四五六七八九十]+、\s*\S*")
SECTION_RE = re.compile(r"^（[一二三四五六七八九十]+）\s*\S*|^\d+\.\d+(?!\d)\s*\S*")
SUBSECTION_RE = re.compile(r"^\d+[\.．、]\s*\S*")


A_BLIP = "{%s}blip" % NS["a"]


def _paragraph_images(p):
    embeds = []
    for elem in p._p.iter():
        if elem.tag == A_BLIP:
            embeds.append(elem.get(f"{{{NS['r']}}}embed"))
    return embeds


def extract_docx(path: Path, media_dir: Path):
    """Parse a docx into paragraphs, tables and image files."""
    media_dir.mkdir(parents=True, exist_ok=True)
    document = Document(path)
    rel_map = {}
    with ZipFile(path) as archive:
        try:
            rels = archive.read("word/_rels/document.xml.rels").decode("utf-8")
            rel_map = dict(re.findall(r'Id="([^"]+)"[^>]*Target="(media/[^"]+)"', rels))
        except KeyError:
            pass
        for name in archive.namelist():
            if name.startswith("word/media/"):
                data = archive.read(name)
                target = media_dir / Path(name).name
                target.write_bytes(data)

    paragraphs = []
    for p in document.paragraphs:
        embeds = _paragraph_images(p)
        images = [media_dir / Path(rel_map[e]).name for e in embeds if e in rel_map]
        paragraphs.append({"text": p.text.strip(), "images": images})

    tables = []
    for table in document.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append(rows)
    return document, paragraphs, tables


def pair_figures(paragraphs):
    """Group consecutive image paragraphs and attach the following 图X caption."""
    figures = []
    i = 0
    while i < len(paragraphs):
        if paragraphs[i]["images"]:
            paths = paragraphs[i]["images"]
            j = i + 1
            while j < len(paragraphs) and paragraphs[j]["images"]:
                paths += paragraphs[j]["images"]
                j += 1
            caption = ""
            if j < len(paragraphs) and paragraphs[j]["text"].startswith("图"):
                caption = paragraphs[j]["text"]
                j += 1
            figures.append({"paths": paths, "caption": caption})
            i = j
        else:
            i += 1
    return figures


def heading_level(text):
    if CHAPTER_RE.match(text):
        return 1
    if SECTION_RE.match(text):
        return 2
    if SUBSECTION_RE.match(text):
        return 3
    return 0


def build_thesis_content(document, paragraphs, tables):
    """Map source document to the 论文 template schema."""
    non_empty = [p for p in paragraphs if p["text"]]
    title = ""
    author = ""
    abstract = ""
    keywords = ""
    body_start = 0

    # Title: first meaningful paragraph before 摘要
    for i, p in enumerate(non_empty):
        if p["text"].startswith("摘要"):
            body_start = i
            abstract = p["text"].replace("摘要：", "").replace("摘要:", "").strip()
            if i + 1 < len(non_empty) and non_empty[i + 1]["text"].startswith("关键词"):
                keywords = non_empty[i + 1]["text"].replace("关键词：", "").replace("关键词:", "").strip()
                body_start = i + 2
            else:
                body_start = i + 1
            break
    title_candidates = [p["text"] for p in non_empty[:body_start] if not p["text"].startswith("摘要")]
    if title_candidates:
        title = title_candidates[-1]
        if len(title_candidates) >= 2:
            # 前面可能是作者或“幼儿教育·研究论文”之类，取最长的一行作标题
            title = max(title_candidates, key=len)

    # Author from info table
    for rows in tables:
        for row in rows:
            cells = [c.strip() for c in row]
            if cells and cells[0] in ("姓名", "姓名：", "作者", "作者："):
                author = cells[1] if len(cells) > 1 else ""
                break

    chapters = []
    current_chapter = None
    current_section = None
    current_subsection = None
    figures = []

    # Extract figures first (image paragraphs + captions), then skip them in body loop
    figure_by_index = {}
    figs = pair_figures(paragraphs)
    # Build body skipping image/caption paragraphs
    body_paragraphs = [p for p in paragraphs if not p["images"] and not p["text"].startswith("图")]

    for p in body_paragraphs:
        text = p["text"]
        if not text:
            continue
        if text == title or text.startswith("摘要") or text.startswith("关键词"):
            continue
        if text == author:
            continue
        level = heading_level(text)
        if level == 1:
            current_chapter = {"heading": text, "sections": []}
            chapters.append(current_chapter)
            current_section = None
            current_subsection = None
        elif level == 2 and current_chapter is not None:
            current_section = {"heading": text, "subsections": []}
            current_chapter["sections"].append(current_section)
            current_subsection = None
        elif level == 3 and current_section is not None:
            current_subsection = {"heading": text, "paragraphs": []}
            current_section["subsections"].append(current_subsection)
        else:
            if current_subsection is not None:
                current_subsection.setdefault("paragraphs", []).append(text)
            elif current_section is not None:
                current_section.setdefault("paragraphs", []).append(text)
            elif current_chapter is not None:
                current_chapter.setdefault("paragraphs", []).append(text)

    # Attach figures to the subsection whose heading is closest before the figure
    figure_specs = []
    for fig in figs:
        if len(fig["paths"]) == 1:
            figure_specs.append({"path": str(fig["paths"][0]), "caption": fig["caption"]})
        else:
            figure_specs.append({"paths": [str(x) for x in fig["paths"]], "caption": fig["caption"], "width_cm": 13, "gap_cm": 0.2})

    content = {
        "title": title,
        "author": author,
        "abstract": abstract,
        "keywords": keywords,
        "chapters": chapters,
        "figures": figure_specs,
    }
    # Info tables (skip tables that look like the author info table? keep as three-line tables)
    info_tables = []
    for rows in tables:
        if not rows:
            continue
        headers = [c.strip() for c in rows[0]]
        data = [[c.strip() for c in row] for row in rows[1:]]
        if headers and data:
            info_tables.append({"caption": "", "headers": headers, "rows": data})
    if info_tables:
        content["tables"] = info_tables
    return content


def build_official_content(document, paragraphs, tables, include_department=True):
    """Map source document to 传统公文 / 活动方案 / 活动总结 schema."""
    non_empty = [p["text"] for p in paragraphs if p["text"]]
    if not non_empty:
        return {}

    # 标题块：取前两行作为单位名 + 标题；若有副标题则为第三行
    organization = non_empty[0]
    title = non_empty[1] if len(non_empty) > 1 else ""
    subtitle = ""
    idx = 2
    if len(non_empty) > 2 and not heading_level(non_empty[2]):
        subtitle = non_empty[2]
        idx = 3

    # 开头段：标题后第一段
    opening = non_empty[idx] if idx < len(non_empty) else ""
    idx += 1

    sections = []
    current_section = None
    current_subsection = None

    for text in non_empty[idx:]:
        if text in (organization, title, subtitle, opening):
            continue
        level = heading_level(text)
        if level == 1:
            current_section = {"heading": text, "subsections": []}
            sections.append(current_section)
            current_subsection = None
        elif level == 2 and current_section is not None:
            current_subsection = {"heading": text, "items": []}
            current_section["subsections"].append(current_subsection)
        elif level == 3 and current_subsection is not None:
            current_subsection.setdefault("items", []).append({"heading": text, "paragraphs": []})
        else:
            # 落款行
            if re.fullmatch(r"[\d年月日]{6,}", text) or re.match(r"^\d{4}年\d{1,2}月\d{1,2}日", text):
                signature_date = text
            if current_subsection is not None and current_subsection.get("items"):
                current_subsection["items"][-1].setdefault("paragraphs", []).append(text)
            elif current_section is not None and current_section.get("subsections"):
                current_section["subsections"][-1].setdefault("paragraphs", []).append(text)
            elif current_section is not None:
                current_section.setdefault("paragraphs", []).append(text)

    # 落款：取末尾 2~3 个“非正文”段落
    tail = [t for t in non_empty[idx:] if not heading_level(t)]
    signature = organization
    department = ""
    date = ""
    if tail:
        date = tail[-1]
    if len(tail) >= 2:
        department = tail[-2]
    if len(tail) >= 3:
        signature = tail[-3]

    content = {
        "organization": organization,
        "title": title,
        "subtitle": subtitle,
        "opening": opening,
        "sections": sections,
        "attachment": "",
        "signature": signature,
        "date": date,
    }
    if include_department:
        content["department"] = department
    return content


def build_weekly_content(document, paragraphs, tables):
    """Map source document to 行政周报 schema."""
    non_empty = [p["text"] for p in paragraphs if p["text"]]
    title_line1 = non_empty[0] if non_empty else ""
    title_line2 = non_empty[1] if len(non_empty) > 1 else ""
    # title_line1 = organization + semester + 行政周报
    m = re.match(r"^(.*?)(\d{4}年(?:春季|秋季)学期)行政周报$", title_line1)
    organization = m.group(1) if m else (title_line1.replace("行政周报", "") if title_line1 else "")
    semester = m.group(2) if m else ""
    m2 = re.match(r"^第(\d+)周（(.+)）$", title_line2)
    week = m2.group(1) if m2 else ""
    date_range = m2.group(2) if m2 else ""

    sections = []
    current_section = None
    current_part = None
    for text in non_empty[2:]:
        if re.match(r"^[一二三四五六七八九十]+、", text):
            current_section = {"name": text, "summary": [], "problems": [], "plan": []}
            sections.append(current_section)
            current_part = None
        elif "本周工作总结" in text:
            current_part = "summary"
        elif "存在问题" in text:
            current_part = "problems"
        elif "下周重点工作" in text:
            current_part = "plan"
        elif current_section is not None and current_part:
            item = re.sub(r"^\s*\d+\s*[．.、]\s*", "", text)
            if item and item != "无":
                current_section[current_part].append(item)

    for sec in sections:
        for part in ("summary", "problems", "plan"):
            if not sec[part]:
                sec[part] = ["无"]
    return {
        "organization": organization,
        "semester": semester,
        "week": week,
        "date_range": date_range,
        "sections": sections,
    }


def table_cell(tables, label):
    for rows in tables:
        for row in rows:
            for i, cell in enumerate(row):
                if cell.strip() == label and i + 1 < len(row):
                    return row[i + 1].strip()
    return ""


def build_training_notice_content(document, paragraphs, tables):
    non_empty = [p["text"] for p in paragraphs if p["text"]]
    organization = non_empty[0] if non_empty else ""
    document_title = next((t for t in non_empty if t.endswith("通知")), "")
    purpose = next((t for t in non_empty if "现将" in t or "有关事项通知如下" in t), "")
    date = next((t for t in reversed(non_empty) if re.match(r"^\d{4}年\d{1,2}月\d{1,2}日$", t)), "")
    tail = [t for t in non_empty if re.match(r"^\d{4}年\d{1,2}月\d{1,2}日$", t)]
    department = tail[-2] if len(tail) >= 2 else ""
    location = ""
    training_topic = ""
    trainer = ""
    participants = ""
    req1 = ""
    req2 = ""
    for t in non_empty:
        if "活动地点" in t:
            location = t.split("：")[-1].split(":")[-1].strip()
        elif "活动内容" in t:
            training_topic = t.split("：")[-1].split(":")[-1].strip()
        elif "培训人" in t:
            trainer = t.split("：")[-1].split(":")[-1].strip()
        elif "活动对象" in t:
            participants = t.split("：")[-1].split(":")[-1].strip()
        elif "准时参会" in t or "学习记录" in t:
            req1 = t
        elif "对照反思" in t or "对照自查" in t:
            req2 = t

    date_short = table_cell(tables, "培训时间") or table_cell(tables, "时间")
    hours = table_cell(tables, "学时")
    training_content = table_cell(tables, "培训内容") or table_cell(tables, "培 / 训 / 内 / 容")
    training_reflection = table_cell(tables, "培训心得")

    return {
        "organization": organization,
        "document_title": document_title or "培训通知",
        "purpose": purpose,
        "date": date,
        "location": location,
        "training_topic": training_topic,
        "trainer": trainer,
        "participants": participants,
        "requirement_1": req1,
        "requirement_2": req2,
        "department": department,
        "date_short": date_short,
        "hours": hours,
        "training_content": training_content,
        "training_reflection": training_reflection,
    }


def build_training_record_content(document, paragraphs, tables):
    return build_training_notice_content(document, paragraphs, tables)


def build_media_content(document, paragraphs, tables, kind):
    """培训活动影像 / 活动影像 通用信息表解析。"""
    date_short = table_cell(tables, "时间") or table_cell(tables, "培训时间")
    location = table_cell(tables, "培训地点") or table_cell(tables, "活动地点")
    if kind == "培训活动影像":
        participants = table_cell(tables, "参 培 人") or table_cell(tables, "参培人")
        training_topic = table_cell(tables, "培训内容") or table_cell(tables, "培训主题")
        return {
            "date_short": date_short,
            "location": location,
            "participants": participants,
            "training_topic": training_topic,
        }
    organizer = table_cell(tables, "负责人")
    activity_name = table_cell(tables, "活动内容")
    title = next((p["text"] for p in paragraphs if p["text"].endswith("活动影像")), "")
    organization = title.replace("活动影像", "") if title else ""
    return {
        "organization": organization,
        "date_short": date_short,
        "location": location,
        "organizer": organizer,
        "activity_name": activity_name,
    }


def build_reformat_content(path: Path, document_type: str, media_dir: Path):
    document, paragraphs, tables = extract_docx(path, media_dir)
    if document_type in ("论文", "演讲稿", "发言稿", "长文"):
        return build_thesis_content(document, paragraphs, tables)
    if document_type in ("传统公文", "计划", "总结", "方案", "汇报", "报告", "请示"):
        return build_official_content(document, paragraphs, tables, include_department=False)
    if document_type in ("传统公文活动方案", "活动方案", "活动总结"):
        return build_official_content(document, paragraphs, tables, include_department=True)
    if document_type in ("行政周报", "周报"):
        return build_weekly_content(document, paragraphs, tables)
    if document_type == "培训通知":
        return build_training_notice_content(document, paragraphs, tables)
    if document_type == "培训活动记录":
        return build_training_record_content(document, paragraphs, tables)
    if document_type == "培训活动影像":
        return build_media_content(document, paragraphs, tables, "培训活动影像")
    if document_type == "活动影像":
        return build_media_content(document, paragraphs, tables, "活动影像")
    raise ValueError(f"reformat 暂不支持该类型: {document_type}")

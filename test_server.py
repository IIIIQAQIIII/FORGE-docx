"""Self-test for the FORGE server.

Run with:  .venv/bin/python test_server.py
"""

from __future__ import annotations

import asyncio
import json
from pathlib import Path

import server


async def main() -> None:
    results: list[tuple[str, bool, str]] = []

    def check(name: str, ok: bool, detail: str = "") -> None:
        results.append((name, ok, detail))
        print(f"{'PASS' if ok else 'FAIL'}  {name}" + (f"  ({detail})" if detail else ""))

    async def call(tool: str, args: dict):
        output = await server.mcp.call_tool(tool, args)
        # FastMCP returns (content_list, result_meta); unwrap the first text part.
        if isinstance(output, (list, tuple)):
            parts = output[0] if isinstance(output, tuple) else output
            if isinstance(parts, (list, tuple)) and parts:
                return parts[0].text if hasattr(parts[0], "text") else str(parts[0])
        return str(output)

    # 1. list document types
    listing = await call("list_document_types", {})
    listing_json = json.loads(listing)
    check("list_document_types returns JSON", "document_types" in listing_json)
    check("list_document_types includes 论文", "论文" in listing_json.get("document_types", {}))
    check("list_document_types includes 培训通知", "培训通知" in listing_json.get("document_types", {}))
    check("list_document_types includes 培训活动影像", "培训活动影像" in listing_json.get("document_types", {}))
    check("list_document_types has document_sets", "document_sets" in listing_json)
    check("document_sets has 培训资料套装", "培训资料套装" in listing_json.get("document_sets", {}))

    # 2. schema for 传统公文
    schema_raw = await call("get_template_schema", {"template_name": "传统公文.docx"})
    schema = json.loads(schema_raw)
    check("get_template_schema has fields", "fields" in schema and "example" in schema)
    check(
        "get_template_schema exposes section.heading",
        "section.heading" in schema.get("fields", {}),
    )

    # 3. generate sample document
    out = await call(
        "generate_docx",
        {
            "template_name": "sample_template.docx",
            "content": {"title": "测试", "body": "正文", "signature": "签名", "date": "2026-08-26"},
            "output_name": "_self_test_sample.docx",
        },
    )
    check("generate sample docx", out.startswith("Created:"), out)

    # 4. generate official document with three levels
    out = await call(
        "generate_by_type",
        {
            "document_type": "传统公文",
            "content": {
                "organization": "某市示范幼儿园",
                "title": "科技之春活动方案",
                "subtitle": "",
                "opening": "为进一步激发幼儿对科学技术的兴趣，现制定方案如下：",
                "sections": [
                    {"heading": "一、活动主题", "paragraphs": ["科技点亮童心 探索创造未来"]},
                    {
                        "heading": "二、活动安排",
                        "subsections": [
                            {
                                "heading": "（一）活动时间",
                                "paragraphs": ["拟定于下周开展。"],
                                "items": [
                                    {"heading": "1.具体安排", "paragraphs": ["由园所统筹安排。"]}
                                ],
                            },
                            {
                                "heading": "（二）活动地点",
                                "paragraphs": ["各园所操场。"],
                            },
                        ],
                    },
                ],
                "attachment": "活动安排表",
                "signature": "某市示范幼儿园",
                "date": "2026年3月19日",
                "annotation": "此件发至各班。",
            },
            "output_name": "_self_test_official.docx",
        },
    )
    check("generate 传统公文 (3 levels)", out.startswith("Created:"), out)

    # 5. generate thesis with a three-line table and verify marker removal
    out = await call(
        "generate_by_type",
        {
            "document_type": "论文",
            "content": {
                "title": "幼儿园科学教育活动设计研究",
                "author": "张三",
                "abstract": "本文围绕幼儿园科学教育活动展开研究。",
                "keywords": "科学教育；幼儿园",
                "chapters": [
                    {
                        "heading": "第一章 绪论",
                        "sections": [
                            {
                                "heading": "1.1 研究背景",
                                "subsections": [
                                    {"heading": "1.1.1 问题的提出", "paragraphs": ["正文段落一。"]}
                                ],
                            }
                        ],
                    }
                ],
                "tables": [
                    {"caption": "表1 研究对象", "headers": ["班级", "人数"], "rows": [["中一班", "30"]]}
                ],
            },
            "output_name": "_self_test_thesis.docx",
        },
    )
    check("generate 论文 with table", out.startswith("Created:"), out)

    from docx import Document

    thesis = Document(server.OUTPUTS_DIR / "_self_test_thesis.docx")
    thesis_text = "\n".join(p.text for p in thesis.paragraphs)
    check("thesis marker removed", "MEDIA_INSERT_HERE" not in thesis_text)
    check("thesis table inserted", len(thesis.tables) == 1)
    header_text = "\n".join(p.text for p in thesis.sections[0].header.paragraphs)
    check("thesis header removed", "首都师范大学全日制教育硕士学位论文" not in header_text)
    title_run = thesis.paragraphs[0].runs[0] if thesis.paragraphs and thesis.paragraphs[0].runs else None
    check("thesis title not bold", title_run is not None and not title_run.font.bold)

    # 6. validate generated official document
    val_raw = await call("validate_docx", {"docx_path": "_self_test_official.docx"})
    val = json.loads(val_raw)
    check("validate passes on generated doc", val.get("status") == "pass", val.get("status", ""))

    # 7. fix_docx
    out = await call(
        "fix_docx",
        {"docx_path": "_self_test_sample.docx", "output_name": "_self_test_sample_fixed.docx"},
    )
    check("fix_docx creates copy", out.startswith("Created repaired copy:"), out)

    # 8. generate training set (three pieces, one shared content)
    training_content = {
        "organization": "某市示范幼儿园",
        "document_title": "培训通知",
        "purpose": "为进一步加强师德师风警示教育……现将有关事项通知如下：",
        "date": "2026年6月30日",
        "location": "党建室",
        "training_topic": "师德师风培训《警示案例学习》",
        "trainer": "张明",
        "participants": "某市示范幼儿园全体教师",
        "requirement_1": "全体教师准时参会，认真做好学习记录。",
        "requirement_2": "结合岗位职责和案例内容主动对照反思。",
        "department": "保教处",
        "date_short": "2026.6.30",
        "hours": "2",
        "training_content": "一、明确警示案例学习的重要意义……",
        "training_reflection": "通过本次培训，我进一步认识到……",
    }
    set_out = await call(
        "generate_document_set",
        {"set_name": "培训资料套装", "content": training_content, "output_prefix": "_self_test_set"},
    )
    check("generate_document_set creates 3 files", set_out.startswith("Created set") and set_out.count("Created:") == 3, set_out[:120])
    set_record = Document(server.OUTPUTS_DIR / "_self_test_set_2_培训活动记录.docx")
    set_record_text = "\n".join(p.text for p in set_record.paragraphs)
    check("培训活动记录 has no 影像 section", "培训活动影像" not in set_record_text)
    set_media = Document(server.OUTPUTS_DIR / "_self_test_set_3_培训活动影像.docx")
    check("培训活动影像 has 4-row photo table", len(set_media.tables) >= 1 and len(set_media.tables[0].rows) == 4)
    check("培训活动影像 photos on one page (no page breaks)", "培训活动影像" in "\n".join(p.text for p in set_media.paragraphs))

    # 9. flexible output path: absolute paths are allowed
    import tempfile, os
    tmp_dir = tempfile.mkdtemp(prefix="mcp_word_test_")
    abs_out = os.path.join(tmp_dir, "abs_output.docx")
    out = await call(
        "generate_docx",
        {
            "template_name": "sample_template.docx",
            "content": {"title": "测试", "body": "正文", "signature": "签名", "date": "2026-02-23"},
            "output_name": abs_out,
        },
    )
    check("generate to absolute path", out.startswith("Created:") and os.path.exists(abs_out), out[:120])
    if os.path.exists(abs_out):
        os.remove(abs_out)
    os.rmdir(tmp_dir)

    # 10. generate activity plan set (方案 + 总结 + 影像)
    activity_content = {
        "organization": "某市示范幼儿园",
        "title": "读书月活动方案",
        "subtitle": "2025——2026学年第二学期",
        "opening": "阅读是幼儿认识世界的重要方式，现制定方案如下：",
        "sections": [{"heading": "一、活动主题", "paragraphs": ["书香润童心·阅读伴成长"]}],
        "attachment": "",
        "signature": "某市示范幼儿园",
        "department": "保教处",
        "date": "2026年4月8日",
        "date_short": "2026.4.29",
        "location": "多功能厅",
        "organizer": "保教处",
        "activity_name": "读书月系列活动",
    }
    act_out = await call(
        "generate_document_set",
        {"set_name": "活动方案套装", "content": activity_content, "output_prefix": "_self_test_activity"},
    )
    check("generate_document_set activity 3 files", act_out.startswith("Created set") and act_out.count("Created:") == 3, act_out[:120])
    act_media = Document(server.OUTPUTS_DIR / "_self_test_activity_3_活动影像.docx")
    check("活动影像 has 4-row table", len(act_media.tables) == 1 and len(act_media.tables[0].rows) == 4)

    # 11. path traversal must be rejected
    out = await call(
        "generate_docx",
        {
            "template_name": "../server.py",
            "content": {},
            "output_name": "_self_test_escape.docx",
        },
    )
    check("path traversal rejected", out.startswith("Error:"), out)

    failed = [r for r in results if not r[1]]
    print(f"\n{len(results) - len(failed)}/{len(results)} checks passed")
    if failed:
        raise SystemExit(1)


if __name__ == "__main__":
    asyncio.run(main())

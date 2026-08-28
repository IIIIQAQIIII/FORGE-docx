"""Mission 05 — Batch Assemble Engine core tests."""

import struct
import zlib
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from assembly_engine.service import assemble_documents
from document_ir import ParagraphBlock, TableBlock, read_docx


def _png_bytes(color=b"\xAA\xBB\xCC"):
    def chunk(tag, data):
        c = tag + data
        return struct.pack(">I", len(data)) + c + struct.pack(">I", zlib.crc32(c))

    ihdr = struct.pack(">IIBBBBB", 8, 8, 8, 2, 0, 0, 0)
    idat = zlib.compress(b"\x00" + color * 8)
    return b"\x89PNG\r\n\x1a\n" + chunk(b"IHDR", ihdr) + chunk(b"IDAT", idat) + chunk(b"IEND", b"")


def _make_doc(path, texts):
    doc = Document()
    for text in texts:
        doc.add_paragraph(text)
    doc.save(path)


def _zip_read(path, name):
    with ZipFile(path) as z:
        return z.read(name).decode("utf-8")


def _inject_numbering_part(path, lvl_text="%1."):
    from lxml import etree

    numbering_xml = (
        '<w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:abstractNum w:abstractNumId="0"><w:multiLevelType w:val="hybridMultilevel"/>'
        '<w:lvl w:ilvl="0"><w:start w:val="1"/><w:numFmt w:val="decimal"/>'
        f'<w:lvlText w:val="{lvl_text}"/><w:lvlJc w:val="left"/>'
        '<w:pPr><w:ind w:left="420" w:hanging="420"/></w:pPr></w:lvl>'
        '</w:abstractNum><w:num w:numId="1"><w:abstractNumId w:val="0"/></w:num></w:numbering>'
    ).encode("utf-8")
    with ZipFile(path) as zin:
        parts = {info.filename: zin.read(info.filename) for info in zin.infolist()}
    parts["word/numbering.xml"] = numbering_xml
    ct_ns = "http://schemas.openxmlformats.org/package/2006/content-types"
    rels_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
    ct_root = etree.fromstring(parts["[Content_Types].xml"])
    override = etree.SubElement(ct_root, f"{{{ct_ns}}}Override")
    override.set("PartName", "/word/numbering.xml")
    override.set("ContentType", "application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml")
    parts["[Content_Types].xml"] = etree.tostring(ct_root, xml_declaration=True, encoding="UTF-8", standalone=True)
    rels_root = etree.fromstring(parts["word/_rels/document.xml.rels"])
    max_n = 0
    for rel in rels_root.findall(f"{{{rels_ns}}}Relationship"):
        rid = rel.get("Id") or ""
        if rid.startswith("rId"):
            try:
                max_n = max(max_n, int(rid[3:]))
            except ValueError:
                pass
    rel = etree.SubElement(rels_root, f"{{{rels_ns}}}Relationship")
    rel.set("Id", f"rId{max_n + 1}")
    rel.set("Type", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering")
    rel.set("Target", "numbering.xml")
    parts["word/_rels/document.xml.rels"] = etree.tostring(
        rels_root, xml_declaration=True, encoding="UTF-8", standalone=True
    )
    with ZipFile(path, "w", ZIP_DEFLATED) as zout:
        for name, data in parts.items():
            zout.writestr(name, data)


def _make_numbered_doc(path, texts, lvl_text="%1."):
    doc = Document()
    for text in texts:
        p = doc.add_paragraph(text)
        ppr = p._p.get_or_add_pPr()
        numpr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        numid = OxmlElement("w:numId")
        numid.set(qn("w:val"), "1")
        numpr.append(ilvl)
        numpr.append(numid)
        ppr.append(numpr)
    doc.save(path)
    _inject_numbering_part(path, lvl_text)


def _block_texts(path):
    ir = read_docx(path)
    return [getattr(b, "text", "") for b in ir.blocks]


def test_1_two_docs_input_order(tmp_path):
    a = tmp_path / "a.docx"
    b = tmp_path / "b.docx"
    out = tmp_path / "assembled.docx"
    _make_doc(a, ["AAAA"])
    _make_doc(b, ["BBBB"])
    res = assemble_documents([str(a), str(b)], str(out), explicit_format_hint="正式公文")
    assert res["status"] == "ok"
    texts = _block_texts(out)
    assert texts[0] == "AAAA"
    assert texts[2] == "BBBB"


def test_2_order_mode_filename(tmp_path):
    a = tmp_path / "a.docx"
    b = tmp_path / "b.docx"
    out = tmp_path / "assembled.docx"
    _make_doc(a, ["AAAA"])
    _make_doc(b, ["BBBB"])
    res = assemble_documents([str(b), str(a)], str(out), explicit_format_hint="正式公文", order_mode="filename")
    assert res["status"] == "ok"
    texts = _block_texts(out)
    assert texts[0] == "AAAA"
    assert texts[2] == "BBBB"


def test_3_each_source_content_kept(tmp_path):
    a = tmp_path / "a.docx"
    b = tmp_path / "b.docx"
    out = tmp_path / "assembled.docx"
    _make_doc(a, ["AAA", "BBB"])
    _make_doc(b, ["CCC", "DDD"])
    res = assemble_documents([str(a), str(b)], str(out), explicit_format_hint="正式公文")
    assert res["status"] == "ok"
    assert all(item["per_item_preservation"]["passed"] for item in res["items"])
    texts = [t for t in _block_texts(out) if t]
    assert texts == ["AAA", "BBB", "CCC", "DDD"]


def test_4_source_sha_unchanged(tmp_path):
    a = tmp_path / "a.docx"
    b = tmp_path / "b.docx"
    out = tmp_path / "assembled.docx"
    _make_doc(a, ["AAA"])
    _make_doc(b, ["BBB"])
    before = {p.name: p.read_bytes() for p in (a, b)}
    assemble_documents([str(a), str(b)], str(out), explicit_format_hint="正式公文")
    assert a.read_bytes() == before["a.docx"]
    assert b.read_bytes() == before["b.docx"]


def test_5_page_break_between_items(tmp_path):
    a = tmp_path / "a.docx"
    b = tmp_path / "b.docx"
    out = tmp_path / "assembled.docx"
    _make_doc(a, ["AAA"])
    _make_doc(b, ["BBB"])
    assemble_documents([str(a), str(b)], str(out), explicit_format_hint="正式公文")
    xml = _zip_read(out, "word/document.xml")
    assert 'w:type="page"' in xml


def test_6_continuous_page_number(tmp_path):
    a = tmp_path / "a.docx"
    b = tmp_path / "b.docx"
    out = tmp_path / "assembled.docx"
    _make_doc(a, ["AAA"])
    _make_doc(b, ["BBB"])
    assemble_documents([str(a), str(b)], str(out), explicit_format_hint="正式公文")
    xml = _zip_read(out, "word/document.xml")
    assert 'w:pgNumType' not in xml
    with ZipFile(out) as z:
        footers = [n for n in z.namelist() if n.startswith("word/footer") and n.endswith(".xml")]
    assert footers
    assert "PAGE" in _zip_read(out, footers[0])


def test_7_duplicate_text_no_mismatch(tmp_path):
    a = tmp_path / "a.docx"
    b = tmp_path / "b.docx"
    out = tmp_path / "assembled.docx"
    _make_doc(a, ["相同内容"])
    _make_doc(b, ["相同内容"])
    res = assemble_documents([str(a), str(b)], str(out), explicit_format_hint="正式公文")
    assert res["status"] == "ok"
    assert all(item["per_item_preservation"]["passed"] for item in res["items"])
    non_empty = [t for t in _block_texts(out) if t]
    assert non_empty == ["相同内容", "相同内容"]


def test_8_media_name_conflict_two_different_images(tmp_path):
    dir_a = tmp_path / "a"
    dir_b = tmp_path / "b"
    dir_a.mkdir()
    dir_b.mkdir()
    png_a = dir_a / "image1.png"
    png_b = dir_b / "image1.png"
    png_a.write_bytes(_png_bytes(b"\x11\x11\x11"))
    png_b.write_bytes(_png_bytes(b"\x22\x22\x22"))
    a = tmp_path / "a.docx"
    b = tmp_path / "b.docx"
    out = tmp_path / "assembled.docx"
    da = Document()
    da.add_picture(str(png_a))
    da.save(a)
    db = Document()
    db.add_picture(str(png_b))
    db.save(b)
    res = assemble_documents([str(a), str(b)], str(out), explicit_format_hint="正式公文")
    assert res["status"] == "ok"
    ir = read_docx(out)
    assert ir.statistics.image_count == 2
    shas = {m.sha256 for m in ir.media}
    assert len(shas) == 2
    import hashlib

    assert hashlib.sha256(png_a.read_bytes()).hexdigest() in shas
    assert hashlib.sha256(png_b.read_bytes()).hexdigest() in shas


def test_9_media_sha_kept(tmp_path):
    png = tmp_path / "image1.png"
    png.write_bytes(_png_bytes())
    a = tmp_path / "a.docx"
    out = tmp_path / "assembled.docx"
    d = Document()
    d.add_paragraph("带图")
    d.add_picture(str(png))
    d.save(a)
    res = assemble_documents([str(a)], str(out), explicit_format_hint="正式公文")
    assert res["status"] == "ok"
    ir = read_docx(out)
    import hashlib

    assert ir.media[0].sha256 == hashlib.sha256(png.read_bytes()).hexdigest()


def test_10_hyperlink_target_kept(tmp_path):
    a = tmp_path / "a.docx"
    out = tmp_path / "assembled.docx"
    doc = Document()
    p = doc.add_paragraph()
    rel_id = doc.part.relate_to(
        "https://example.com",
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "点击查看"
    run.append(t)
    hyperlink.append(run)
    p._p.append(hyperlink)
    doc.save(a)
    res = assemble_documents([str(a)], str(out), explicit_format_hint="正式公文")
    assert res["status"] == "ok"
    ir = read_docx(out)
    block = next(b for b in ir.blocks if isinstance(b, ParagraphBlock) and b.text == "点击查看")
    assert block.inline[0].type == "hyperlink"
    assert block.inline[0].target == "https://example.com"


def test_11_numbering_id_remap(tmp_path):
    a = tmp_path / "a.docx"
    b = tmp_path / "b.docx"
    out = tmp_path / "assembled.docx"
    _make_numbered_doc(a, ["A1", "A2"], lvl_text="%1.")
    _make_numbered_doc(b, ["B1", "B2"], lvl_text="(%1)")
    res = assemble_documents([str(a), str(b)], str(out), explicit_format_hint="正式公文")
    assert res["status"] == "ok"
    xml = _zip_read(out, "word/document.xml")
    import re

    num_ids = re.findall(r"<w:numId w:val=\"(\d+)\"/>", xml)
    assert len(num_ids) == 4
    assert len(set(num_ids[:2])) == 1  # A paragraphs share one new numId
    assert len(set(num_ids[2:])) == 1  # B paragraphs share another new numId
    assert num_ids[0] != num_ids[2]
    numbering_xml = _zip_read(out, "word/numbering.xml")
    assert 'w:val="%1."' in numbering_xml
    assert 'w:val="(%1)"' in numbering_xml


def test_12_style_collision_safe(tmp_path):
    a = tmp_path / "a.docx"
    b = tmp_path / "b.docx"
    out = tmp_path / "assembled.docx"

    da = Document()
    sa = da.styles.add_style("CustomStyle", WD_STYLE_TYPE.PARAGRAPH)
    sa.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    pa = da.add_paragraph("A 使用样式")
    pa.style = sa
    da.save(a)

    db = Document()
    sb = db.styles.add_style("CustomStyle", WD_STYLE_TYPE.PARAGRAPH)
    sb.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
    pb = db.add_paragraph("B 使用样式")
    pb.style = sb
    db.save(b)

    res = assemble_documents([str(a), str(b)], str(out), explicit_format_hint="正式公文")
    assert res["status"] == "ok"
    xml = _zip_read(out, "word/document.xml")
    assert 'w:val="CustomStyle"' in xml
    assert 'w:val="Assembly_CustomStyle_0"' in xml
    styles_xml = _zip_read(out, "word/styles.xml")
    assert 'w:styleId="CustomStyle"' in styles_xml
    assert 'w:styleId="Assembly_CustomStyle_0"' in styles_xml


def test_13_table_gridspan_vmerge_preserved(tmp_path):
    a = tmp_path / "a.docx"
    out = tmp_path / "assembled.docx"
    doc = Document()
    t1 = doc.add_table(rows=1, cols=2)
    t1.cell(0, 0).merge(t1.cell(0, 1))
    t1.cell(0, 0).text = "合并列"
    t2 = doc.add_table(rows=2, cols=1)
    t2.cell(0, 0).text = "上半"
    t2.cell(1, 0).text = "下半"
    t2.cell(0, 0).merge(t2.cell(1, 0))
    doc.save(a)
    res = assemble_documents([str(a)], str(out), explicit_format_hint="正式公文")
    assert res["status"] == "ok"
    ir = read_docx(out)
    tables = [b for b in ir.blocks if isinstance(b, TableBlock)]
    assert len(tables) == 2
    assert tables[0].rows[0][0].metadata.get("grid_span") == "2"
    assert tables[1].rows[0][0].metadata.get("v_merge") == "restart"
    assert tables[1].rows[1][0].metadata.get("v_merge") == "continue"


def test_14_table_cell_image_preserved(tmp_path):
    png = tmp_path / "cell.png"
    png.write_bytes(_png_bytes())
    a = tmp_path / "a.docx"
    out = tmp_path / "assembled.docx"
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    cell.paragraphs[0].text = "图在单元格"
    cell.add_paragraph().add_run().add_picture(str(png), width=Cm(3))
    doc.save(a)
    res = assemble_documents([str(a)], str(out), explicit_format_hint="正式公文")
    assert res["status"] == "ok"
    ir = read_docx(out)
    table = next(b for b in ir.blocks if isinstance(b, TableBlock))
    assert ir.statistics.image_count == 1
    assert table.rows[0][0].blocks[1].inline[0].type == "image"
    import hashlib

    assert ir.media[0].sha256 == hashlib.sha256(png.read_bytes()).hexdigest()


def test_15_review_required_item_success_with_warnings(tmp_path, monkeypatch):
    import reformat_engine.service as reformat_service
    from document_ir import ParagraphBlock as PB

    a = tmp_path / "a.docx"
    out = tmp_path / "assembled.docx"
    _make_doc(a, ["普通段落"])
    real_annotate = reformat_service.annotate_document

    def fake_annotate(ir):
        real_annotate(ir)
        for block in ir.blocks:
            if isinstance(block, PB) and block.text.strip():
                block.semantic_role = "unknown"
                block.role_confidence = 0.1
                break
        return ir

    monkeypatch.setattr(reformat_service, "annotate_document", fake_annotate)
    res = assemble_documents([str(a)], str(out), explicit_format_hint="正式公文")
    assert res["status"] == "ok"
    assert res["items"][0]["status"] == "success_with_warnings"
    assert any("SOURCE_FORMAT_PRESERVED_FOR_UNRESOLVED_BLOCK" in w for w in res["items"][0]["warnings"])


def test_16_opaque_item_atomic_failure(tmp_path):
    a = tmp_path / "a.docx"
    out = tmp_path / "assembled.docx"
    doc = Document()
    doc.add_paragraph("正常")
    doc.save(a)
    with ZipFile(a) as zin:
        parts = {info.filename: zin.read(info.filename) for info in zin.infolist()}
    docxml = parts["word/document.xml"].decode("utf-8")
    docxml = docxml.replace(
        "<w:sectPr",
        '<w:customXml><w:t>opaque content</w:t></w:customXml><w:sectPr',
        1,
    )
    parts["word/document.xml"] = docxml.encode("utf-8")
    with ZipFile(a, "w", ZIP_DEFLATED) as zout:
        for name, data in parts.items():
            zout.writestr(name, data)
    res = assemble_documents([str(a)], str(out), explicit_format_hint="正式公文")
    assert res["status"] == "ASSEMBLY_INCOMPLETE"
    assert res["failed"] == 1
    assert not out.exists()


def test_17_preservation_failed_item_no_output(tmp_path, monkeypatch):
    import reformat_engine.service as reformat_service

    a = tmp_path / "a.docx"
    b = tmp_path / "b.docx"
    out = tmp_path / "assembled.docx"
    _make_doc(a, ["AAA"])
    _make_doc(b, ["BBB"])
    real_render = reformat_service.render_reformat
    calls = {"n": 0}

    def fake_render(source, plan, output_path):
        calls["n"] += 1
        if calls["n"] == 2:
            return {
                "status": "error",
                "source_path": str(source),
                "output_path": str(output_path),
                "target_profile_id": plan.target_profile_id,
                "operations": {"planned": 0, "applied": 0, "preserved": 0, "deferred": 0},
                "warnings": [],
                "errors": ["CONTENT_PRESERVATION_FAILED"],
                "content_preservation": {"passed": False},
            }
        return real_render(source, plan, output_path)

    monkeypatch.setattr(reformat_service, "render_reformat", fake_render)
    res = assemble_documents([str(a), str(b)], str(out), explicit_format_hint="正式公文")
    assert res["status"] == "ASSEMBLY_INCOMPLETE"
    assert res["failed"] == 1
    assert not out.exists()


def test_18_output_mode_assembled(tmp_path):
    a = tmp_path / "a.docx"
    b = tmp_path / "b.docx"
    out = tmp_path / "assembled.docx"
    _make_doc(a, ["AAA"])
    _make_doc(b, ["BBB"])
    res = assemble_documents([str(a), str(b)], str(out), explicit_format_hint="正式公文", output_mode="assembled")
    assert res["status"] == "ok"
    assert out.exists()
    assert res["normalized_outputs"] == []


def test_19_output_mode_separate(tmp_path):
    a = tmp_path / "a.docx"
    b = tmp_path / "b.docx"
    out = tmp_path / "separate.docx"
    _make_doc(a, ["AAA"])
    _make_doc(b, ["BBB"])
    res = assemble_documents([str(a), str(b)], str(out), explicit_format_hint="正式公文", output_mode="separate")
    assert res["status"] == "ok"
    assert not out.exists()
    assert len(res["normalized_outputs"]) == 2
    assert all(Path(p).exists() for p in res["normalized_outputs"])


def test_20_output_mode_both(tmp_path):
    a = tmp_path / "a.docx"
    b = tmp_path / "b.docx"
    out = tmp_path / "both.docx"
    _make_doc(a, ["AAA"])
    _make_doc(b, ["BBB"])
    res = assemble_documents([str(a), str(b)], str(out), explicit_format_hint="正式公文", output_mode="both")
    assert res["status"] == "ok"
    assert out.exists()
    assert len(res["normalized_outputs"]) == 2
    assert all(Path(p).exists() for p in res["normalized_outputs"])


def test_21_payload_fingerprint_order_sensitive(tmp_path):
    a = tmp_path / "a.docx"
    b = tmp_path / "b.docx"
    out1 = tmp_path / "ab.docx"
    out2 = tmp_path / "ba.docx"
    _make_doc(a, ["AAA"])
    _make_doc(b, ["BBB"])
    res1 = assemble_documents([str(a), str(b)], str(out1), explicit_format_hint="正式公文")
    res2 = assemble_documents([str(b), str(a)], str(out2), explicit_format_hint="正式公文")
    assert res1["assembly_payload_sha256"] != res2["assembly_payload_sha256"]


def test_22_payload_fingerprint_stable(tmp_path):
    a = tmp_path / "a.docx"
    b = tmp_path / "b.docx"
    out1 = tmp_path / "ab1.docx"
    out2 = tmp_path / "ab2.docx"
    _make_doc(a, ["AAA"])
    _make_doc(b, ["BBB"])
    res1 = assemble_documents([str(a), str(b)], str(out1), explicit_format_hint="正式公文")
    res2 = assemble_documents([str(a), str(b)], str(out2), explicit_format_hint="正式公文")
    assert res1["assembly_payload_sha256"] == res2["assembly_payload_sha256"]


def test_23_temp_cleanup(tmp_path, monkeypatch):
    import tempfile as tempfile_module

    real_mkdtemp = tempfile_module.mkdtemp
    created = []

    def fake_mkdtemp(prefix=None, **kwargs):
        path = tmp_path / (prefix + "test")
        path.mkdir()
        created.append(path)
        return str(path)

    monkeypatch.setattr(tempfile_module, "mkdtemp", fake_mkdtemp)
    a = tmp_path / "a.docx"
    out = tmp_path / "assembled.docx"
    _make_doc(a, ["AAA"])
    assemble_documents([str(a)], str(out), explicit_format_hint="正式公文")
    assert out.exists()
    for path in created:
        assert not path.exists()

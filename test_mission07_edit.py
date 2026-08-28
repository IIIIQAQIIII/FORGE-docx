"""Mission 07 — Inspect + Source-Preserving Edit Engine tests."""

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from document_ir import ParagraphBlock, TableBlock, read_docx
from edit_engine.service import edit_document, inspect_document
from edit_engine.validation import payload_from_ir, validate_edit_contract


def _sha(path):
    import hashlib

    return hashlib.sha256(path.read_bytes()).hexdigest()


def _make_doc(path, texts):
    doc = Document()
    for text in texts:
        doc.add_paragraph(text)
    doc.save(path)


def _edits_from_ir(source, ops):
    ir = read_docx(source)
    return ir


# ---------------------------------------------------------------- inspect


def test_inspect_pagination(tmp_path):
    src = tmp_path / "src.docx"
    _make_doc(src, [f"p{i}" for i in range(5)])
    result = inspect_document(src, offset=0, limit=2)
    assert result["total_blocks"] == 5
    assert result["next_offset"] == 2
    assert len(result["blocks"]) == 2
    result2 = inspect_document(src, offset=4, limit=2)
    assert result2["next_offset"] is None
    assert len(result2["blocks"]) == 1


def test_inspect_query(tmp_path):
    src = tmp_path / "src.docx"
    _make_doc(src, ["开头"] * 50 + ["目标段落TARGET"] + ["结尾"] * 10)
    result = inspect_document(src, query="TARGET")
    assert result["total_blocks"] == 1
    assert "TARGET" in result["blocks"][0]["text_preview"]


def test_duplicate_text_locators_distinct(tmp_path):
    src = tmp_path / "src.docx"
    _make_doc(src, ["重复文字", "重复文字"])
    result = inspect_document(src)
    locs = [b["source_locator"] for b in result["blocks"]]
    assert len(set(locs)) == 2
    assert locs[0] != locs[1]


def test_nested_table_cell_paragraph_locator(tmp_path):
    src = tmp_path / "src.docx"
    doc = Document()
    doc.add_paragraph("前")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "cell text"
    doc.add_paragraph("后")
    doc.save(src)
    result = inspect_document(src)
    table_block = next(b for b in result["blocks"] if b["block_type"] == "table")
    cell_p = table_block["table"]["cells"][0]["paragraphs"][0]
    assert "/table/row/0/cell/0/p/0" in cell_p["source_locator"]


def test_stale_source_sha_rejected(tmp_path):
    src = tmp_path / "src.docx"
    _make_doc(src, ["hello"])
    result = edit_document(
        src,
        expected_source_sha256="0" * 64,
        edits=[{"op": "replace_text", "target_block_id": "b0", "old_text": "hello", "new_text": "hi", "expected_occurrences": 1}],
    )
    assert result["status"] == "error"
    assert "SOURCE_CHANGED" in result["errors"][0]


# ---------------------------------------------------------------- replace


def test_exact_replace(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    _make_doc(src, ["hello world"])
    result = edit_document(
        src,
        expected_source_sha256=_sha(src),
        edits=[{"op": "replace_text", "target_block_id": "b0", "old_text": "hello", "new_text": "你好", "expected_occurrences": 1}],
        output_path=out,
    )
    assert result["status"] == "ok"
    assert out.exists()
    texts = [getattr(b, "text", "") for b in read_docx(out).blocks]
    assert texts[0] == "你好 world"


def test_replace_across_split_runs(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("Hel")
    p.add_run("lo")
    p.add_run(" world")
    doc.save(src)
    result = edit_document(
        src,
        expected_source_sha256=_sha(src),
        edits=[{"op": "replace_text", "target_block_id": "b0", "old_text": "Hello", "new_text": "你好", "expected_occurrences": 1}],
        output_path=out,
    )
    assert result["status"] == "ok"
    texts = [getattr(b, "text", "") for b in read_docx(out).blocks]
    assert texts[0] == "你好 world"


def test_occurrence_mismatch(tmp_path):
    src = tmp_path / "src.docx"
    _make_doc(src, ["foo foo"])
    result = edit_document(
        src,
        expected_source_sha256=_sha(src),
        edits=[{"op": "replace_text", "target_block_id": "b0", "old_text": "foo", "new_text": "x", "expected_occurrences": 1}],
        dry_run=True,
    )
    assert any("EDIT_OCCURRENCE_MISMATCH" in b for b in result["blockers"])


def test_hyperlink_boundary_rejected(tmp_path):
    src = tmp_path / "src.docx"
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("AAA")
    rel_id = doc.part.relate_to(
        "https://example.com",
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "BBB"
    run.append(t)
    hyperlink.append(run)
    p._p.append(hyperlink)
    p.add_run("CCC")
    doc.save(src)
    result = edit_document(
        src,
        expected_source_sha256=_sha(src),
        edits=[{"op": "replace_text", "target_block_id": "b0", "old_text": "AAABBB", "new_text": "X", "expected_occurrences": 1}],
        dry_run=True,
    )
    assert any("EDIT_COMPLEX_BOUNDARY_UNSUPPORTED" in b for b in result["blockers"])


def test_image_boundary_rejected(tmp_path):
    import struct
    import zlib

    def png():
        def chunk(tag, data):
            c = tag + data
            return struct.pack(">I", len(data)) + c + struct.pack(">I", zlib.crc32(c))

        ihdr = struct.pack(">IIBBBBB", 8, 8, 8, 2, 0, 0, 0)
        idat = zlib.compress(b"\x00" + b"\xAA\xBB\xCC" * 8)
        return b"\x89PNG\r\n\x1a\n" + chunk(b"IHDR", ihdr) + chunk(b"IDAT", idat) + chunk(b"IEND", b"")

    src = tmp_path / "src.docx"
    img = tmp_path / "img.png"
    img.write_bytes(png())
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("AAA")
    p.add_run().add_picture(str(img), width=Cm(2))
    p.add_run("BBB")
    doc.save(src)
    result = edit_document(
        src,
        expected_source_sha256=_sha(src),
        edits=[{"op": "replace_text", "target_block_id": "b0", "old_text": "AAABBB", "new_text": "X", "expected_occurrences": 1}],
        dry_run=True,
    )
    assert any("EDIT_COMPLEX_BOUNDARY_UNSUPPORTED" in b for b in result["blockers"])


# ---------------------------------------------------------------- insert/delete


def test_insert_and_append(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    _make_doc(src, ["B"])
    sha = _sha(src)
    result = edit_document(
        src,
        expected_source_sha256=sha,
        edits=[
            {"op": "insert_paragraph_before", "target_block_id": "b0", "text": "A"},
            {"op": "insert_paragraph_after", "target_block_id": "b0", "text": "C"},
            {"op": "append_paragraph", "text": "D"},
        ],
        output_path=out,
    )
    assert result["status"] == "ok"
    texts = [getattr(b, "text", "") for b in read_docx(out).blocks]
    assert texts == ["A", "B", "C", "D"]


def test_safe_delete(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    _make_doc(src, ["A", "B"])
    result = edit_document(
        src,
        expected_source_sha256=_sha(src),
        edits=[{"op": "delete_paragraph", "target_block_id": "b0"}],
        output_path=out,
    )
    assert result["status"] == "ok"
    texts = [getattr(b, "text", "") for b in read_docx(out).blocks]
    assert texts == ["B"]


def test_unsafe_delete_rejected(tmp_path):
    src = tmp_path / "src.docx"
    doc = Document()
    p = doc.add_paragraph("文本")
    p.add_run().add_picture(str(_make_png(tmp_path)), width=Cm(2))
    doc.save(src)
    result = edit_document(
        src,
        expected_source_sha256=_sha(src),
        edits=[{"op": "delete_paragraph", "target_block_id": "b0"}],
        dry_run=True,
    )
    assert any("EDIT_UNSAFE_DELETE_TARGET" in b for b in result["blockers"])


def _make_png(tmp_path):
    import struct
    import zlib

    def chunk(tag, data):
        c = tag + data
        return struct.pack(">I", len(data)) + c + struct.pack(">I", zlib.crc32(c))

    ihdr = struct.pack(">IIBBBBB", 8, 8, 8, 2, 0, 0, 0)
    idat = zlib.compress(b"\x00" + b"\xAA\xBB\xCC" * 8)
    data = b"\x89PNG\r\n\x1a\n" + chunk(b"IHDR", ihdr) + chunk(b"IDAT", idat) + chunk(b"IEND", b"")
    path = tmp_path / "img.png"
    path.write_bytes(data)
    return path


def test_conflicting_edits_rejected(tmp_path):
    src = tmp_path / "src.docx"
    _make_doc(src, ["A", "B"])
    result = edit_document(
        src,
        expected_source_sha256=_sha(src),
        edits=[
            {"op": "delete_paragraph", "target_block_id": "b0"},
            {"op": "replace_text", "target_block_id": "b0", "old_text": "A", "new_text": "X", "expected_occurrences": 1},
        ],
        dry_run=True,
    )
    assert any("EDIT_OPERATION_CONFLICT" in b for b in result["blockers"])


def test_dry_run_no_output(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    _make_doc(src, ["A"])
    result = edit_document(
        src,
        expected_source_sha256=_sha(src),
        edits=[{"op": "replace_text", "target_block_id": "b0", "old_text": "A", "new_text": "B", "expected_occurrences": 1}],
        output_path=out,
        dry_run=True,
    )
    assert result["status"] == "dry_run"
    assert not out.exists()


def test_edit_source_sha_unchanged_and_contract_pass(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    _make_doc(src, ["旧文字", "其他不变"])
    before = src.read_bytes()
    result = edit_document(
        src,
        expected_source_sha256=_sha(src),
        edits=[{"op": "replace_text", "target_block_id": "b0", "old_text": "旧文字", "new_text": "新文字", "expected_occurrences": 1}],
        output_path=out,
    )
    assert result["status"] == "ok"
    assert result["preservation"]["passed"] is True
    assert src.read_bytes() == before


def test_unauthorized_mutation_contract_violation(tmp_path, monkeypatch):
    import edit_engine.service as service

    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    _make_doc(src, ["旧文字"])
    real_render = service.render_edit

    def rogue_render(source_path, output_path, plan):
        result = real_render(source_path, output_path, plan)
        # unauthorized extra paragraph after render
        d = Document(output_path)
        d.add_paragraph("未授权的新内容")
        d.save(output_path)
        return result

    monkeypatch.setattr(service, "render_edit", rogue_render)
    result = edit_document(
        src,
        expected_source_sha256=_sha(src),
        edits=[{"op": "replace_text", "target_block_id": "b0", "old_text": "旧文字", "new_text": "新文字", "expected_occurrences": 1}],
        output_path=out,
    )
    assert result["status"] == "error"
    assert "EDIT_CONTRACT_VIOLATION" in result["errors"][0]
    assert not out.exists()


def test_table_media_unchanged_by_edit(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    png = _make_png(tmp_path)
    doc = Document()
    doc.add_paragraph("要改的段落")
    doc.add_picture(str(png), width=Cm(2))
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "表格A"
    table.rows[0].cells[1].text = "表格B"
    doc.save(src)
    result = edit_document(
        src,
        expected_source_sha256=_sha(src),
        edits=[{"op": "replace_text", "target_block_id": "b0", "old_text": "要改的段落", "new_text": "已改段落", "expected_occurrences": 1}],
        output_path=out,
    )
    assert result["status"] == "ok"
    src_ir = read_docx(src)
    out_ir = read_docx(out)
    assert [m.sha256 for m in src_ir.media] == [m.sha256 for m in out_ir.media]
    from reformat_engine.renderer import _table_structure_signature

    assert _table_structure_signature(src_ir) == _table_structure_signature(out_ir)


def test_edit_then_reformat_e2e(tmp_path, monkeypatch):
    monkeypatch.setenv("FORGE_HOME", str(tmp_path / "forge_home"))
    from open_format.profile_store import save_user_profile

    save_user_profile("edit_profile", "编辑后格式", "", "custom", "generic_document", {"body": {"font": "黑体"}})

    src = tmp_path / "src.docx"
    edited = tmp_path / "edited.docx"
    final = tmp_path / "final.docx"
    _make_doc(src, ["旧标题文字", "正文内容"])
    edit_result = edit_document(
        src,
        expected_source_sha256=_sha(src),
        edits=[{"op": "replace_text", "target_block_id": "b0", "old_text": "旧标题文字", "new_text": "新标题文字", "expected_occurrences": 1}],
        output_path=edited,
    )
    assert edit_result["status"] == "ok"

    from reformat_engine.service import reformat_document

    ref_result = reformat_document(edited, output_path=final, saved_profile_id="edit_profile")
    assert ref_result["status"] == "ok"
    texts = [getattr(b, "text", "") for b in read_docx(final).blocks]
    assert "新标题文字" in texts
    assert "正文内容" in texts

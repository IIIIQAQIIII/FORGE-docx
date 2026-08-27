"""A local MCP server for template-based Word documents.

Tools:
- list_document_types   : available friendly document types and their templates
- get_template_schema   : the JSON fields a template expects, plus an example
- generate_docx         : fill a template and save the finished .docx
- generate_by_type      : like generate_docx but uses a friendly type name
- validate_docx         : basic content / layout / leftover-placeholder checks
- fix_docx              : safe light repairs (NBSP and accidental trailing spaces)
"""

from __future__ import annotations

import json
import math
import re
from pathlib import Path
from typing import Any
from zipfile import ZipFile

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docxtpl import DocxTemplate
from mcp.server.fastmcp import FastMCP

PROJECT_DIR = Path(__file__).resolve().parent
TEMPLATES_DIR = PROJECT_DIR / "templates"
OUTPUTS_DIR = PROJECT_DIR / "outputs"

for directory in (TEMPLATES_DIR, OUTPUTS_DIR):
    directory.mkdir(exist_ok=True)

mcp = FastMCP("FORGE Document Engine")

# ---------------------------------------------------------------------------
# Document type registry
# ---------------------------------------------------------------------------

# 类型注册表：同一格式可以有多个入口名，便于 AI 按用户意图直接选择。
DOCUMENT_TYPES = {
    # 标准传统公文（计划/总结/方案/汇报/报告/请示等一般公文）
    "传统公文": "传统公文.docx",
    "传统公文通用": "传统公文.docx",
    "计划": "传统公文.docx",
    "总结": "传统公文.docx",
    "方案": "传统公文.docx",
    "汇报": "传统公文.docx",
    "报告": "传统公文.docx",
    "请示": "传统公文.docx",
    # 活动方案专用（落款含部门）
    "传统公文活动方案": "传统公文-活动方案.docx",
    "活动方案": "传统公文-活动方案.docx",
    # 论文格式也适用于其他文字量大的长文
    "论文": "论文.docx",
    "演讲稿": "论文.docx",
    "发言稿": "论文.docx",
    "长文": "论文.docx",
    # 行政周报
    "行政周报": "行政周报.docx",
    "周报": "行政周报.docx",
    # 活动方案套装部件（方案/总结/影像）
    "活动总结": "活动总结.docx",
    "活动影像": "活动影像.docx",
    # 培训资料套装
    "培训通知": "培训通知.docx",
    "培训活动记录": "培训活动记录.docx",
    "培训活动影像": "培训活动影像.docx",
    # 完整三合一版本（保留兼容）
    "培训通知记录": "培训通知记录.docx",
    "通知培训资料": "培训通知记录.docx",
    "sample": "sample_template.docx",
}

# 格式判断指南：AI 根据用户想要生成的内容自行选择格式。
FORMAT_GUIDE = [
    {
        "when": "用户要生成计划、总结、方案、汇报、报告、请示等一般公文（例如“评价计划”“学期总结”“工作汇报”）",
        "use": "传统公文（generate_by_type 的 document_type 可填：传统公文/计划/总结/方案/汇报/报告/请示）",
        "because": "标准公文格式：标题居中、正文仿宋28磅、三级层次、落款中线靠右；无红色红头。注意：只要文种是“计划/总结/方案/汇报/报告”等，就用传统公文，不要因其他修饰词误判为论文。",
    },
    {
        "when": "用户要生成活动方案（落款需要单位+部门+日期三行）",
        "use": "传统公文活动方案（或 document_type=活动方案）",
        "because": "活动方案专用模板，落款三行共享中线、整体靠右。",
    },
    {
        "when": "用户要生成论文、演讲稿、发言稿等文字量大的长文",
        "use": "论文（或 document_type=演讲稿/发言稿/长文）",
        "because": "论文格式：无页眉、标题不加粗、正文小四宋体18磅、可插三线表/图片；演讲稿可省略摘要和关键词。",
    },
    {
        "when": "用户要生成行政周报、周报",
        "use": "行政周报（或 document_type=周报）",
        "because": "周报格式：两行标题（单位+学期+行政周报 / 第X周（日期）），按部门分节，每段自动编号，空缺写“无”。",
    },
    {
        "when": "用户要生成活动方案、活动总结、活动影像，或一套完整活动资料（方案+总结+影像）",
        "use": "活动方案套装（generate_document_set(\"活动方案套装\")，或单独选择活动方案/活动总结/活动影像）",
        "because": "活动方案套装 = 方案 + 总结 + 影像；方案/总结共用通用活动文档格式，影像为标题+信息表+两张照片页。",
    },
    {
        "when": "用户要生成培训通知、培训活动记录、培训活动影像，或一套完整培训资料",
        "use": "培训资料套装三件（generate_document_set）或单独选择培训通知/培训活动记录/培训活动影像",
        "because": "培训通知含红色红头+双红线；活动记录为表格；活动影像为照片页。",
    },
]

# 类型说明（避免混淆：红色红头 + 双红线只属于通知类模板）。
DOCUMENT_TYPE_DESCRIPTIONS = {
    "传统公文": "标准传统公文格式：标题由单位名称+事由文种两行组成，黑色2号方正小标宋居中；无红色红头、无双红线。",
    "传统公文通用": "同“传统公文”，标准传统公文格式，无红色红头。",
    "计划": "计划 → 用标准传统公文格式（无红色红头）。",
    "总结": "总结 → 用标准传统公文格式（无红色红头）。",
    "方案": "方案 → 用标准传统公文格式（无红色红头）。",
    "汇报": "汇报 → 用标准传统公文格式（无红色红头）。",
    "报告": "报告 → 用标准传统公文格式（无红色红头）。",
    "请示": "请示 → 用标准传统公文格式（无红色红头）。",
    "传统公文活动方案": "活动方案专用模板（已脱敏）：标题居中、正文三层结构、落款单位+部门+日期三行居中靠右；无红色红头。",
    "活动方案": "同“传统公文活动方案”，落款含部门三行。",
    "论文": "论文版式：无页眉，标题与各级标题不加粗，正文小四宋体18磅，可插入三线表/图片。",
    "演讲稿": "演讲稿 → 用论文格式（可省略摘要、关键词）。",
    "发言稿": "发言稿 → 用论文格式（可省略摘要、关键词）。",
    "长文": "文字量大的长文 → 用论文格式。",
    "行政周报": "行政周报格式：两行标题（单位+学期+行政周报 / 第X周（日期）），按部门分节，条目自动编号，空缺写“无”。",
    "周报": "同“行政周报”。",
    "活动总结": "活动方案套装第2件：活动总结（标题+副标题+开头+章节+落款单位/部门/日期）。",
    "活动影像": "活动方案套装第3件：活动影像（标题+信息表：时间/活动地点/负责人/活动内容+两张照片同一页）。",
    "培训通知": "培训资料套装第1件：仅通知，含红色红头 + 双红线（VML 红线）。",
    "培训活动记录": "培训资料套装第2件：仅活动记录表（不含影像页）。",
    "培训活动影像": "培训资料套装第3件：仅活动影像页，两张照片同一页。",
    "培训通知记录": "完整三合一版本：通知（含红色红头+双红线）+ 活动记录 + 活动影像，三部分各另起新页。",
    "通知培训资料": "同“培训通知记录”，完整三合一版本。",
    "sample": "验证用最小模板。",
}

# 一套培训资料由三个模板组成，按此顺序生成/装订。
DOCUMENT_SETS: dict[str, dict[str, Any]] = {
    "培训资料套装": {
        "description": "一套完整的培训资料：培训通知 → 培训活动记录 → 培训活动影像。",
        "templates": ["培训通知.docx", "培训活动记录.docx", "培训活动影像.docx"],
    },
    "活动方案套装": {
        "description": "一套完整的活动资料：活动方案 → 活动总结 → 活动影像。",
        "templates": ["传统公文-活动方案.docx", "活动总结.docx", "活动影像.docx"],
    },
}

# 学年度/学期规律：所有落款时间必须落在对应学期内，且取工作日。
# 一学期默认 18 周，均为工作日；分上下两学期。
# 例：2025——2026学年度
#   第一学期：2025年9月～2026年1月
#   第二学期：2026年2月～2026年7月
SEMESTER_RULES: dict[str, Any] = {
    "weeks_per_semester": 18,
    "weeks_are_workdays": True,
    "first_semester_months": (9, 1),   # 同年9月～次年1月
    "second_semester_months": (2, 7),  # 次年2月～次年7月
    "academic_year_template": "{start_year}——{end_year}学年度",
}

# Human-readable placeholder descriptions and example payloads.
# The placeholder lists themselves are extracted from the real docx files,
# so this table only supplies the meaning and a ready-to-use example.
TEMPLATE_INFO: dict[str, dict[str, Any]] = {
    "sample_template.docx": {
        "notes": ["验证用最小模板，用于确认 MCP 服务与 docxtpl 渲染正常。"],
        "placeholders": {
            "title": "文档标题",
            "body": "正文内容",
            "signature": "落款署名",
            "date": "落款日期",
        },
        "example": {
            "title": "测试文档",
            "body": "这是一段用于验证 MCP 服务的正文内容。",
            "signature": "测试单位",
            "date": "2026年8月26日",
        },
    },
    "传统公文.docx": {
        "notes": [
            "传统公文格式（无红色红头、无双红线）：标题由发文单位名称 + 事由文种组成，两行均为2号方正小标宋简体、居中、固定32磅行距。",
            "3号仿宋正文（固定28磅）、黑体/楷体/仿宋三级层次；sections 支持三层嵌套：一、→（一）→ 1.，没有的层级直接省略对应键即可。",
            "subtitle、attachment 可选；不填时对应整段会自动消失（公文不含附注）。",
            "页码为4号宋体阿拉伯数字，左右各一条一字线，单页右空一字、双页左空一字。",
            "红色红头与双红线只用于通知类模板（培训通知/培训通知记录）。",
        ],
        "placeholders": {
            "organization": "发文单位全称（标题第一行，黑色2号方正小标宋）",
            "title": "公文标题（如：科技之春活动方案）",
            "subtitle": "标题下方备注/副标题，可省略",
            "opening": "正文开头段（如：为进一步……现制定方案如下：）",
            "section.heading": "一级标题（一、XXX），黑体",
            "section.paragraphs": "一级标题下的正文段落数组",
            "section.subsections": "二级层次数组（（一）XXX），楷体",
            "subsection.heading": "二级标题",
            "subsection.paragraphs": "二级标题下的正文段落数组",
            "subsection.items": "三级层次数组（1.XXX），仿宋",
            "item.heading": "三级标题",
            "item.paragraphs": "三级标题下的正文段落数组",
            "attachment": "附件名称，可省略",
            "signature": "发文单位署名",
            "date": "成文日期",
        },
        "example": {
            "organization": "示例幼儿园",
            "title": "科技之春活动方案",
            "subtitle": "",
            "opening": "为进一步激发幼儿对科学技术的兴趣，结合园所实际，拟于近期开展“科技之春”系列活动。现制定方案如下：",
            "sections": [
                {
                    "heading": "一、活动主题",
                    "paragraphs": ["科技点亮童心 探索创造未来"],
                },
                {
                    "heading": "二、活动安排",
                    "subsections": [
                        {
                            "heading": "（一）活动时间",
                            "paragraphs": ["拟定于下周至下下周分园开展。"],
                            "items": [
                                {"heading": "1.具体安排", "paragraphs": ["由三幼、八幼根据园所实际统筹安排。"]},
                            ],
                        }
                    ],
                },
            ],
            "attachment": "活动安排表",
            "signature": "示例幼儿园",
            "date": "2026年3月19日",
        },
    },
    "传统公文-活动方案.docx": {
        "notes": [
            "活动方案模板（已脱敏，可自由填写）：标题块为2号方正小标宋居中；正文3号仿宋固定28磅，一、黑体/（一）楷体/1.仿宋三层。",
            "落款为单位+部门+日期三行，共享中线、整体靠右。",
            "sections 支持三层嵌套；subtitle、attachment 可选（公文不含附注）。",
        ],
        "placeholders": {
            "organization": "发文单位全称（标题第一行）",
            "title": "活动方案标题",
            "subtitle": "副标题，可省略",
            "opening": "开头段",
            "section.heading": "一级标题（一、XXX），黑体",
            "section.paragraphs": "一级标题下的正文段落数组",
            "section.subsections": "二级层次数组（（一）XXX），楷体",
            "subsection.heading": "二级标题",
            "subsection.paragraphs": "二级标题下的正文段落数组",
            "subsection.items": "三级层次数组（1.XXX），仿宋",
            "item.heading": "三级标题",
            "item.paragraphs": "三级标题下的正文段落数组",
            "attachment": "附件名称，可省略",
            "signature": "落款单位",
            "department": "落款部门",
            "date": "成文日期",
        },
        "example": {
            "organization": "某幼儿园",
            "title": "某活动方案",
            "subtitle": "",
            "opening": "为开展某活动，现制定方案如下：",
            "sections": [
                {"heading": "一、活动主题", "paragraphs": ["某主题"]},
                {"heading": "二、活动安排", "subsections": [
                    {"heading": "（一）活动时间", "paragraphs": ["某时间"], "items": [
                        {"heading": "1.具体安排", "paragraphs": ["某安排"]}
                    ]},
                ]},
            ],
            "attachment": "",
            "signature": "某幼儿园",
            "department": "某部门",
            "date": "2026年3月2日",
        },
    },
    "论文.docx": {
        "optional": ["author", "abstract", "keywords"],
        "notes": [
            "论文/长文模板：A4，左边距30mm、右边距25mm、上边距30mm、下边距25mm。",
            "一级标题黑体16pt、二级黑体14pt、三级黑体12pt，段前段后1行；正文小四宋体、行距18磅。",
            "论文模板已按当前要求删除页眉文字和横线；页脚保留居中页码。",
            "如提供 tables/figures，会自动插入三线表与居中图片；图、表标题五号黑体，表格文字小五宋体。",
        ],
        "placeholders": {
            "title": "论文题目",
            "author": "作者姓名",
            "abstract": "摘要内容",
            "keywords": "关键词（分号分隔）",
            "chapter.heading": "一级标题（第一章 XXX）",
            "chapter.sections": "二级标题数组",
            "section.heading": "二级标题（1.1 XXX）",
            "section.subsections": "三级标题数组",
            "subsection.heading": "三级标题（1.1.1 XXX）",
            "subsection.paragraphs": "正文段落数组",
            "subsection.figures": "可选：插入该子节正文之后的图片数组。单图 [{path, caption, width_cm}]；拼图 [{paths: [图1,图2,图3], caption, width_cm, gap_cm}]，自动并排为一张图",
            "tables": "三线表数组：[{caption, headers, rows}]",
            "figures": "图片数组：[{path, caption, width_cm}]",
        },
        "example": {
            "title": "幼儿园科学教育活动设计研究",
            "author": "张三",
            "abstract": "本文围绕幼儿园科学教育活动的设计与实施展开研究……",
            "keywords": "科学教育；幼儿园；活动设计",
            "chapters": [
                {
                    "heading": "第一章 绪论",
                    "sections": [
                        {
                            "heading": "1.1 研究背景",
                            "subsections": [
                                {
                                    "heading": "1.1.1 问题的提出",
                                    "paragraphs": ["正文段落一。", "正文段落二。"],
                                }
                            ],
                        }
                    ],
                }
            ],
            "tables": [
                {"caption": "表1 研究对象基本情况", "headers": ["班级", "人数", "年龄"], "rows": [["中一班", "30", "4-5岁"]]},
            ],
            "figures": [
                {"path": "outputs/example_figure.png", "caption": "图1 活动流程", "width_cm": 13},
            ],
        },
    },
    "行政周报.docx": {
        "notes": [
            "行政周报：标题两行——第一行“单位+学期+行政周报”，第二行“第X周（起止日期）”。",
            "按部门分节（一、党支部；二、办公室……），每节固定三块：本周工作总结（含亮点）、存在问题及改进措施、下周重点工作计划。",
            "条目由模板自动编号（1. 2. 3. …）；某一块空缺时传 [\"无\"]，生成时只显示“无”且不加序号。",
        ],
        "placeholders": {
            "organization": "单位全称（如：示例幼儿园）",
            "semester": "学期（如：2026年秋季学期）",
            "week": "周次（如：17）",
            "date_range": "本周起止日期（如：12.21-12.25）",
            "section.name": "部门名称（如：一、党支部）",
            "section.summary": "本周工作总结条目数组；空缺写 [\"无\"]",
            "section.problems": "存在问题及改进措施条目数组；空缺写 [\"无\"]",
            "section.plan": "下周重点工作计划条目数组；空缺写 [\"无\"]",
        },
        "example": {
            "organization": "示例幼儿园",
            "semester": "2026年秋季学期",
            "week": "17",
            "date_range": "12.21-12.25",
            "sections": [
                {"name": "一、党支部", "summary": ["持续推进党建重点工作。", "完成意识形态工作报告报送。"], "problems": ["无"], "plan": ["继续做好党建与创建工作推进。", "梳理党支部工作亮点和过程性资料。"]},
                {"name": "二、办公室", "summary": ["参加行政会。", "收集并查看家长问卷调查表填写情况，形成分析报告。"], "problems": ["近期需各部门配合提交资料较多，存在未按时提交、格式错误较多等问题，需多次督促核对。"], "plan": ["敲定总结大会细节并召开学期总结大会。", "撰写园务工作总结。"]},
            ],
        },
    },
    "活动总结.docx": {
        "notes": [
            "活动方案套装第 2 件：活动总结（标题+副标题+开头+章节+落款单位/部门/日期）。",
            "结构可与活动方案一致，AI 根据用户要求撰写总结内容。",
        ],
        "placeholders": {
            "organization": "发文单位全称（标题第一行）",
            "title": "总结标题（如：读书月活动总结）",
            "subtitle": "副标题（如：2025——2026学年第二学期），可省略",
            "opening": "开头段",
            "section.heading": "一级标题（一、XXX），黑体",
            "section.paragraphs": "正文段落数组",
            "section.subsections": "二级层次数组（（一）XXX），楷体",
            "subsection.heading": "二级标题",
            "subsection.paragraphs": "二级标题下的正文段落",
            "subsection.items": "三级层次数组（1.XXX），仿宋",
            "item.heading": "三级标题",
            "item.paragraphs": "三级标题下的正文段落",
            "attachment": "附件，可省略",
            "signature": "落款单位",
            "department": "落款部门",
            "date": "成文日期",
        },
        "example": {
            "organization": "示例幼儿园",
            "title": "读书月活动总结",
            "subtitle": "2025——2026学年第二学期",
            "opening": "为进一步营造浓厚的园所阅读氛围，培养幼儿良好的阅读习惯，我园于2026年4月组织开展了读书月活动，现将活动总结如下：",
            "sections": [
                {"heading": "一、活动开展情况", "paragraphs": ["各班结合年龄特点开展了形式多样的阅读活动。"]},
                {"heading": "二、活动成效与亮点", "paragraphs": ["活动设计体现了年龄梯度，家园协同作用明显。"]},
            ],
            "attachment": "",
            "signature": "示例幼儿园",
            "department": "保教部",
            "date": "2026年4月30日",
        },
    },
    "活动影像.docx": {
        "notes": [
            "活动方案套装第 3 件：活动影像（标题+信息表+两张照片同一页）。",
            "信息表字段：时间/活动地点/负责人/活动内容；照片行保留灰色占位图。",
        ],
        "placeholders": {
            "organization": "单位全称（标题用）",
            "date_short": "影像时间（如：2026.4.29）",
            "location": "活动地点",
            "organizer": "负责人",
            "activity_name": "活动内容",
        },
        "example": {
            "organization": "示例幼儿园",
            "date_short": "2026.4.29",
            "location": "幼儿园多功能厅、各班级",
            "organizer": "保教部",
            "activity_name": "读书月系列活动",
        },
    },
    "培训通知.docx": {
        "notes": [
            "培训资料套装第 1 件：仅培训通知。",
            "含红色红头（小初方正小标宋简体红色）+ 双红线（VML 两条红色横线），这是通知类模板的专属格式。",
            "正文段间距固定 28 磅；通知默认保持一页。如因排版或字数需要超过一页，必须先与操作者确认，确认后用 force_multipage_notice=true 重新生成。",
            "红字单位名称自动适配：字号固定小初 36pt；三幼/八幼保持模板默认参数，其他单位按名称长度自动调整缩放/间距，保证一行美观。",
        ],
        "placeholders": {
            "organization": "单位全称",
            "document_title": "通知标题（如：培训通知）",
            "purpose": "通知开头段",
            "date": "通知落款日期",
            "location": "活动地点",
            "training_topic": "培训主题",
            "trainer": "培训人",
            "participants": "参培人员",
            "requirement_1": "培训要求1",
            "requirement_2": "培训要求2",
            "department": "落款部门",
        },
        "example": {
            "organization": "示例幼儿园",
            "document_title": "培训通知",
            "purpose": "为进一步加强师德师风警示教育……现将有关事项通知如下：",
            "date": "2026年6月30日",
            "location": "党建室",
            "training_topic": "师德师风培训《警示案例学习》",
            "trainer": "陈老师",
            "participants": "示例幼儿园全体教师",
            "requirement_1": "全体教师准时参会，认真做好学习记录。",
            "requirement_2": "结合岗位职责和案例内容主动对照反思，增强底线意识、规矩意识和责任意识。",
            "department": "保教部",
        },
    },
    "培训活动影像.docx": {
        "notes": ["培训资料套装第 3 件：仅活动影像页（4 行照片表，两张照片在同一页）。"],
        "placeholders": {
            "date_short": "培训日期",
            "location": "培训地点",
            "participants": "参培人员",
            "training_topic": "培训主题",
        },
        "example": {
            "date_short": "2026.6.30",
            "location": "党建室",
            "participants": "全体教师",
            "training_topic": "师德师风培训《警示案例学习》",
        },
    },
    "培训通知记录.docx": {
        "notes": [
            "完整三合一版本：通知 + 活动记录 + 活动影像，三部分各另起新页。",
            "仅“通知”部分含红色红头 + 双红线；其余部分为常规表格版式。",
        ],
        "placeholders": {
            "organization": "单位全称",
            "document_title": "通知标题（如：培训通知）",
            "purpose": "通知开头段",
            "date": "通知落款日期",
            "location": "活动地点",
            "training_topic": "培训主题",
            "trainer": "培训人",
            "participants": "参培人员",
            "requirement_1": "培训要求1",
            "requirement_2": "培训要求2",
            "department": "落款部门",
            "date_short": "表格内日期（如：2026.6.30）",
            "hours": "学时",
            "training_content": "培训内容（表格长文本）",
            "training_reflection": "培训心得（表格长文本）",
        },
        "example": {
            "organization": "示例幼儿园",
            "document_title": "培训通知",
            "purpose": "为进一步加强师德师风警示教育……现将有关事项通知如下：",
            "date": "2026年6月30日",
            "location": "党建室",
            "training_topic": "师德师风培训《警示案例学习》",
            "trainer": "陈老师",
            "participants": "示例幼儿园全体教师",
            "requirement_1": "全体教师准时参会，认真做好学习记录。",
            "requirement_2": "结合岗位职责和案例内容主动对照反思，增强底线意识、规矩意识和责任意识。",
            "department": "保教部",
            "date_short": "2026.6.30",
            "hours": "2",
            "training_content": "一、明确警示案例学习的重要意义……",
            "training_reflection": "通过本次培训，我进一步认识到……",
        },
    },
    "培训活动记录.docx": {
        "notes": ["培训资料套装第 2 件：仅活动记录（标题 + 5 行记录表，不含影像页）。"],
        "placeholders": {
            "organization": "单位全称",
            "date_short": "培训日期",
            "location": "培训地点",
            "participants": "参培人员",
            "trainer": "培训人",
            "hours": "学时",
            "training_topic": "培训主题",
            "training_content": "培训内容（表格长文本）",
            "training_reflection": "培训心得（表格长文本）",
        },
        "example": {
            "organization": "示例幼儿园",
            "date_short": "2026.6.30",
            "location": "党建室",
            "participants": "全体教师",
            "trainer": "陈老师",
            "hours": "2",
            "training_topic": "师德师风培训《警示案例学习》",
            "training_content": "一、明确警示案例学习的重要意义……",
            "training_reflection": "通过本次培训，我进一步认识到……",
        },
    },
}


# ---------------------------------------------------------------------------
# Path helpers
# ---------------------------------------------------------------------------

class PathError(ValueError):
    pass


def resolve_path(value: str, default_dir: Path) -> Path:
    """Resolve a template path; relative paths are based on default_dir.

    Templates must stay inside default_dir (no ../ escape).
    """
    path = Path(value).expanduser()
    if not path.is_absolute():
        path = default_dir / path
    path = path.resolve()
    root = default_dir.resolve()
    if not path.is_relative_to(root):
        raise PathError(f"path escapes allowed directory: {path!s} (must stay inside {root})")
    return path


def resolve_output_path(value: str) -> Path:
    """Resolve an output path flexibly.

    - Absolute path (or ~/...) is used as-is, so the user can save anywhere.
    - A bare filename (no directory part) defaults to the project outputs/ dir.
    - A relative path with directories is resolved against the project dir.
    """
    path = Path(value).expanduser()
    if path.is_absolute():
        return path.resolve()
    if path.parent == Path("."):
        return (OUTPUTS_DIR / path).resolve()
    return (PROJECT_DIR / path).resolve()


def resolve_input_docx(value: str) -> Path:
    """Resolve an existing-docx path for validate/fix.

    Absolute paths are used as-is; relative paths default to outputs/.
    """
    path = Path(value).expanduser()
    if path.is_absolute():
        return path.resolve()
    if path.parent == Path("."):
        return (OUTPUTS_DIR / path).resolve()
    return (PROJECT_DIR / path).resolve()


def _error(message: str) -> str:
    return f"Error: {message}"


def _paragraph_line_height_pt(p) -> float:
    """Return exact line height of a paragraph in points (default 28pt)."""
    ppr = p._p.pPr
    if ppr is not None:
        spacing = ppr.find(qn("w:spacing"))
        if spacing is not None:
            line = spacing.get(qn("w:line"))
            rule = spacing.get(qn("w:lineRule"))
            if line and rule == "exact":
                return int(line) / 20.0
    return 28.0


def _adjust_notice_masthead(path: Path, organization: str) -> None:
    """培训通知红字：字号固定小初 36pt 不变，只按单位名称长度调整缩放/间距。

    示例幼儿园 / 示例幼儿园 保持模板默认参数；
    其他单位保持字号不变，自动调整字符缩放（w:w）与紧缩值（w:spacing），
    保证红字一行放下且美观。
    """
    if not organization or organization in ("示例幼儿园", "示例幼儿园"):
        return
    document = Document(path)
    if not document.paragraphs:
        return
    mast = document.paragraphs[0]
    n = max(len(organization), 1)
    # 版心宽 15.6cm；字号固定 36pt；默认缩放 74%、紧缩 2.3 磅。
    available_pt = (15.6 / 2.54) * 72
    font_pt = 36.0
    default_condense_pt = 2.3

    # 先只调缩放：总宽 = n * (字号*缩放 - 紧缩)
    scale = (available_pt * 0.95 / n + default_condense_pt) / font_pt * 100
    scale = max(50.0, min(74.0, scale))
    condense_twips = -46  # 默认紧缩 2.3 磅
    # 如果缩放已到 50% 仍放不下，再加大紧缩。
    if scale <= 50.0:
        extra = available_pt * 0.95 / n - font_pt * 0.50
        condense_twips = int(round(extra * 20))  # 负数，单位 twips
        condense_twips = max(condense_twips, -120)  # 紧缩最多 6 磅

    for run in mast.runs:
        rpr = run._element.get_or_add_rPr()
        width = rpr.find(qn("w:w"))
        if width is None:
            width = OxmlElement("w:w")
            rpr.append(width)
        width.set(qn("w:val"), str(int(round(scale))))
        spacing = rpr.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            rpr.append(spacing)
        spacing.set(qn("w:val"), str(condense_twips))
    document.save(path)


def _notice_fits_one_page(path: Path) -> tuple[bool, float, float]:
    """Estimate whether the notice part fits on one page.

    Returns (fits, used_points, available_points). For the full training set,
    the notice part ends at the first page break.
    """
    document = Document(path)
    section = document.sections[0]
    text_height_cm = section.page_height.cm - section.top_margin.cm - section.bottom_margin.cm
    available_pt = text_height_cm * 28.3465
    used_pt = 0.0
    for paragraph in document.paragraphs:
        # 完整版中，第一个分页符之后就是活动记录部分，不再计入通知页。
        if paragraph._p.xpath('.//w:br[@w:type="page"]'):
            break
        text = paragraph.text.strip()
        lines = max(1, math.ceil(len(text) / 28)) if text else 1
        used_pt += lines * _paragraph_line_height_pt(paragraph)
    return used_pt <= available_pt, used_pt, available_pt


# ---------------------------------------------------------------------------
# Placeholder extraction
# ---------------------------------------------------------------------------

_PLACEHOLDER_RE = re.compile(r"\{\{([^{}]+)\}\}")
_LOOP_RE = re.compile(r"\{%p\s*for\s+([A-Za-z_]\w*)\s+in\s+([^%]+?)\s*%\}")


def extract_placeholders(template_path: Path) -> list[str]:
    """Extract docxtpl placeholders from document.xml and headers/footers.

    Only real JSON keys are returned: ``{{ expr }}`` placeholders plus top-level
    loop collection names (for example ``sections``). Bare loop variables such as
    ``paragraph`` are filtered out.
    """
    found: list[str] = []
    loop_vars: set[str] = set()
    with ZipFile(template_path) as archive:
        xml_parts = ["word/document.xml"] + [
            name
            for name in archive.namelist()
            if re.match(r"word/(header|footer)\d*\.xml", name)
        ]
        for part in xml_parts:
            try:
                xml = archive.read(part).decode("utf-8")
            except KeyError:
                continue
            for match in _PLACEHOLDER_RE.finditer(xml):
                expr = match.group(1).strip()
                if expr and expr not in found:
                    found.append(expr)
            for match in _LOOP_RE.finditer(xml):
                var = match.group(1).strip()
                seq = match.group(2).strip()
                loop_vars.add(var)
                if re.fullmatch(r"[A-Za-z_]\w*", seq) and seq not in found:
                    found.append(seq)
    return [field for field in found if field not in loop_vars]


# ---------------------------------------------------------------------------
# Tools
# ---------------------------------------------------------------------------

@mcp.tool()
def list_document_types() -> str:
    """List available document types, their template files, and the fields they expect."""
    templates_on_disk = sorted(p.name for p in TEMPLATES_DIR.glob("*.docx"))
    report: dict[str, Any] = {
        "document_types": {
            type_name: {
                "template": template_file,
                "description": DOCUMENT_TYPE_DESCRIPTIONS.get(type_name, ""),
                "fields": extract_placeholders(TEMPLATES_DIR / template_file),
            }
            for type_name, template_file in DOCUMENT_TYPES.items()
        },
        "document_sets": {
            set_name: {
                "description": set_info["description"],
                "templates": set_info["templates"],
            }
            for set_name, set_info in DOCUMENT_SETS.items()
        },
        "templates_on_disk": templates_on_disk,
        "format_guide": FORMAT_GUIDE,
        "semester_rule": {
            "weeks_per_semester": SEMESTER_RULES["weeks_per_semester"],
            "weeks_are_workdays": SEMESTER_RULES["weeks_are_workdays"],
            "first_semester": "同年9月～次年1月",
            "second_semester": "次年2月～次年7月",
            "academic_year_notation": "2025——2026学年度",
            "note": "所有落款时间必须落在对应学期时间段内，且取工作日；具体学期可用 get_semester_info(起始年份) 计算。",
        },
    }
    return json.dumps(report, ensure_ascii=False, indent=2)


@mcp.tool()
def get_template_schema(template_name: str) -> str:
    """Return the JSON schema (fields + descriptions + example) for one template.

    template_name is a file in templates/ (for example, 论文.docx or 传统公文.docx).
    """
    try:
        template_path = resolve_path(template_name, TEMPLATES_DIR)
    except PathError as exc:
        return _error(str(exc))
    if not template_path.is_file():
        return _error(f"template not found: {template_path}")

    fields = extract_placeholders(template_path)
    info = TEMPLATE_INFO.get(template_path.name, {})
    descriptions = info.get("placeholders", {})
    # Some capabilities (for example 论文.docx tables/figures) are inserted by
    # the server rather than present as docxtpl placeholders; document them too.
    for field in descriptions:
        if field not in fields:
            fields.append(field)
    optional = set(info.get("optional", [])) | {"subtitle", "attachment", "annotation", "tables", "figures"}
    schema = {
        "template": template_path.name,
        "document_types": [name for name, tpl in DOCUMENT_TYPES.items() if tpl == template_path.name],
        "fields": {
            field: {
                "description": descriptions.get(field, "（该模板中的变量）"),
                "required": ("." not in field and field not in optional),
            }
            for field in fields
        },
        "example": info.get("example", {}),
        "notes": info.get("notes", []),
    }
    return json.dumps(schema, ensure_ascii=False, indent=2)


@mcp.tool()
def generate_docx(template_name: str, content: dict[str, Any], output_name: str, force_multipage_notice: bool = False) -> str:
    """Fill a DOCX template with structured JSON content and save the Word file.

    template_name is a file in templates/ (for example, sample_template.docx).
    content keys must match the {{ variables }} in that template.
    output_name must end in .docx. It can be an absolute path (save anywhere),
    a relative path with directories (resolved against the project), or a bare
    filename (saved in outputs/).
    Use get_template_schema(template_name) to see required fields and an example.

    培训通知默认保持一页：如果内容预计会超过一页，默认会拒绝生成并提示；
    只有在操作者明确确认允许超过一页时，才传 force_multipage_notice=true 重新生成。
    """
    if not output_name.endswith(".docx"):
        return _error("output_name must end in .docx")

    try:
        template_path = resolve_path(template_name, TEMPLATES_DIR)
    except PathError as exc:
        return _error(str(exc))
    output_path = resolve_output_path(output_name)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    if not template_path.is_file():
        return _error(f"template not found: {template_path}")

    try:
        document = DocxTemplate(template_path)
        document.render(dict(content))
        document.save(output_path)
        if template_path.name == "论文.docx":
            try:
                _insert_thesis_media(output_path, content)
            except Exception:
                output_path.unlink(missing_ok=True)
                raise
        if template_path.name in ("培训通知.docx", "培训通知记录.docx"):
            _adjust_notice_masthead(output_path, str(content.get("organization", "")))
            fits, used_pt, available_pt = _notice_fits_one_page(output_path)
            if not fits and not force_multipage_notice:
                output_path.unlink(missing_ok=True)
                return _error(
                    f"通知内容预计超过一页（估计需 {used_pt:.0f}pt，可用 {available_pt:.0f}pt）。"
                    "通知默认保持一页：请先精简正文；如操作者确认允许超过一页，"
                    "请用 force_multipage_notice=true 重新生成。"
                )
    except Exception as exc:  # return a useful tool error to the client
        return _error(f"generating document: {exc}")

    return f"Created: {output_path}"


@mcp.tool()
def generate_by_type(document_type: str, content: dict[str, Any], output_name: str, force_multipage_notice: bool = False) -> str:
    """Generate a document by friendly type name (for example, 论文 or 传统公文).

    Call list_document_types() first to see every available type.
    output_name may be an absolute path, a relative path, or a bare filename (outputs/).
    培训通知默认保持一页；确认允许超过一页时传 force_multipage_notice=true。
    如果用户未明确说明是整套还是单个，请先与用户确认后再调用本工具。
    """
    template_name = DOCUMENT_TYPES.get(document_type)
    if not template_name:
        available = "、".join(sorted(set(DOCUMENT_TYPES)))
        return _error(f"unknown document type {document_type!r}. Available: {available}")
    return generate_docx(template_name, content, output_name, force_multipage_notice)


@mcp.tool()
def generate_document_set(set_name: str, content: dict[str, Any], output_prefix: str, force_multipage_notice: bool = False) -> str:
    """Generate every document in a named set with the same content.

    set_name is one of list_document_types() -> document_sets (for example,
    "培训资料套装"). output_prefix may be a bare name (saved in outputs/) or a
    path like /path/to/Desktop/培训资料; files are named <prefix>_1_培训通知.docx,
    _2_培训活动记录.docx, _3_培训活动影像.docx.
    培训通知默认保持一页；确认允许超过一页时传 force_multipage_notice=true。
    如果用户未明确说明是整套还是单个，请先与用户确认后再调用本工具。
    """
    set_info = DOCUMENT_SETS.get(set_name)
    if not set_info:
        available = "、".join(DOCUMENT_SETS)
        return _error(f"unknown document set {set_name!r}. Available: {available}")

    created: list[str] = []
    for index, template_name in enumerate(set_info["templates"], start=1):
        output_name = f"{output_prefix}_{index}_{Path(template_name).stem}.docx"
        result = generate_docx(template_name, content, output_name, force_multipage_notice)
        if result.startswith("Error:"):
            return _error(f"set generation stopped at {template_name}: {result}")
        created.append(result)
    return "Created set '" + set_name + "':\n" + "\n".join(created)


@mcp.tool()
def get_semester_info(start_year: int) -> str:
    """Return academic-year/semester date rules for choosing document dates.

    start_year is the first year of the academic year, for example 2025 for
    "2025——2026学年度". The first semester runs from September of start_year to
    January of the next year; the second semester runs from February to July of
    the next year. Each semester defaults to 18 workday weeks.
    All signature/date fields in generated documents must fall inside the
    relevant semester and be workdays.
    """
    end_year = start_year + 1
    info = {
        "academic_year": SEMESTER_RULES["academic_year_template"].format(start_year=start_year, end_year=end_year),
        "weeks_per_semester": SEMESTER_RULES["weeks_per_semester"],
        "weeks_are_workdays": SEMESTER_RULES["weeks_are_workdays"],
        "first_semester": {
            "period": f"{start_year}年9月～{end_year}年1月",
            "weeks": SEMESTER_RULES["weeks_per_semester"],
            "signature_hint": f"落款取 {start_year}年9月 至 {end_year}年1月 内的工作日（如学期初计划可落款 9 月上旬工作日）。",
        },
        "second_semester": {
            "period": f"{end_year}年2月～{end_year}年7月",
            "weeks": SEMESTER_RULES["weeks_per_semester"],
            "signature_hint": f"落款取 {end_year}年2月 至 {end_year}年7月 内的工作日（如学期初计划可落款 2 月下旬或 3 月上旬工作日）。",
        },
        "rule": "所有落款时间均须落在对应学期时间段内，且取工作日。",
    }
    return json.dumps(info, ensure_ascii=False, indent=2)


@mcp.tool()
def recommend_document_type(description: str) -> str:
    """Recommend the document type for a user's request.

    Pass the user's request (for example "写一个评价计划" or "写一篇演讲稿").
    The returned recommendation is deterministic keyword-based guidance; if the
    user explicitly names a format, follow the user instead.
    """
    text = description or ""
    recommendation = None
    reason = ""

    def choose(doc_type, why):
        nonlocal recommendation, reason
        recommendation = doc_type
        reason = why

    # 0. 周报
    if any(k in text for k in ["周报"]):
        choose("行政周报", "描述包含“周报”，使用行政周报格式（两行标题、部门分节、条目自动编号、空缺写“无”）。")
    # 1. 培训资料类：整套 or 单件
    elif any(k in text for k in ["培训资料", "培训套装", "一整套", "一套"]):
        choose("培训资料套装", "描述包含“培训资料/一套”等关键词，使用培训资料套装（通知+活动记录+活动影像）。")
    elif "培训通知" in text:
        choose("培训通知", "描述包含“培训通知”，使用培训通知模板（红色红头+双红线）。")
    elif "培训记录" in text or "培训活动记录" in text:
        choose("培训活动记录", "描述包含“培训记录”，使用培训活动记录模板（记录表）。")
    elif "培训影像" in text or "培训活动影像" in text:
        choose("培训活动影像", "描述包含“培训影像”，使用培训活动影像模板（照片页）。")
    # 2. 长文类：演讲稿/发言稿/论文
    elif any(k in text for k in ["演讲稿", "发言稿", "讲话稿", "论文"]):
        choose("论文", "描述包含演讲稿/发言稿/论文等关键词，使用论文格式（无页眉、标题不加粗）。")
    # 3. 活动资料套装（方案+总结+影像）
    if any(k in text for k in ["活动资料", "活动方案", "活动总结", "活动影像"]):
        choose("活动方案套装", "描述包含活动方案/总结/影像等关键词，使用活动方案套装（方案+总结+影像三件）。")
    elif "活动方案" in text:
        choose("传统公文活动方案", "描述包含“活动方案”，使用活动方案专用模板（落款单位+部门+日期）。")
    # 4. 一般公文：计划/总结/方案/汇报/报告/请示/通知/意见/决定/通报/函/纪要
    elif any(k in text for k in ["计划", "总结", "方案", "汇报", "报告", "请示", "通知", "意见", "决定", "通报", "函", "纪要"]):
        choose("传统公文", "描述包含公文文种关键词（计划/总结/方案/汇报/报告等），使用标准传统公文格式（无红色红头）。")
    # 5. 默认：标准传统公文最通用
    else:
        choose("传统公文", "未识别到明确文种关键词，默认使用标准传统公文格式；如用户明确指定格式，以用户为准。")

    if recommendation in ("培训资料套装", "活动方案套装"):
        set_info = DOCUMENT_SETS[recommendation]
        result = {
            "description": description,
            "recommended_document_type": recommendation,
            "templates": set_info["templates"],
            "use_tool": f"generate_document_set(\"{recommendation}\", content, output_prefix)",
            "reason": reason,
            "note": "如果用户未明确说明要整套还是单个，请先与用户确认；确认后再用 generate_document_set 生成整套，或用 generate_by_type 生成单个部件。",
        }
        return json.dumps(result, ensure_ascii=False, indent=2)
    if recommendation == "培训资料套装":
        set_info = DOCUMENT_SETS[recommendation]
        result: dict[str, Any] = {
            "description": description,
            "recommended_document_type": recommendation,
            "templates": set_info["templates"],
            "use_tool": f"generate_document_set(\"培训资料套装\", content, output_prefix)",
            "reason": reason,
            "note": "如果用户已明确指定格式，以用户指定为准。",
        }
    else:
        template_file = DOCUMENT_TYPES.get(recommendation or "传统公文", "传统公文.docx")
        result = {
            "description": description,
            "recommended_document_type": recommendation,
            "template": template_file,
            "reason": reason,
            "note": "如果用户已明确指定格式，以用户指定为准；拿不准时可用 get_template_schema 查看字段后再生成。",
        }
    return json.dumps(result, ensure_ascii=False, indent=2)


@mcp.tool()
def validate_docx(docx_path: str) -> str:
    """Inspect a DOCX and report content, layout, and leftover-placeholder checks.

    docx_path may be an absolute path or a filename inside outputs/.
    """
    path = resolve_input_docx(docx_path)
    if not path.is_file():
        return _error(f"document not found: {path}")

    try:
        document = Document(path)
    except Exception as exc:
        return _error(f"opening document: {exc}")

    warnings: list[str] = []
    paragraphs = [p for p in document.paragraphs if p.text.strip()]
    if not paragraphs:
        warnings.append("The document has no visible paragraph text.")
    if not document.sections:
        warnings.append("The document has no page section.")

    # Leftover jinja placeholders mean the template was not fully rendered.
    leftovers = _leftover_placeholders(path)
    if leftovers:
        warnings.append(f"Unrendered template placeholders remain: {leftovers}")

    # Column-count consistency per table (python-docx pads cells, compare gridCol).
    for index, table in enumerate(document.tables):
        grid_cols = len(table._tbl.xpath("./w:tblGrid/w:gridCol"))
        if grid_cols and len(table.columns) != grid_cols:
            warnings.append(
                f"Table {index + 1}: column count {len(table.columns)} differs from table grid {grid_cols}."
            )

    layout: list[dict[str, float]] = []
    for section in document.sections:
        layout.append(
            {
                "top_cm": round(section.top_margin.cm, 2),
                "bottom_cm": round(section.bottom_margin.cm, 2),
                "left_cm": round(section.left_margin.cm, 2),
                "right_cm": round(section.right_margin.cm, 2),
            }
        )

    report = {
        "file": str(path),
        "status": "warning" if warnings else "pass",
        "paragraphs_with_text": len(paragraphs),
        "tables": len(document.tables),
        "sections": len(document.sections),
        "inline_images": len(document.inline_shapes),
        "margins_cm": layout,
        "warnings": warnings,
    }
    return json.dumps(report, ensure_ascii=False, indent=2)


@mcp.tool()
def fix_docx(docx_path: str, output_name: str) -> str:
    """Apply safe basic repairs and save a new DOCX copy.

    Repairs: replace non-breaking spaces with normal spaces and strip accidental
    trailing whitespace from runs. Fonts, styles, tables, and page layout stay intact.
    """
    if not output_name.endswith(".docx"):
        return _error("output_name must end in .docx")

    source_path = resolve_input_docx(docx_path)
    output_path = resolve_output_path(output_name)

    if not source_path.is_file():
        return _error(f"document not found: {source_path}")

    try:
        document = Document(source_path)
        changed = 0
        for paragraph in document.paragraphs:
            for run in paragraph.runs:
                new_text = run.text.replace("\u00a0", " ").rstrip()
                if new_text != run.text:
                    run.text = new_text
                    changed += 1
        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(output_path)
    except Exception as exc:
        return _error(f"fixing document: {exc}")

    return f"Created repaired copy: {output_path} (cleaned {changed} text runs)"


# ---------------------------------------------------------------------------
# Thesis media insertion
# ---------------------------------------------------------------------------

def _leftover_placeholders(path: Path) -> list[str]:
    """Return jinja placeholders still present in the saved docx.

    Checks document.xml and headers/footers (docxtpl renders all of them).
    """
    leftovers: list[str] = []
    with ZipFile(path) as archive:
        parts = ["word/document.xml"] + [
            name
            for name in archive.namelist()
            if re.match(r"word/(header|footer)\d*\.xml", name)
        ]
        for part in parts:
            try:
                xml = archive.read(part).decode("utf-8")
            except KeyError:
                continue
            for match in re.finditer(r"\{\{[^{}]+\}\}|\{%[^%]*%\}", xml):
                token = match.group(0).replace("&lt;", "<").replace("&gt;", ">")
                if token not in leftovers:
                    leftovers.append(token)
    return leftovers


def _set_run_font(run, east_asia: str, size_pt: float, bold: bool = False, latin: str = "Times New Roman") -> None:
    """Set latin + east-asian fonts on a run (python-docx only sets w:ascii by default)."""
    run.font.name = latin
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east_asia)


def _set_cell_font(cell, size_pt: float = 9.0, east_asia: str = "宋体") -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            _set_run_font(run, east_asia, size_pt)


def _apply_three_line_table(table) -> None:
    """Use the academic three-line table convention (top/header-bottom/table-bottom only)."""
    rows = table.rows
    for ri, row in enumerate(rows):
        for cell in row.cells:
            tcpr = cell._tc.get_or_add_tcPr()
            borders = tcpr.first_child_found_in("w:tcBorders")
            if borders is None:
                borders = OxmlElement("w:tcBorders")
                tcpr.append(borders)
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                tag = qn(f"w:{edge}")
                element = borders.find(tag)
                if element is None:
                    element = OxmlElement(f"w:{edge}")
                    borders.append(element)
                element.set(qn("w:val"), "nil")
            emphatic_edges: list[str] = []
            if ri == 0:
                emphatic_edges += ["top", "bottom"]  # top rule + below-header rule
            if ri == len(rows) - 1:
                emphatic_edges += ["bottom"]  # table bottom rule
            for edge in emphatic_edges:
                element = borders.find(qn(f"w:{edge}"))
                element.set(qn("w:val"), "single")
                # Top and table-bottom rules are 1.5pt (sz=12); the
                # below-header rule stays at 0.75pt (sz=6).
                if edge == "top" or (ri == len(rows) - 1 and edge == "bottom"):
                    element.set(qn("w:sz"), "12")
                else:
                    element.set(qn("w:sz"), "6")
                element.set(qn("w:color"), "000000")


def _set_table_no_borders(table) -> None:
    """去掉表格所有边框（用于拼图表格）。"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    borders = tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "nil")


def _image_size_px(path: Path):
    """读取 JPEG/PNG 图片尺寸（不依赖第三方库）。"""
    with open(path, "rb") as fh:
        data = fh.read()
    if data[:3] == b"\xff\xd8\xff":
        i = 2
        while i + 9 < len(data):
            if data[i] != 0xFF:
                i += 1
                continue
            marker = data[i + 1]
            if marker in (0xC0, 0xC1, 0xC2, 0xC3, 0xC5, 0xC6, 0xC7, 0xC9, 0xCA, 0xCB, 0xCD, 0xCE, 0xCF):
                height = int.from_bytes(data[i + 5 : i + 7], "big")
                width = int.from_bytes(data[i + 7 : i + 9], "big")
                return width, height
            seg_len = int.from_bytes(data[i + 2 : i + 4], "big")
            i += 2 + seg_len
    elif data[:8] == b"\x89PNG\r\n\x1a\n":
        width = int.from_bytes(data[16:20], "big")
        height = int.from_bytes(data[20:24], "big")
        return width, height
    return None


def _set_table_cell_margins(table, left_cm: float, right_cm: float) -> None:
    """设置表格单元格左右边距（用于拼图间距）。"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    cell_mar = tblPr.first_child_found_in("w:tblCellMar")
    if cell_mar is None:
        cell_mar = OxmlElement("w:tblCellMar")
        tblPr.append(cell_mar)
    for side, value_cm in (("left", left_cm), ("right", right_cm)):
        tag = qn(f"w:{side}")
        element = cell_mar.find(tag)
        if element is None:
            element = OxmlElement(f"w:{side}")
            cell_mar.append(element)
        element.set(qn("w:w"), str(int(value_cm * 567)))
        element.set(qn("w:type"), "dxa")


def _insert_thesis_media(path: Path, content: dict[str, Any]) -> None:
    """Insert thesis tables/figures.

    - 全局 tables / figures 插入到 MEDIA_INSERT_HERE 标记处（关键词之后）。
    - 每个 subsection 可以带 figures，插入到该子节正文之后的 FIGM 标记处，
      实现图片紧跟在正文之后（图题五号黑体、居中）。
    """
    doc = Document(path)

    def add_figures_after(anchor_element, figures) -> None:
        current_anchor = anchor_element
        for figure in figures:
            paths = figure.get("paths") or ([figure.get("path")] if figure.get("path") else [])
            if not paths:
                continue
            total_width_cm = float(figure.get("width_cm", 13))
            if len(paths) == 1:
                figure_path = resolve_path(str(paths[0]), PROJECT_DIR)
                if not figure_path.is_file():
                    raise ValueError(f"figure image not found: {figure_path}")
                picture_p = doc.add_paragraph()
                picture_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                picture_p.paragraph_format.space_before = Pt(0)
                picture_p.paragraph_format.space_after = Pt(0)
                run = picture_p.add_run()
                run.add_picture(str(figure_path), width=Cm(total_width_cm))
                current_anchor.addnext(picture_p._p)
                current_anchor = picture_p._p
            else:
                # 拼图：1×N 无边框表格，所有图片统一高度、顶部底部对齐。
                gap_cm = float(figure.get("gap_cm", 0.2))
                n = len(paths)
                resolved_paths = []
                for image_path in paths:
                    resolved = resolve_path(str(image_path), PROJECT_DIR)
                    if not resolved.is_file():
                        raise ValueError(f"figure image not found: {resolved}")
                    resolved_paths.append(resolved)
                dims = [_image_size_px(resolved) for resolved in resolved_paths]
                if any(d is None for d in dims):
                    # 读不到尺寸时退化为等分宽度
                    widths = [(total_width_cm - gap_cm * (n - 1)) / n] * n
                else:
                    aspects = [w / h for w, h in dims]
                    available = max(2.0, total_width_cm - gap_cm * n)
                    height_cm = available / sum(aspects)
                    height_cm = max(2.0, min(8.0, height_cm))
                    widths = [height_cm * aspect for aspect in aspects]
                    scale = available / sum(widths)
                    widths = [w * scale for w in widths]
                table = doc.add_table(rows=1, cols=n)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.autofit = False
                _set_table_no_borders(table)
                _set_table_cell_margins(table, gap_cm / 2, gap_cm / 2)
                for index, resolved in enumerate(resolved_paths):
                    cell = table.rows[0].cells[index]
                    cell.width = Cm(max(0.5, widths[index]))
                    paragraph = cell.paragraphs[0]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    run = paragraph.add_run()
                    run.add_picture(str(resolved), width=Cm(max(0.5, widths[index])))
                current_anchor.addnext(table._tbl)
                current_anchor = table._tbl
            caption = figure.get("caption", "")
            if caption:
                caption_p = doc.add_paragraph(caption)
                caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_p.paragraph_format.space_before = Pt(0)
                caption_p.paragraph_format.space_after = Pt(0)
                for r in caption_p.runs:
                    _set_run_font(r, "黑体", 10.5)
                current_anchor.addnext(caption_p._p)
                current_anchor = caption_p._p

    # 1) 子节内联图片：每个 subsection 的 figures 插入对应 FIGM 标记处。
    for chapter in content.get("chapters", []):
        for section in chapter.get("sections", []):
            for subsection in section.get("subsections", []):
                figures = subsection.get("figures", [])
                marker_text = (
                    f"FIGM<{chapter.get('heading', '')}>|"
                    f"<{section.get('heading', '')}>|"
                    f"<{subsection.get('heading', '')}>"
                )
                marker = next((p for p in doc.paragraphs if p.text.strip() == marker_text), None)
                if marker is None:
                    continue
                if figures:
                    add_figures_after(marker._p, figures)
                marker._element.getparent().remove(marker._element)

    # 清理所有残留的 FIGM 标记。
    for paragraph in list(doc.paragraphs):
        if paragraph.text.strip().startswith("FIGM<"):
            paragraph._element.getparent().remove(paragraph._element)

    # 2) 全局 MEDIA_INSERT_HERE：表格 + 图片。
    marker = next((p for p in doc.paragraphs if p.text == "MEDIA_INSERT_HERE"), None)
    if marker is None:
        doc.save(path)
        return
    anchor = marker._p

    def add_after(element) -> None:
        nonlocal anchor
        anchor.addnext(element)
        anchor = element

    for table_spec in content.get("tables", []):
        caption = table_spec.get("caption", "")
        if caption:
            caption_p = doc.add_paragraph(caption)
            caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_p.paragraph_format.space_before = Pt(0)
            caption_p.paragraph_format.space_after = Pt(0)
            for run in caption_p.runs:
                _set_run_font(run, "黑体", 10.5)
            add_after(caption_p._p)
        headers = table_spec.get("headers", [])
        rows = table_spec.get("rows", [])
        table = doc.add_table(rows=1, cols=max(1, len(headers)))
        table.style = "Table Grid"
        for i, value in enumerate(headers):
            table.rows[0].cells[i].text = str(value)
        for row in rows:
            cells = table.add_row().cells
            for i, value in enumerate(row):
                if i < len(cells):
                    cells[i].text = str(value)
        _apply_three_line_table(table)
        for row in table.rows:
            for cell in row.cells:
                _set_cell_font(cell, 9.0, "宋体")
        add_after(table._tbl)

    add_figures_after(anchor, content.get("figures", []))
    marker._element.getparent().remove(marker._element)
    doc.save(path)


# ---------------------------------------------------------------------------
# Prompts (slash-style quick commands in supporting clients)
# ---------------------------------------------------------------------------

@mcp.prompt(title="生成文档（自动判断格式）")
def prompt_generate_any() -> str:
    """根据用户要生成的内容自动判断格式并生成。"""
    return (
        "用户要生成一份文档。请自行判断格式：\n"
        "1. 先调用 recommend_document_type(用户要求)，按其推荐结果选择格式；"
        "例如“写个评价计划”会推荐“传统公文”而不是论文。\n"
        "1.5. 如果推荐结果是培训资料套装或活动方案套装，而用户没有明确说“要整套”还是“只要单个”，"
        "必须先询问用户确认；不要擅自替用户决定。\n"
        "2. 再调用 list_document_types() 查看 document_types 与 format_guide 作为补充。\n"
        "3. 需要落款时间时调用 get_semester_info(起始年份)，确保日期在对应学期内且为工作日。\n"
        "4. 调用 get_template_schema(对应模板文件) 查看字段。\n"
        "5. 根据用户要求自行撰写完整内容，调用 generate_by_type 或 generate_document_set 生成；"
        "只给文件名则默认保存到 outputs/。\n"
        "6. 调用 validate_docx 校验。"
    )


@mcp.prompt(title="查看模板与规则")
def prompt_list_templates() -> str:
    """查看可用文档类型、模板字段、格式判断指南和学年度/学期落款规则。"""
    return (
        "请调用 list_document_types() 查看可用文档类型、模板字段、format_guide（格式判断指南）和学期规则；"
        "需要落款时间时调用 get_semester_info(起始年份) 计算。"
        "根据用户要求自行撰写内容，并套用相应模板格式；生成文件默认保存到 outputs/ 目录。"
    )


@mcp.prompt(title="生成传统公文（计划/总结/汇报等）")
def prompt_generate_official() -> str:
    """生成标准传统公文（计划、总结、方案、汇报、报告、请示等，无红色红头）。"""
    return (
        "用户要生成一般公文（计划/总结/方案/汇报/报告/请示等）。步骤：\n"
        "1. 调用 get_semester_info(起始年份) 确定落款时间（须在对应学期内且为工作日）。\n"
        "2. 调用 get_template_schema(\"传统公文.docx\") 查看字段。\n"
        "3. 根据用户要求自行撰写完整公文内容：单位名称、标题、副标题（可选）、开头段、"
        "各章节标题和段落、附件（可选）、落款单位、日期，内容要贴合用户场景。\n"
        "4. 调用 generate_by_type(\"传统公文\", content, output_name) 生成；"
        "output_name 只给文件名时默认保存到 outputs/，也可给绝对路径。\n"
        "5. 调用 validate_docx(output_name) 校验。"
    )


@mcp.prompt(title="生成论文/演讲稿等长文")
def prompt_generate_thesis() -> str:
    """生成论文、演讲稿、发言稿等长文（论文格式：无页眉、标题不加粗）。"""
    return (
        "用户要生成论文/演讲稿/发言稿等长文。步骤：\n"
        "1. 调用 get_template_schema(\"论文.docx\") 查看字段。\n"
        "2. 根据用户要求自行撰写完整内容：题目、作者（可选）、摘要（可选）、关键词（可选）、"
        "各章节目录与正文段落；演讲稿通常只写题目和正文。\n"
        "3. 调用 generate_by_type(\"论文\", content, output_name) 生成；"
        "output_name 只给文件名时默认保存到 outputs/，也可给绝对路径。\n"
        "4. 调用 validate_docx(output_name) 校验。"
    )


@mcp.prompt(title="生成培训资料套装")
def prompt_generate_training_set() -> str:
    """生成一套培训资料：通知 + 活动记录 + 活动影像。"""
    return (
        "用户要生成一套培训资料（培训通知 + 培训活动记录 + 培训活动影像）。步骤：\n"
        "1. 调用 get_semester_info(起始年份) 确定培训时间（须在对应学期内且为工作日）。\n"
        "2. 调用 get_template_schema(\"培训通知.docx\") 等查看字段。\n"
        "3. 根据用户要求自行撰写完整培训内容：单位、通知标题、培训主题、时间、地点、培训人、"
        "参培人员、培训要求、落款部门与日期、培训内容长文本、培训心得长文本。\n"
        "4. 调用 generate_document_set(\"培训资料套装\", content, output_prefix) 一次生成三件；"
        "output_prefix 只给前缀时默认保存到 outputs/，也可给绝对路径前缀。\n"
        "5. 逐件调用 validate_docx 校验。"
    )


if __name__ == "__main__":
    mcp.run(transport="stdio")

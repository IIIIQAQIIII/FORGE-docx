"""Mission 07 — Scale: large doc inspect/edit + assembly checkpoint/resume."""

import time
from pathlib import Path

import pytest
from docx import Document

from assembly_engine.service import assemble_documents
from document_ir import read_docx
from edit_engine.service import edit_document, inspect_document


def _sha(path):
    import hashlib

    return hashlib.sha256(path.read_bytes()).hexdigest()


# ---------------------------------------------------------------- large doc


def test_large_doc_paginated_inspect_and_targeted_edit(tmp_path):
    src = tmp_path / "large.docx"
    doc = Document()
    for i in range(5000):
        text = f"段落{i}：这是一段用于分页检视的正文内容。" if i % 10 else f"第{i}节标题"
        p = doc.add_paragraph(text)
        if i % 10 == 0:
            p.runs[0].font.bold = True
    doc.add_paragraph("TARGET_END 末尾目标段落")
    doc.add_table(rows=1, cols=2).rows[0].cells[0].text = "末尾表格"
    doc.save(src)

    t0 = time.perf_counter()
    page1 = inspect_document(src, offset=0, limit=100)
    assert page1["total_blocks"] > 5000
    assert page1["next_offset"] == 100
    assert len(page1["blocks"]) == 100

    found = inspect_document(src, query="TARGET_END")
    assert found["total_blocks"] == 1
    target = found["blocks"][0]
    sha = found["source_file_sha256"]
    t_inspect = time.perf_counter() - t0

    out = tmp_path / "large_edited.docx"
    t1 = time.perf_counter()
    result = edit_document(
        src,
        expected_source_sha256=sha,
        edits=[
            {
                "op": "replace_text",
                "source_locator": target["source_locator"],
                "old_text": "TARGET_END",
                "new_text": "EDITED_END",
                "expected_occurrences": 1,
            }
        ],
        output_path=out,
    )
    t_edit = time.perf_counter() - t1
    assert result["status"] == "ok"
    assert result["preservation"]["passed"] is True

    out_ir = read_docx(out)
    texts = [getattr(b, "text", "") for b in out_ir.blocks if getattr(b, "text", "")]
    assert any("EDITED_END" in t for t in texts)
    assert not any("TARGET_END" in t for t in texts)
    assert len([t for t in texts if t.startswith("段落")]) == 4500
    assert len([t for t in texts if t.startswith("第") and "节标题" in t]) == 500

    print(f"\nlarge-doc inspect: {t_inspect:.3f}s, targeted edit: {t_edit:.3f}s")


# ---------------------------------------------------------------- checkpoint helpers


def _make_batch(tmp_path, n):
    src_dir = tmp_path / "sources"
    src_dir.mkdir(exist_ok=True)
    paths = []
    for i in range(n):
        p = src_dir / f"doc_{i:03d}.docx"
        d = Document()
        d.add_paragraph(f"文档{i}内容")
        if i % 10 == 0:
            d.add_table(rows=1, cols=1).rows[0].cells[0].text = f"表{i}"
        d.save(p)
        paths.append(str(p))
    return paths


def test_checkpoint_manifest_persistence_and_cleanup(tmp_path, monkeypatch):
    monkeypatch.setenv("FORGE_HOME", str(tmp_path / "forge_home"))
    sources = _make_batch(tmp_path, 5)
    out = tmp_path / "assembled.docx"
    result = assemble_documents(sources, str(out), explicit_format_hint="正式公文", checkpoint=True)
    assert result["status"] == "ok"
    assert out.exists()

    from assembly_engine.service import _manifest_path

    manifest_path = _manifest_path(result["warnings"][0].split("=")[1])
    assert manifest_path.is_file()
    import json

    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    assert manifest["status"] == "completed"
    assert len(manifest["completed"]) == 5


def test_interrupted_batch_no_output_and_resume_reuses(tmp_path, monkeypatch):
    monkeypatch.setenv("FORGE_HOME", str(tmp_path / "forge_home"))
    import assembly_engine.service as assembly_service

    sources = _make_batch(tmp_path, 100)
    out = tmp_path / "assembled.docx"
    real_reformat = assembly_service.reformat_single
    first_calls = {"n": 0}

    def failing_reformat(source_path, output_path, explicit_profile_id=None):
        first_calls["n"] += 1
        if first_calls["n"] == 31:
            raise RuntimeError("simulated interruption at item 30")
        return real_reformat(source_path=source_path, output_path=output_path, explicit_profile_id=explicit_profile_id)

    monkeypatch.setattr(assembly_service, "reformat_single", failing_reformat)
    with pytest.raises(RuntimeError):
        assemble_documents(sources, str(out), explicit_format_hint="正式公文", checkpoint=True)
    assert first_calls["n"] == 31
    assert not out.exists()

    # manifest should keep 30 completed checkpoints
    from assembly_engine.service import _load_manifest, _jobs_dir

    job_ids = [p.name for p in _jobs_dir().iterdir() if p.is_dir()]
    assert job_ids
    job_id = job_ids[0]
    manifest = _load_manifest(job_id)
    assert len(manifest["completed"]) == 30

    # resume with a counting spy
    resume_calls = {"n": 0}

    def counting_reformat(source_path, output_path, explicit_profile_id=None):
        resume_calls["n"] += 1
        return real_reformat(source_path=source_path, output_path=output_path, explicit_profile_id=explicit_profile_id)

    monkeypatch.setattr(assembly_service, "reformat_single", counting_reformat)
    result = assemble_documents(sources, str(out), explicit_format_hint="正式公文", checkpoint=True, job_id=job_id, resume=True)
    assert result["status"] == "ok"
    assert result["total"] == 100
    assert result["processed"] == 100
    assert result["failed"] == 0
    assert out.exists()
    # only remaining 70 items re-normalized
    assert resume_calls["n"] == 70


def test_resume_config_change_detected(tmp_path, monkeypatch):
    monkeypatch.setenv("FORGE_HOME", str(tmp_path / "forge_home"))
    sources = _make_batch(tmp_path, 3)
    out = tmp_path / "assembled.docx"
    first = assemble_documents(sources, str(out), explicit_format_hint="正式公文", checkpoint=True)
    assert first["status"] == "ok"
    job_id = first["warnings"][0].split("=")[1]
    changed = assemble_documents(sources, str(out), explicit_format_hint="论文格式", checkpoint=True, job_id=job_id, resume=True)
    assert changed["status"] == "error"
    assert changed["errors"][0] == "JOB_CONFIGURATION_CHANGED"


def test_resume_changed_source_detected(tmp_path, monkeypatch):
    monkeypatch.setenv("FORGE_HOME", str(tmp_path / "forge_home"))
    sources = _make_batch(tmp_path, 3)
    out = tmp_path / "assembled.docx"
    first = assemble_documents(sources, str(out), explicit_format_hint="正式公文", checkpoint=True)
    assert first["status"] == "ok"
    job_id = first["warnings"][0].split("=")[1]
    # change a completed source
    d = Document(sources[0])
    d.add_paragraph("修改后的内容")
    d.save(sources[0])
    resumed = assemble_documents(sources, str(out), explicit_format_hint="正式公文", checkpoint=True, job_id=job_id, resume=True)
    assert resumed["status"] == "error"
    assert resumed["errors"][0].startswith("CHECKPOINT_SOURCE_CHANGED")

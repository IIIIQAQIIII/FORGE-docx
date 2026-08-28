"""Mission 03-C2 Extended Semantic Role Annotation 测试。"""

import struct
import zlib

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from document_ir import ParagraphBlock, read_docx
from semantics.annotator import annotate_document


def _png_bytes():
    def chunk(tag, data):
        c = tag + data
        return struct.pack(">I", len(data)) + c + struct.pack(">I", zlib.crc32(c))

    ihdr = struct.pack(">IIBBBBB", 8, 8, 8, 2, 0, 0, 0)
    idat = zlib.compress(b"\x00" + b"\xAA\xBB\xCC" * 8)
    return b"\x89PNG\r\n\x1a\n" + chunk(b"IHDR", ihdr) + chunk(b"IDAT", idat) + chunk(b"IEND", b"")


def _ir(paragraphs, with_table=False, with_image=False):
    doc = Document()
    for text, style in paragraphs:
        p = doc.add_paragraph(text)
        if style.get("center"):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if style.get("bold"):
            for run in p.runs:
                run.font.bold = True
        if style.get("size"):
            for run in p.runs:
                run.font.size = Pt(style["size"])
    if with_table:
        table = doc.add_table(rows=1, cols=1)
        table.rows[0].cells[0].text = "数据"
    if with_image:
        png = "/tmp/_semantic_img.png"
        with open(png, "wb") as f:
            f.write(_png_bytes())
        doc.add_picture(png, width=Cm(3))
    path = "/tmp/_semantic_ext.docx"
    doc.save(path)
    ir = read_docx(path)
    annotate_document(ir)
    return ir


def test_1_full_date():
    ir = _ir([("2026年8月28日", {})])
    assert ir.blocks[0].semantic_role == "date"


def test_2_date_with_extra_text_not_date():
    ir = _ir([("活动时间为2026年8月28日上午9点。", {})])
    assert ir.blocks[0].semantic_role != "date"


def test_3_end_org_and_date_signature():
    ir = _ir([("XX幼儿园", {}), ("2026年8月28日", {})])
    assert ir.blocks[0].semantic_role == "signature"
    assert ir.blocks[1].semantic_role == "date"


def test_4_start_org_and_title():
    ir = _ir(
        [
            ("XX幼儿园", {}),
            ("秋季运动会活动方案", {"center": True, "size": 22, "bold": True}),
        ]
    )
    assert ir.blocks[0].semantic_role == "organization"
    assert ir.blocks[1].semantic_role == "title"


def test_5_same_org_different_position():
    start = _ir([("XX幼儿园", {}), ("活动方案", {"center": True, "size": 22, "bold": True})])
    end = _ir([("XX幼儿园", {}), ("2026年8月28日", {})])
    assert start.blocks[0].semantic_role == "organization"
    assert end.blocks[0].semantic_role == "signature"


def test_6_author_label():
    ir = _ir([("作者：张三", {})])
    assert ir.blocks[0].semantic_role == "author"


def test_7_author_label_variant():
    ir = _ir([("撰稿：张三", {})])
    assert ir.blocks[0].semantic_role == "author"


def test_8_body_person_not_author():
    ir = _ir([("张三老师组织幼儿开展活动。", {})])
    assert ir.blocks[0].semantic_role != "author"


def test_9_figure_caption():
    ir = _ir([("图1 活动现场", {})])
    assert ir.blocks[0].semantic_role == "caption"
    assert ir.blocks[0].metadata["caption_type"] == "figure"


def test_10_table_caption():
    ir = _ir([("表1 幼儿发展数据", {})])
    assert ir.blocks[0].semantic_role == "caption"
    assert ir.blocks[0].metadata["caption_type"] == "table"


def test_11_image_adjacent_caption_evidence():
    ir = _ir([("图1 活动现场", {})], with_image=True)
    cap = ir.blocks[0]
    assert cap.semantic_role == "caption"
    assert "adjacent_to_image" in cap.role_evidence


def test_12_short_paragraph_not_caption():
    ir = _ir([("活动内容", {})])
    assert ir.blocks[0].semantic_role != "caption"


def test_13_full_document_roles():
    ir = _ir(
        [
            ("XX幼儿园", {}),
            ("秋季运动会活动方案", {"center": True, "size": 22, "bold": True}),
            ("作者：张三", {}),
            ("一、活动目标", {}),
            ("这是正文内容，用来测试正文角色的识别情况。", {}),
            ("XX幼儿园", {}),
            ("2026年8月28日", {}),
        ]
    )
    roles = [b.semantic_role for b in ir.blocks]
    assert roles[0] == "organization"
    assert roles[1] == "title"
    assert roles[2] == "author"
    assert roles[3] == "heading_1"
    assert roles[4] == "body"
    assert roles[5] == "signature"
    assert roles[6] == "date"


def test_14_fingerprints_preserved():
    ir = _ir(
        [
            ("XX幼儿园", {}),
            ("秋季运动会活动方案", {"center": True, "size": 22, "bold": True}),
            ("作者：张三", {}),
            ("一、活动目标", {}),
            ("正文内容", {}),
            ("XX幼儿园", {}),
            ("2026年8月28日", {}),
        ]
    )
    before = ir.content_fingerprint
    annotate_document(ir)
    after = ir.content_fingerprint
    assert before.text_sha256 == after.text_sha256
    assert before.structure_sha256 == after.structure_sha256
    assert before.media_sha256 == after.media_sha256
    assert before.content_sequence_sha256 == after.content_sequence_sha256


def test_15_c1_regression_heading():
    ir = _ir([("一、活动目标", {})])
    assert ir.blocks[0].semantic_role == "heading_1"


def test_16_title_observation_without_bold_large_font():
    ir = _ir([("2026年秋季教师案例汇编", {"center": True})])
    block = ir.blocks[0]
    # 观察：当前 C1 模型依赖 bold/larger font，这里应仍为 body/unknown
    assert block.semantic_role in ("body", "unknown")
    # 记录 evidence 供观察
    print(f"\n[title observation] role={block.semantic_role} confidence={block.role_confidence} evidence={block.role_evidence}")

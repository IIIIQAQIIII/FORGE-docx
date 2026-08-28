"""Mission 03-C1 Core Semantic Role Annotation 测试。"""

import struct
import zlib

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from document_ir import ParagraphBlock, TableBlock, read_docx
from semantics.annotator import annotate_document


def _png_bytes():
    def chunk(tag, data):
        c = tag + data
        return struct.pack(">I", len(data)) + c + struct.pack(">I", zlib.crc32(c))

    ihdr = struct.pack(">IIBBBBB", 8, 8, 8, 2, 0, 0, 0)
    idat = zlib.compress(b"\x00" + b"\xAA\xBB\xCC" * 8)
    return b"\x89PNG\r\n\x1a\n" + chunk(b"IHDR", ihdr) + chunk(b"IDAT", idat) + chunk(b"IEND", b"")


def _ir_with_paragraphs(paragraphs):
    """paragraphs: list of (text, style dict)"""
    doc = Document()
    for text, style in paragraphs:
        p = doc.add_paragraph(text)
        if style.get("center"):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if style.get("bold"):
            for run in p.runs:
                run.font.bold = True
        if style.get("size"):
            for run in p.runs:
                run.font.size = Pt(style["size"])
    path = "/tmp/_semantic.docx"
    doc.save(path)
    ir = read_docx(path)
    annotate_document(ir)
    return ir


def test_1_heading_1():
    ir = _ir_with_paragraphs([("一、活动目标", {})])
    assert ir.blocks[0].semantic_role == "heading_1"


def test_2_heading_2():
    ir = _ir_with_paragraphs([("（一）活动准备", {})])
    assert ir.blocks[0].semantic_role == "heading_2"


def test_3_heading_3():
    ir = _ir_with_paragraphs([("1. 场地准备", {})])
    assert ir.blocks[0].semantic_role == "heading_3"


def test_4_date_not_heading_3():
    ir = _ir_with_paragraphs([("2026.9.1", {})])
    assert ir.blocks[0].semantic_role != "heading_3"


def test_5_title():
    ir = _ir_with_paragraphs(
        [("幼儿园秋季运动会活动方案", {"center": True, "size": 22, "bold": True})]
    )
    assert ir.blocks[0].semantic_role == "title"


def test_6_center_not_necessarily_title():
    ir = _ir_with_paragraphs([("活动时间：2026年9月10日", {"center": True})])
    assert ir.blocks[0].semantic_role != "title"


def test_7_subtitle():
    ir = _ir_with_paragraphs(
        [
            ("幼儿园秋季运动会活动方案", {"center": True, "size": 22, "bold": True}),
            ("——2026年秋季学期", {"center": True}),
        ]
    )
    roles = [b.semantic_role for b in ir.blocks]
    assert roles[0] == "title"
    assert roles[1] == "subtitle"


def test_8_body():
    ir = _ir_with_paragraphs([("这是一段比较长的普通正文段落，用来测试正文角色是否能够被正确识别出来。", {})])
    assert ir.blocks[0].semantic_role == "body"


def test_9_empty():
    ir = _ir_with_paragraphs([("", {})])
    assert ir.blocks[0].semantic_role == "empty"


def test_10_image_not_empty(tmp_path):
    p = tmp_path / "img.docx"
    png = tmp_path / "img.png"
    png.write_bytes(_png_bytes())
    doc = Document()
    para = doc.add_paragraph()
    para.add_run().add_picture(str(png), width=Cm(2))
    doc.save(p)
    ir = read_docx(p)
    annotate_document(ir)
    assert ir.blocks[0].semantic_role != "empty"


def test_11_consecutive_heading_context_boost():
    ir = _ir_with_paragraphs(
        [
            ("一、活动目标", {}),
            ("二、活动时间", {}),
            ("三、活动地点", {}),
        ]
    )
    # 第二个标题应包含上下文 evidence
    evidence = ir.blocks[1].role_evidence
    assert "context_heading_before" in evidence
    assert ir.blocks[1].semantic_role == "heading_1"


def test_12_long_heading_prefix_not_confident():
    long_text = "一、" + "这是一个非常长的标题内容" * 10
    ir = _ir_with_paragraphs([(long_text, {})])
    assert ir.blocks[0].semantic_role == "heading_1"
    # 长段落的 confidence 不能因前缀接近 1.0
    assert ir.blocks[0].role_confidence < 0.95


def test_13_numbering_metadata(tmp_path):
    p = tmp_path / "num.docx"
    doc = Document()
    doc.add_paragraph("dummy")
    doc.save(p)
    from zipfile import ZIP_DEFLATED, ZipFile

    with ZipFile(p) as zin:
        parts = {info.filename: zin.read(info.filename) for info in zin.infolist()}
    xml = parts["word/document.xml"].decode("utf-8")
    xml = xml.replace(
        "<w:p>",
        '<w:p><w:pPr><w:numPr><w:ilvl w:val="0"/><w:numId w:val="3"/></w:numPr></w:pPr>',
        1,
    )
    parts["word/document.xml"] = xml.encode("utf-8")
    with ZipFile(p, "w", ZIP_DEFLATED) as zout:
        for name, data in parts.items():
            zout.writestr(name, data)

    ir = read_docx(p)
    block = [b for b in ir.blocks if isinstance(b, ParagraphBlock)][0]
    assert block.metadata["numbering"]["num_id"] == "3"
    assert block.metadata["numbering"]["ilvl"] == "0"


def test_14_fingerprints_preserved_after_annotation(tmp_path):
    p = tmp_path / "fp.docx"
    png = tmp_path / "img.png"
    png.write_bytes(_png_bytes())
    doc = Document()
    doc.add_paragraph("一、标题")
    table = doc.add_table(rows=1, cols=1)
    table.rows[0].cells[0].text = "内容"
    doc.add_paragraph("")
    para = doc.add_paragraph()
    para.add_run().add_picture(str(png), width=Cm(2))
    doc.add_paragraph("正文内容")
    doc.save(p)

    ir = read_docx(p)
    before = ir.content_fingerprint
    annotate_document(ir)
    after = ir.content_fingerprint
    assert before.text_sha256 == after.text_sha256
    assert before.structure_sha256 == after.structure_sha256
    assert before.media_sha256 == after.media_sha256
    assert before.content_sequence_sha256 == after.content_sequence_sha256


def test_15_regression_fingerprints_after_annotation(tmp_path):
    p = tmp_path / "reg.docx"
    doc = Document()
    doc.add_paragraph("A")
    doc.add_paragraph("B")
    doc.save(p)
    ir1 = read_docx(p)
    annotate_document(ir1)
    ir2 = read_docx(p)
    assert ir1.content_fingerprint.text_sha256 == ir2.content_fingerprint.text_sha256
    assert ir1.content_fingerprint.structure_sha256 == ir2.content_fingerprint.structure_sha256

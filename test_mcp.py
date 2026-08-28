"""pytest 测试：覆盖 FORGE MCP 的核心工具。"""

import asyncio
import json
import os
import tempfile

import server


def call_sync(tool: str, args: dict) -> str:
    async def _call():
        output = await server.mcp.call_tool(tool, args)
        if isinstance(output, (list, tuple)):
            parts = output[0] if isinstance(output, tuple) else output
            if isinstance(parts, (list, tuple)) and parts:
                return parts[0].text if hasattr(parts[0], "text") else str(parts[0])
        return str(output)

    return asyncio.run(_call())


def test_list_document_types():
    data = json.loads(call_sync("list_document_types", {}))
    assert "document_types" in data
    assert "document_sets" in data
    assert "semester_rule" in data
    assert "培训资料套装" in data["document_sets"]
    assert "活动方案套装" in data["document_sets"]


def test_get_template_schema():
    data = json.loads(call_sync("get_template_schema", {"template_name": "论文.docx"}))
    assert "fields" in data
    assert "example" in data
    assert "title" in data["fields"]


def test_semester_info():
    data = json.loads(call_sync("get_semester_info", {"start_year": 2025}))
    assert data["academic_year"] == "2025——2026学年度"
    assert data["first_semester"]["period"] == "2025年9月～2026年1月"
    assert data["second_semester"]["period"] == "2026年2月～2026年7月"


def test_recommend_official_plan():
    data = json.loads(call_sync("recommend_document_type", {"description": "写个评价计划"}))
    assert data["recommended_document_type"] == "传统公文"


def test_generate_sample():
    out = call_sync(
        "generate_docx",
        {
            "template_name": "sample_template.docx",
            "content": {"title": "测试", "body": "正文", "signature": "签名", "date": "2026年2月23日"},
            "output_name": "_pytest_sample.docx",
        },
    )
    assert out.startswith("Created:")
    os.remove(server.OUTPUTS_DIR / "_pytest_sample.docx")


def test_generate_official():
    out = call_sync(
        "generate_by_type",
        {
            "document_type": "传统公文",
            "content": {
                "organization": "某市示范幼儿园",
                "title": "测试计划",
                "subtitle": "",
                "opening": "为开展测试，现制定计划如下：",
                "sections": [{"heading": "一、目标", "paragraphs": ["完成测试。"]}],
                "attachment": "",
                "signature": "某市示范幼儿园",
                "date": "2026年2月23日",
            },
            "output_name": "_pytest_official.docx",
        },
    )
    assert out.startswith("Created:")
    os.remove(server.OUTPUTS_DIR / "_pytest_official.docx")


def test_format_profiles_registry():
    data = json.loads(call_sync("list_format_profiles", {}))
    assert data["format_model_version"] == "2.0"
    assert "generic_document" in data["format_profiles"]
    assert "official_standard" in data["resolved_profiles"]
    resolved = data["resolved_profiles"]["activity_plan_standard"]
    assert resolved["signature"]["department"] is True
    assert resolved["body"]["font"] == "仿宋_GB2312"


def test_path_traversal_rejected():
    out = call_sync(
        "generate_docx",
        {"template_name": "../server.py", "content": {}, "output_name": "_pytest_escape.docx"},
    )
    assert out.startswith("Error:")


def test_validate_two_layers():
    from docx import Document

    path = server.OUTPUTS_DIR / "_pytest_bad.docx"
    doc = Document()
    doc.add_paragraph("帐号错误和中文,标点")
    doc.save(path)
    data = json.loads(call_sync("validate_docx", {"docx_path": "_pytest_bad.docx"}))
    os.remove(path)
    assert data["status"] == "warning"
    assert any("帐号" in issue for issue in data["text_checks"])


def test_reformat_preserves_text(tmp_path):
    from docx import Document

    src = tmp_path / "src.docx"
    doc = Document()
    doc.add_paragraph("基于STEM教育理念的实践研究")
    doc.add_paragraph("摘要：测试摘要")
    doc.add_paragraph("关键词：测试")
    doc.add_paragraph("一、引言")
    doc.add_paragraph("（一）研究背景")
    doc.add_paragraph("1.理论意义")
    doc.add_paragraph("这是一段不能丢失的正文内容。")
    doc.add_paragraph("五、结论")
    doc.add_paragraph("结论直接段落，不能丢失。")
    doc.save(src)

    out = call_sync(
        "reformat_docx",
        {"docx_path": str(src), "document_type": "论文", "output_name": "_pytest_reformat.docx"},
    )
    assert out.startswith("Created:")
    from docx import Document as D

    rendered = D(server.OUTPUTS_DIR / "_pytest_reformat.docx")
    text = "\n".join(p.text for p in rendered.paragraphs)
    assert "这是一段不能丢失的正文内容。" in text
    assert "结论直接段落，不能丢失。" in text
    os.remove(server.OUTPUTS_DIR / "_pytest_reformat.docx")

"""Mission 03-B Document Fidelity Hardening 测试。"""

import hashlib
import struct
import zlib
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.shared import Cm

from document_ir import OpaqueBlock, ParagraphBlock, TableBlock, read_docx

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
W = "{%s}" % W_NS


def _png_bytes():
    def chunk(tag, data):
        c = tag + data
        return struct.pack(">I", len(data)) + c + struct.pack(">I", zlib.crc32(c))

    ihdr = struct.pack(">IIBBBBB", 8, 8, 8, 2, 0, 0, 0)
    idat = zlib.compress(b"\x00" + b"\xAA\xBB\xCC" * 8)
    return b"\x89PNG\r\n\x1a\n" + chunk(b"IHDR", ihdr) + chunk(b"IDAT", idat) + chunk(b"IEND", b"")


def _rewrite_zip(path, edit):
    with ZipFile(path) as zin:
        parts = {info.filename: zin.read(info.filename) for info in zin.infolist()}
    parts = edit(parts)
    with ZipFile(path, "w", ZIP_DEFLATED) as zout:
        for name, data in parts.items():
            zout.writestr(name, data)


def _insert_body_xml(path, xml_fragment, before_sectpr=True):
    def edit(parts):
        xml = parts["word/document.xml"].decode("utf-8")
        if before_sectpr:
            xml = xml.replace("<w:sectPr", xml_fragment + "<w:sectPr", 1)
        else:
            xml = xml.replace("</w:body>", xml_fragment + "</w:body>", 1)
        parts["word/document.xml"] = xml.encode("utf-8")
        return parts

    _rewrite_zip(path, edit)


def _add_rel(path, rid, target):
    def edit(parts):
        rels = parts["word/_rels/document.xml.rels"].decode("utf-8")
        rel = f'<Relationship Id="{rid}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink" Target="{target}" TargetMode="External"/>'
        rels = rels.replace("</Relationships>", rel + "</Relationships>", 1)
        parts["word/_rels/document.xml.rels"] = rels.encode("utf-8")
        return parts

    _rewrite_zip(path, edit)


def _sha256_file(path):
    with open(path, "rb") as f:
        return hashlib.sha256(f.read()).hexdigest()


def test_1_text_image_text_inline_order(tmp_path):
    p = tmp_path / "f1.docx"
    png = tmp_path / "img.png"
    png.write_bytes(_png_bytes())
    doc = Document()
    para = doc.add_paragraph()
    para.add_run("TEXT A")
    para.add_run().add_picture(str(png), width=Cm(2))
    para.add_run("TEXT B")
    doc.save(p)

    ir = read_docx(p)
    block = ir.blocks[0]
    kinds = [i.type for i in block.inline]
    assert kinds == ["text", "image", "text"]
    assert block.inline[0].text == "TEXT A"
    assert block.inline[2].text == "TEXT B"
    assert block.inline[1].media_id is not None


def test_2_text_two_images_text(tmp_path):
    p = tmp_path / "f2.docx"
    png = tmp_path / "img.png"
    png.write_bytes(_png_bytes())
    doc = Document()
    para = doc.add_paragraph()
    para.add_run("A")
    para.add_run().add_picture(str(png), width=Cm(2))
    para.add_run("B")
    para.add_run().add_picture(str(png), width=Cm(2))
    para.add_run("C")
    doc.save(p)

    ir = read_docx(p)
    kinds = [i.type for i in ir.blocks[0].inline]
    assert kinds == ["text", "image", "text", "image", "text"]


def test_3_hyperlink_text_and_target(tmp_path):
    p = tmp_path / "f3.docx"
    doc = Document()
    doc.add_paragraph("dummy")
    doc.save(p)
    _add_rel(p, "rId99", "https://example.com")
    _insert_body_xml(
        p,
        '<w:p><w:hyperlink r:id="rId99"><w:r><w:t>链接文字</w:t></w:r></w:hyperlink></w:p>',
        before_sectpr=False,
    )
    ir = read_docx(p)
    hyperlinks = [
        i for b in ir.blocks if isinstance(b, ParagraphBlock) for i in b.inline if i.type == "hyperlink"
    ]
    assert len(hyperlinks) == 1
    assert hyperlinks[0].text == "链接文字"
    assert hyperlinks[0].target == "https://example.com"


def test_4_tab_and_line_break(tmp_path):
    p = tmp_path / "f4.docx"
    doc = Document()
    doc.add_paragraph("dummy")
    doc.save(p)
    _insert_body_xml(
        p,
        '<w:p><w:r><w:t>A</w:t></w:r><w:r><w:tab/></w:r><w:r><w:t>B</w:t></w:r><w:r><w:br/></w:r><w:r><w:t>C</w:t></w:r></w:p>',
        before_sectpr=False,
    )
    ir = read_docx(p)
    target = [b for b in ir.blocks if isinstance(b, ParagraphBlock) and b.text == "ABC"][0]
    kinds = [i.type for i in target.inline]
    assert kinds == ["text", "tab", "text", "line_break", "text"]


def test_5_page_break_vs_line_break(tmp_path):
    p = tmp_path / "f5.docx"
    doc = Document()
    doc.add_paragraph("dummy")
    doc.save(p)
    _insert_body_xml(
        p,
        '<w:p><w:r><w:br w:type="page"/></w:r><w:r><w:br/></w:r></w:p>',
        before_sectpr=False,
    )
    ir = read_docx(p)
    target = [b for b in ir.blocks if isinstance(b, ParagraphBlock) and b.inline][-1]
    kinds = [i.type for i in target.inline]
    assert kinds == ["page_break", "line_break"]


def test_6_table_cell_multi_paragraph(tmp_path):
    p = tmp_path / "f6.docx"
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    cell.text = "第一段"
    cell.add_paragraph("第二段")
    doc.save(p)

    ir = read_docx(p)
    table_block = [b for b in ir.blocks if isinstance(b, TableBlock)][0]
    paragraphs = table_block.rows[0][0].blocks
    assert [x.text for x in paragraphs] == ["第一段", "第二段"]


def test_7_table_cell_image_media_id(tmp_path):
    p = tmp_path / "f7.docx"
    png = tmp_path / "img.png"
    png.write_bytes(_png_bytes())
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    cell.paragraphs[0].add_run().add_picture(str(png), width=Cm(2))
    doc.save(p)

    ir = read_docx(p)
    table_block = [b for b in ir.blocks if isinstance(b, TableBlock)][0]
    cell_paragraph = table_block.rows[0][0].blocks[0]
    images = [i for i in cell_paragraph.inline if i.type == "image"]
    assert len(images) == 1
    assert images[0].media_id is not None
    assert ir.statistics.image_count == 1


def test_8_merged_cell_metadata(tmp_path):
    p = tmp_path / "f8.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).merge(table.cell(0, 1))
    doc.save(p)

    ir = read_docx(p)
    table_block = [b for b in ir.blocks if isinstance(b, TableBlock)][0]
    assert table_block.rows[0][0].metadata.get("grid_span") == "2"
    # v_merge：手动注入
    _insert_body_xml(
        p,
        '<w:tbl><w:tr><w:tc><w:tcPr><w:vMerge w:val="restart"/></w:tcPr><w:p/></w:tc></w:tr><w:tr><w:tc><w:tcPr><w:vMerge/></w:tcPr><w:p/></w:tc></w:tr></w:tbl>',
        before_sectpr=False,
    )
    ir2 = read_docx(p)
    vtable = [b for b in ir2.blocks if isinstance(b, TableBlock)][-1]
    assert vtable.rows[0][0].metadata.get("v_merge") == "restart"
    assert vtable.rows[1][0].metadata.get("v_merge") == "continue"


def test_9_sdt_wrapping_paragraph(tmp_path):
    p = tmp_path / "f9.docx"
    doc = Document()
    doc.add_paragraph("dummy")
    doc.save(p)
    _insert_body_xml(
        p,
        '<w:sdt><w:sdtContent><w:p><w:r><w:t>SDT内容</w:t></w:r></w:p></w:sdtContent></w:sdt>',
        before_sectpr=False,
    )
    ir = read_docx(p)
    texts = [b.text for b in ir.blocks if isinstance(b, ParagraphBlock)]
    assert "SDT内容" in texts


def test_10_sdt_paragraph_table_paragraph(tmp_path):
    p = tmp_path / "f10.docx"
    doc = Document()
    doc.add_paragraph("dummy")
    doc.save(p)
    _insert_body_xml(
        p,
        '<w:sdt><w:sdtContent><w:p><w:r><w:t>P1</w:t></w:r></w:p><w:tbl><w:tr><w:tc><w:p><w:r><w:t>T1</w:t></w:r></w:p></w:tc></w:tr></w:tbl><w:p><w:r><w:t>P2</w:t></w:r></w:p></w:sdtContent></w:sdt>',
        before_sectpr=False,
    )
    ir = read_docx(p)
    target_blocks = [b for b in ir.blocks if getattr(b, "metadata", {}).get("container") == "sdt"]
    kinds = [type(b).__name__ for b in target_blocks]
    assert kinds == ["ParagraphBlock", "TableBlock", "ParagraphBlock"]


def test_11_unknown_body_content_opaque(tmp_path):
    p = tmp_path / "f11.docx"
    doc = Document()
    doc.add_paragraph("dummy")
    doc.save(p)
    _insert_body_xml(
        p,
        '<w:customThing><w:r><w:t>隐藏内容</w:t></w:r></w:customThing>',
        before_sectpr=False,
    )
    ir = read_docx(p)
    opaque = [b for b in ir.blocks if isinstance(b, OpaqueBlock)]
    assert len(opaque) == 1
    assert opaque[0].xml_tag == "w:customThing"
    assert opaque[0].extracted_text == "隐藏内容"


def test_12_floating_image_anchor(tmp_path):
    p = tmp_path / "f12.docx"
    png = tmp_path / "img.png"
    png.write_bytes(_png_bytes())
    doc = Document()
    doc.add_picture(str(png), width=Cm(2))
    doc.save(p)

    def edit(parts):
        xml = parts["word/document.xml"].decode("utf-8")
        xml = xml.replace(
            "<wp:inline",
            '<wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="0" behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1"',
        ).replace("</wp:inline>", "</wp:anchor>")
        parts["word/document.xml"] = xml.encode("utf-8")
        return parts

    _rewrite_zip(p, edit)
    ir = read_docx(p)
    images = [
        i for b in ir.blocks if isinstance(b, ParagraphBlock) for i in b.inline if i.type == "image"
    ]
    assert len(images) == 1
    assert images[0].metadata.get("layout") == "anchor"


def test_13_sequence_fingerprint_changes_with_position(tmp_path):
    png = tmp_path / "img.png"
    png.write_bytes(_png_bytes())

    doc1 = Document()
    para = doc1.add_paragraph()
    para.add_run("A")
    para.add_run().add_picture(str(png), width=Cm(2))
    para.add_run("B")
    doc1.save(tmp_path / "seq1.docx")

    doc2 = Document()
    doc2.add_paragraph("A")
    para2 = doc2.add_paragraph()
    para2.add_run().add_picture(str(png), width=Cm(2))
    doc2.add_paragraph("B")
    doc2.save(tmp_path / "seq2.docx")

    ir1 = read_docx(tmp_path / "seq1.docx")
    ir2 = read_docx(tmp_path / "seq2.docx")
    assert ir1.content_fingerprint.content_sequence_sha256 != ir2.content_fingerprint.content_sequence_sha256


def test_14_fingerprints_identical_twice(tmp_path):
    p = tmp_path / "f14.docx"
    png = tmp_path / "img.png"
    png.write_bytes(_png_bytes())
    doc = Document()
    doc.add_paragraph("A")
    doc.add_picture(str(png), width=Cm(2))
    doc.add_paragraph("B")
    doc.save(p)

    ir1 = read_docx(p)
    ir2 = read_docx(p)
    f1, f2 = ir1.content_fingerprint, ir2.content_fingerprint
    assert f1.text_sha256 == f2.text_sha256
    assert f1.structure_sha256 == f2.structure_sha256
    assert f1.media_sha256 == f2.media_sha256
    assert f1.content_sequence_sha256 == f2.content_sequence_sha256


def test_15_source_sha256_unchanged(tmp_path):
    p = tmp_path / "f15.docx"
    doc = Document()
    doc.add_paragraph("A")
    doc.save(p)
    before = _sha256_file(p)
    read_docx(p)
    after = _sha256_file(p)
    assert before == after

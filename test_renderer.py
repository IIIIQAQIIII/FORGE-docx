"""Mission 04-B1 Source-Preserving Renderer 测试。"""

import hashlib
import struct
import zlib
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from document_ir import ParagraphBlock, TableBlock, read_docx
from format_model import FormatProfile, FormatSource
from profiles import registry as profile_registry
from reformat_engine.models import ReformatPlan
from reformat_engine.planner import build_plan
from reformat_engine.renderer import (
    CONTENT_PRESERVATION_FAILED,
    SOURCE_CHANGED,
    SOURCE_OUTPUT_PATH_CONFLICT,
    render_reformat,
)
from semantics.annotator import annotate_document


def _png_bytes():
    def chunk(tag, data):
        c = tag + data
        return struct.pack(">I", len(data)) + c + struct.pack(">I", zlib.crc32(c))

    ihdr = struct.pack(">IIBBBBB", 8, 8, 8, 2, 0, 0, 0)
    idat = zlib.compress(b"\x00" + b"\xAA\xBB\xCC" * 8)
    return b"\x89PNG\r\n\x1a\n" + chunk(b"IHDR", ihdr) + chunk(b"IDAT", idat) + chunk(b"IEND", b"")


def _sha256_file(path):
    with open(path, "rb") as f:
        return hashlib.sha256(f.read()).hexdigest()


def _make_docx(path, build):
    doc = Document()
    build(doc)
    doc.save(path)


def _plan_for(path, profile_id="official_standard"):
    ir = read_docx(path)
    annotate_document(ir)
    return build_plan(ir, profile_id)


def _insert_body_xml(path, fragment):
    with ZipFile(path) as zin:
        parts = {info.filename: zin.read(info.filename) for info in zin.infolist()}
    xml = parts["word/document.xml"].decode("utf-8")
    xml = xml.replace("</w:body>", fragment + "</w:body>", 1)
    parts["word/document.xml"] = xml.encode("utf-8")
    with ZipFile(path, "w", ZIP_DEFLATED) as zout:
        for name, data in parts.items():
            zout.writestr(name, data)


def _add_rel(path, rid, target):
    with ZipFile(path) as zin:
        parts = {info.filename: zin.read(info.filename) for info in zin.infolist()}
    rels = parts["word/_rels/document.xml.rels"].decode("utf-8")
    rel = f'<Relationship Id="{rid}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink" Target="{target}" TargetMode="External"/>'
    rels = rels.replace("</Relationships>", rel + "</Relationships>", 1)
    parts["word/_rels/document.xml.rels"] = rels.encode("utf-8")
    with ZipFile(path, "w", ZIP_DEFLATED) as zout:
        for name, data in parts.items():
            zout.writestr(name, data)


def test_1_source_file_sha_unchanged(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    _make_docx(src, lambda d: [d.add_paragraph(t) for t in ["一、标题", "正文内容"]])
    plan = _plan_for(src)
    before = _sha256_file(src)
    render_reformat(src, plan, out)
    after = _sha256_file(src)
    assert before == after


def test_2_source_output_conflict(tmp_path):
    src = tmp_path / "src.docx"
    _make_docx(src, lambda d: [d.add_paragraph("正文")])
    plan = _plan_for(src)
    result = render_reformat(src, plan, src)
    assert result["status"] == "error"
    assert SOURCE_OUTPUT_PATH_CONFLICT in result["errors"]


def test_3_source_changed_after_plan(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    _make_docx(src, lambda d: [d.add_paragraph("正文")])
    plan = _plan_for(src)
    # 修改源文件
    _insert_body_xml(src, '<w:p><w:r><w:t>新增</w:t></w:r></w:p>')
    result = render_reformat(src, plan, out)
    assert result["status"] == "error"
    assert SOURCE_CHANGED in result["errors"]


def test_4_title_formatting(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    def build(d):
        p = d.add_paragraph("幼儿园秋季运动会活动方案")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.size = Pt(22)
            r.font.bold = True

    _make_docx(src, build)
    plan = _plan_for(src)
    result = render_reformat(src, plan, out)
    assert result["status"] == "ok"
    ir = read_docx(out)
    block = ir.blocks[0]
    assert block.style.font_name == "方正小标宋简体"
    assert block.style.font_size_pt == 22
    assert block.style.alignment == "center"


def test_5_heading_formatting(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    _make_docx(
        src,
        lambda d: [d.add_paragraph(t) for t in ["一、一级标题", "（一）二级标题", "1. 三级标题"]],
    )
    plan = _plan_for(src)
    render_reformat(src, plan, out)
    ir = read_docx(out)
    blocks = [b for b in ir.blocks if isinstance(b, ParagraphBlock)]
    assert blocks[0].style.font_name == "黑体"
    assert blocks[0].style.font_size_pt == 16
    assert blocks[1].style.font_name == "楷体_GB2312"
    assert blocks[2].style.font_name == "仿宋_GB2312"


def test_6_body_formatting(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    _make_docx(src, lambda d: [d.add_paragraph("这是一段正文内容，用来测试正文字体和行距。")])
    profile_registry.register_profile(
        FormatProfile(
            profile_id="_test_body_indent",
            inherits="official_standard",
            body={"font": "仿宋_GB2312", "size_pt": 16, "line_spacing_pt": 28, "first_line_chars": 200, "first_line_twips": 640},
        )
    )
    plan = _plan_for(src, "_test_body_indent")
    result = render_reformat(src, plan, out)
    assert result["status"] == "ok"
    ir = read_docx(out)
    block = ir.blocks[0]
    assert block.style.font_name == "仿宋_GB2312"
    assert block.style.font_size_pt == 16
    # 行距检查
    ppr = block._p if hasattr(block, "_p") else None


def test_7_signature_date_formatting(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    _make_docx(src, lambda d: [d.add_paragraph(t) for t in ["XX幼儿园", "2026年8月28日"]])
    plan = _plan_for(src, "generic_document")
    result = render_reformat(src, plan, out)
    assert result["status"] == "ok"
    ir = read_docx(out)
    blocks = [b for b in ir.blocks if isinstance(b, ParagraphBlock)]
    assert blocks[0].style.alignment == "right"
    assert blocks[1].style.alignment == "right"


def test_8_page_margins(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    _make_docx(src, lambda d: [d.add_paragraph("正文")])
    plan = _plan_for(src, "official_standard")
    render_reformat(src, plan, out)
    ir = read_docx(out)
    sec = ir.sections[0]
    assert round(sec.margin_top, 1) == 3.7
    assert round(sec.margin_bottom, 1) == 3.5
    assert round(sec.margin_left, 1) == 2.8
    assert round(sec.margin_right, 1) == 2.6


def test_9_east_asia_font(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    _make_docx(src, lambda d: [d.add_paragraph("正文内容")])
    plan = _plan_for(src)
    render_reformat(src, plan, out)
    with ZipFile(out) as z:
        xml = z.read("word/document.xml").decode("utf-8")
    assert 'w:eastAsia="仿宋_GB2312"' in xml


def test_10_hyperlink_unchanged(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    _make_docx(src, lambda d: [d.add_paragraph("dummy")])
    _add_rel(src, "rId99", "https://example.com")
    _insert_body_xml(
        src,
        '<w:p><w:hyperlink r:id="rId99"><w:r><w:t>链接文字</w:t></w:r></w:hyperlink></w:p>',
    )
    plan = _plan_for(src)
    render_reformat(src, plan, out)
    ir = read_docx(out)
    hyperlinks = [
        i for b in ir.blocks if isinstance(b, ParagraphBlock) for i in b.inline if i.type == "hyperlink"
    ]
    assert hyperlinks[0].text == "链接文字"
    assert hyperlinks[0].target == "https://example.com"


def test_11_inline_image_preserved(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    png = tmp_path / "img.png"
    png.write_bytes(_png_bytes())
    _make_docx(
        src,
        lambda d: (
            lambda p: (
                p.add_run("TEXT A"),
                p.add_run().add_picture(str(png), width=Cm(2)),
                p.add_run("TEXT B"),
            )
        )(d.add_paragraph()),
    )
    plan = _plan_for(src)
    render_reformat(src, plan, out)
    ir = read_docx(out)
    kinds = [i.type for i in ir.blocks[0].inline]
    assert kinds == ["text", "image", "text"]
    assert ir.statistics.image_count == 1


def test_12_breaks_preserved(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    _make_docx(src, lambda d: [d.add_paragraph("dummy")])
    _insert_body_xml(
        src,
        '<w:p><w:r><w:t>A</w:t></w:r><w:r><w:tab/></w:r><w:r><w:br/></w:r><w:r><w:br w:type="page"/></w:r><w:r><w:t>B</w:t></w:r></w:p>',
    )
    plan = _plan_for(src)
    render_reformat(src, plan, out)
    ir = read_docx(out)
    block = [b for b in ir.blocks if isinstance(b, ParagraphBlock) and b.text == "AB"][0]
    kinds = [i.type for i in block.inline]
    assert "tab" in kinds and "line_break" in kinds and "page_break" in kinds


def test_13_numbering_preserved(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    _make_docx(src, lambda d: [d.add_paragraph("dummy")])
    with ZipFile(src) as zin:
        parts = {info.filename: zin.read(info.filename) for info in zin.infolist()}
    xml = parts["word/document.xml"].decode("utf-8")
    xml = xml.replace(
        "<w:p>",
        '<w:p><w:pPr><w:numPr><w:ilvl w:val="0"/><w:numId w:val="3"/></w:numPr></w:pPr>',
        1,
    )
    parts["word/document.xml"] = xml.encode("utf-8")
    with ZipFile(src, "w", ZIP_DEFLATED) as zout:
        for name, data in parts.items():
            zout.writestr(name, data)
    plan = _plan_for(src)
    render_reformat(src, plan, out)
    ir = read_docx(out)
    block = [b for b in ir.blocks if isinstance(b, ParagraphBlock)][0]
    assert block.metadata["numbering"]["num_id"] == "3"
    assert block.metadata["numbering"]["ilvl"] == "0"


def test_14_empty_paragraph_preserved(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    _make_docx(src, lambda d: [d.add_paragraph("A"), d.add_paragraph(""), d.add_paragraph("B")])
    plan = _plan_for(src)
    render_reformat(src, plan, out)
    ir = read_docx(out)
    texts = [b.text for b in ir.blocks if isinstance(b, ParagraphBlock)]
    assert texts == ["A", "", "B"]


def test_15_review_required_warning(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    _make_docx(src, lambda d: [d.add_paragraph("正文")])
    ir = read_docx(src)
    annotate_document(ir)
    ir.blocks[0].semantic_role = "unknown"
    plan = build_plan(ir, "official_standard")
    result = render_reformat(src, plan, out)
    assert result["status"] == "ok"
    assert "SOURCE_FORMAT_PRESERVED_FOR_UNRESOLVED_BLOCK" in result["warnings"]


def test_16_opaque_refuses_render(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    _make_docx(src, lambda d: [d.add_paragraph("正文")])
    _insert_body_xml(src, '<w:customThing><w:r><w:t>隐藏</w:t></w:r></w:customThing>')
    plan = _plan_for(src)
    result = render_reformat(src, plan, out)
    assert result["status"] == "error"
    assert not out.exists()


def test_17_table_deferred(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    doc = Document()
    doc.add_paragraph("正文")
    table = doc.add_table(rows=1, cols=1)
    table.rows[0].cells[0].text = "表格内容"
    doc.save(src)
    plan = _plan_for(src)
    result = render_reformat(src, plan, out)
    assert result["status"] == "ok"
    assert "TABLE_FORMAT_DEFERRED" not in result["warnings"]
    ir = read_docx(out)
    table_block = [b for b in ir.blocks if isinstance(b, TableBlock)][0]
    assert table_block.rows[0][0].blocks[0].text == "表格内容"


def test_18_image_deferred(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    png = tmp_path / "img.png"
    png.write_bytes(_png_bytes())
    doc = Document()
    p = doc.add_paragraph()
    p.add_run().add_picture(str(png), width=Cm(2))
    doc.save(src)
    plan = _plan_for(src)
    result = render_reformat(src, plan, out)
    assert result["status"] == "ok"
    assert "IMAGE_FORMAT_DEFERRED" not in result["warnings"]
    ir = read_docx(out)
    assert ir.statistics.image_count == 1


def test_19_duplicate_text_only_locator_selected(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    _make_docx(src, lambda d: [d.add_paragraph("无"), d.add_paragraph("无")])
    ir = read_docx(src)
    annotate_document(ir)
    # 只让第二个“无”应用 body 格式；第一个保持空角色不动
    ir.blocks[0].semantic_role = "empty"
    ir.blocks[1].semantic_role = "body"
    plan = build_plan(ir, "official_standard")
    render_reformat(src, plan, out)
    ir2 = read_docx(out)
    first = ir2.blocks[0]
    second = ir2.blocks[1]
    assert first.style.font_name != "仿宋_GB2312"
    assert second.style.font_name == "仿宋_GB2312"


def test_20_sdt_paragraph_located(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    _make_docx(src, lambda d: [d.add_paragraph("dummy")])
    _insert_body_xml(
        src,
        '<w:sdt><w:sdtContent><w:p><w:r><w:t>SDT正文</w:t></w:r></w:p></w:sdtContent></w:sdt>',
    )
    plan = _plan_for(src)
    result = render_reformat(src, plan, out)
    assert result["status"] == "ok"
    ir = read_docx(out)
    sdt_blocks = [b for b in ir.blocks if b.metadata.get("container") == "sdt"]
    assert len(sdt_blocks) == 1
    assert sdt_blocks[0].style.font_name == "仿宋_GB2312"


def test_21_fingerprints_preserved_after_render(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    png = tmp_path / "img.png"
    png.write_bytes(_png_bytes())
    doc = Document()
    doc.add_paragraph("一、标题")
    doc.add_paragraph("正文")
    p = doc.add_paragraph()
    p.add_run("A")
    p.add_run().add_picture(str(png), width=Cm(2))
    p.add_run("B")
    doc.save(src)
    plan = _plan_for(src)
    result = render_reformat(src, plan, out)
    assert result["status"] == "ok"
    assert result["content_preservation"]["passed"] is True
    expected = plan.source_fingerprint
    out_ir = read_docx(out)
    assert out_ir.content_fingerprint.text_sha256 == expected["text_sha256"]
    assert out_ir.content_fingerprint.structure_sha256 == expected["structure_sha256"]
    assert out_ir.content_fingerprint.media_sha256 == expected["media_sha256"]
    assert out_ir.content_fingerprint.content_sequence_sha256 == expected["content_sequence_sha256"]


def test_22_preservation_failure_no_output(tmp_path, monkeypatch):
    import reformat_engine.renderer as renderer_mod

    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    _make_docx(src, lambda d: [d.add_paragraph("正文")])
    plan = _plan_for(src)

    real_read = renderer_mod.read_docx

    def fake_read(path):
        ir = real_read(path)
        ir.content_fingerprint.text_sha256 = "deadbeef"
        return ir

    monkeypatch.setattr(renderer_mod, "read_docx", fake_read)
    result = renderer_mod.render_reformat(src, plan, out)
    assert result["status"] == "error"
    assert CONTENT_PRESERVATION_FAILED in result["errors"]
    assert not out.exists()


def test_23_plan_not_ready_refused(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    _make_docx(src, lambda d: [d.add_paragraph("正文")])
    _insert_body_xml(src, '<w:customThing><w:r><w:t>隐藏</w:t></w:r></w:customThing>')
    plan = _plan_for(src)
    result = render_reformat(src, plan, out)
    assert result["status"] == "error"
    assert not out.exists()

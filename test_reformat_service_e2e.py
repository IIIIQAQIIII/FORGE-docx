"""Mission 04-C — End-to-end FORGE Reformat 2.0 pipeline tests."""

import struct
import zlib
from zipfile import ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from document_ir import ParagraphBlock, TableBlock, read_docx
from format_model import FormatProfile, FormatSource
from profiles import registry as profile_registry
from reformat_engine.service import reformat_document
from semantics.annotator import annotate_document

# ---------------------------------------------------------------- helpers


def _png_bytes():
    def chunk(tag, data):
        c = tag + data
        return struct.pack(">I", len(data)) + c + struct.pack(">I", zlib.crc32(c))

    ihdr = struct.pack(">IIBBBBB", 8, 8, 8, 2, 0, 0, 0)
    idat = zlib.compress(b"\x00" + b"\xAA\xBB\xCC" * 8)
    return b"\x89PNG\r\n\x1a\n" + chunk(b"IHDR", ihdr) + chunk(b"IDAT", idat) + chunk(b"IEND", b"")


def _register_profile(profile_id, inherits=None, **slots):
    profile_registry.register_profile(
        FormatProfile(profile_id=profile_id, name=profile_id, source=FormatSource(), inherits=inherits, **slots)
    )


def _visible_texts(ir):
    out = []
    for block in ir.blocks:
        if isinstance(block, ParagraphBlock):
            out.append(block.text)
        elif isinstance(block, TableBlock):
            for row in block.rows:
                for cell in row:
                    out.append("\n".join(p.text for p in cell.blocks))
    return out


def _block_types(ir):
    return [b.type for b in ir.blocks]


def _table_signature(ir):
    rows = []
    for block in ir.blocks:
        if isinstance(block, TableBlock):
            for row in block.rows:
                rows.append([cell.metadata.get("grid_span", "") for cell in row])
                rows.append(["\n".join(p.text for p in cell.blocks) for cell in row])
    return rows


def _media_signature(ir):
    return [(m.relationship_id, m.part_name, m.sha256) for m in ir.media]


def _assert_content_identical(src_ir, out_ir):
    assert _visible_texts(src_ir) == _visible_texts(out_ir)
    assert _block_types(src_ir) == _block_types(out_ir)
    assert _table_signature(src_ir) == _table_signature(out_ir)
    assert _media_signature(src_ir) == _media_signature(out_ir)
    sfp, ofp = src_ir.content_fingerprint, out_ir.content_fingerprint
    assert ofp.text_sha256 == sfp.text_sha256
    assert ofp.structure_sha256 == sfp.structure_sha256
    assert ofp.media_sha256 == sfp.media_sha256
    assert ofp.content_sequence_sha256 == sfp.content_sequence_sha256


def _footer_contains_page(out):
    with ZipFile(out) as z:
        for name in z.namelist():
            if name.startswith("word/footer") and name.endswith(".xml"):
                if "PAGE" in z.read(name).decode("utf-8"):
                    return True
    return False


# ---------------------------------------------------------------- fixtures


def _make_fixture_a(path):
    doc = Document()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("科技之春活动方案")
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    doc.add_paragraph("一、活动主题")
    doc.add_paragraph("二、活动安排")
    doc.add_paragraph("为进一步激发幼儿对科学技术的兴趣，结合园所实际开展系列活动。")
    doc.add_paragraph("某市示范幼儿园")
    doc.add_paragraph("2026年3月19日")
    doc.save(path)


def _make_fixture_b(path):
    doc = Document()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("读书月活动影像")
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    doc.add_paragraph("本次活动围绕阅读主题开展，记录精彩瞬间如下。")
    png = path.parent / "_e2e_img.png"
    png.write_bytes(_png_bytes())
    image_p = doc.add_paragraph()
    image_p.add_run().add_picture(str(png), width=Cm(12))
    doc.add_paragraph("图1 活动流程")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "时间"
    table.rows[0].cells[1].text = "2026年4月"
    doc.save(path)


def _make_fixture_c(path):
    doc = Document()
    doc.add_paragraph("讨论事项")
    doc.add_paragraph("讨论事项")
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.add_run("这是一个普通段落，没有明显文种特征。")

    # hyperlink
    rel_id = doc.part.relate_to("https://example.com", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink_p = doc.add_paragraph()
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rstyle = OxmlElement("w:rStyle")
    rstyle.set(qn("w:val"), "Hyperlink")
    rpr.append(rstyle)
    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = "点击查看"
    run.append(t)
    hyperlink.append(run)
    hyperlink_p._p.append(hyperlink)

    # automatic numbering paragraph (numPr; numbering part injected after save)
    numbered_p = doc.add_paragraph("第一项")
    ppr = numbered_p._p.get_or_add_pPr()
    numpr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), "1")
    numpr.append(ilvl)
    numpr.append(numid)
    ppr.append(numpr)
    doc.save(path)
    _inject_numbering_part(path)


def _inject_numbering_part(path):
    from lxml import etree
    from zipfile import ZIP_DEFLATED, ZipFile

    numbering_xml = (
        '<w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:abstractNum w:abstractNumId="0"><w:multiLevelType w:val="hybridMultilevel"/>'
        '<w:lvl w:ilvl="0"><w:start w:val="1"/><w:numFmt w:val="decimal"/><w:lvlText w:val="%1."/>'
        '<w:lvlJc w:val="left"/><w:pPr><w:ind w:left="420" w:hanging="420"/></w:pPr></w:lvl>'
        "</w:abstractNum><w:num w:numId=\"1\"><w:abstractNumId w:val=\"0\"/></w:num></w:numbering>"
    ).encode("utf-8")
    with ZipFile(path) as zin:
        parts = {info.filename: zin.read(info.filename) for info in zin.infolist()}

    parts["word/numbering.xml"] = numbering_xml

    ct_name = "[Content_Types].xml"
    ct_root = etree.fromstring(parts[ct_name])
    ct_ns = "http://schemas.openxmlformats.org/package/2006/content-types"
    override = etree.SubElement(ct_root, f"{{{ct_ns}}}Override")
    override.set("PartName", "/word/numbering.xml")
    override.set(
        "ContentType",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml",
    )
    parts[ct_name] = etree.tostring(ct_root, xml_declaration=True, encoding="UTF-8", standalone=True)

    rels_name = "word/_rels/document.xml.rels"
    rels_root = etree.fromstring(parts[rels_name])
    rels_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
    max_n = 0
    for rel in rels_root.findall(f"{{{rels_ns}}}Relationship"):
        rid = rel.get("Id") or ""
        if rid.startswith("rId"):
            try:
                max_n = max(max_n, int(rid[3:]))
            except ValueError:
                pass
    rel = etree.SubElement(rels_root, f"{{{rels_ns}}}Relationship")
    rel.set("Id", f"rId{max_n + 1}")
    rel.set(
        "Type",
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering",
    )
    rel.set("Target", "numbering.xml")
    parts[rels_name] = etree.tostring(rels_root, xml_declaration=True, encoding="UTF-8", standalone=True)

    with ZipFile(path, "w", ZIP_DEFLATED) as zout:
        for name, data in parts.items():
            zout.writestr(name, data)


# ---------------------------------------------------------------- tests


def test_e2e_a_official_document(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    _make_fixture_a(src)
    src_sha_before = src.read_bytes().hex()
    src_ir = read_docx(src)

    result = reformat_document(src, output_path=out, explicit_profile_id="official_standard")

    assert result["status"] == "ok"
    assert result["classification"] is not None
    assert result["resolution"]["profile_id"] == "official_standard"
    assert result["output"] == str(out)
    assert out.is_file()
    assert src.read_bytes().hex() == src_sha_before

    out_ir = read_docx(out)
    _assert_content_identical(src_ir, out_ir)

    annotate_document(out_ir)
    blocks_by_text = {b.text: b for b in out_ir.blocks if isinstance(b, ParagraphBlock)}
    title = blocks_by_text["科技之春活动方案"]
    assert title.semantic_role == "title"
    assert title.style.font_name == "方正小标宋简体"
    assert title.style.alignment == "center"
    heading = blocks_by_text["一、活动主题"]
    assert heading.semantic_role == "heading_1"
    assert heading.style.font_name == "黑体"
    body = blocks_by_text["为进一步激发幼儿对科学技术的兴趣，结合园所实际开展系列活动。"]
    assert body.style.font_name == "仿宋_GB2312"
    assert body.style.font_size_pt == 16

    section = read_docx(out).sections[0]
    assert abs(section.margin_top - 3.7) < 0.02
    assert _footer_contains_page(out)


def test_e2e_b_activity_with_media(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    _make_fixture_b(src)
    _register_profile(
        "_e2e_activity",
        inherits="activity_plan_standard",
        table={"text": {"font": "黑体", "size_pt": 10}},
        image={"max_width_cm": 8, "preserve_aspect_ratio": True, "allow_upscale": False, "alignment": "center"},
    )
    src_sha_before = src.read_bytes().hex()
    src_ir = read_docx(src)

    result = reformat_document(src, output_path=out, explicit_profile_id="_e2e_activity")

    assert result["status"] == "ok"
    assert result["output"] == str(out)
    assert src.read_bytes().hex() == src_sha_before
    out_ir = read_docx(out)
    _assert_content_identical(src_ir, out_ir)

    # image safe rules: 12cm → 8cm = 2880000 EMU
    with ZipFile(out) as z:
        doc_xml = z.read("word/document.xml").decode("utf-8")
    assert 'cx="2880000"' in doc_xml

    # table formatting applied
    table = [b for b in out_ir.blocks if isinstance(b, TableBlock)][0]
    cell_style = table.rows[0][0].blocks[0].style
    assert cell_style.font_name == "黑体"
    assert cell_style.font_size_pt == 10
    assert _footer_contains_page(out)


def test_e2e_c_weak_document_with_default(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    _make_fixture_c(src)
    src_sha_before = src.read_bytes().hex()
    src_ir = read_docx(src)

    result = reformat_document(src, output_path=out, allow_default=True)

    assert result["status"] == "ok"
    assert result["resolution"]["profile_id"] == "generic_document"
    assert src.read_bytes().hex() == src_sha_before
    out_ir = read_docx(out)
    _assert_content_identical(src_ir, out_ir)
    assert _footer_contains_page(out)
    # body formatting applied to weak paragraphs
    body_texts = [b.text for b in out_ir.blocks if isinstance(b, ParagraphBlock) and b.text == "讨论事项"]
    assert body_texts, "重复段落应保留"

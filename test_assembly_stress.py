"""Mission 05 — 50-document assembly stress test."""

import struct
import time
import zlib
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt

from assembly_engine.service import assemble_documents
from document_ir import ParagraphBlock, read_docx


def _png_bytes(seed):
    def chunk(tag, data):
        c = tag + data
        return struct.pack(">I", len(data)) + c + struct.pack(">I", zlib.crc32(c))

    color = bytes([(seed * 37 + 40) % 256, (seed * 71 + 80) % 256, (seed * 113 + 120) % 256])
    ihdr = struct.pack(">IIBBBBB", 8, 8, 8, 2, 0, 0, 0)
    idat = zlib.compress(b"\x00" + color * 8)
    return b"\x89PNG\r\n\x1a\n" + chunk(b"IHDR", ihdr) + chunk(b"IDAT", idat) + chunk(b"IEND", b"")


def test_50_doc_assembly(tmp_path):
    src_dir = tmp_path / "sources"
    src_dir.mkdir()
    source_paths = []
    expected_block_counts = []
    expected_nonempty_texts = []

    for i in range(50):
        path = src_dir / f"doc_{i:02d}.docx"
        doc = Document()
        title_text = f"第{i}号材料标题"
        p = doc.add_paragraph(title_text)
        if i % 5 == 0:
            p.runs[0].font.bold = True
            p.runs[0].font.size = Pt(16)
        doc.add_paragraph(f"第{i}号材料正文段落一。")
        if i % 3 == 0:
            doc.add_paragraph("重复出现的段落内容。")
        if i % 7 == 0:
            table = doc.add_table(rows=1, cols=2)
            table.rows[0].cells[0].text = f"表{i}-A"
            table.rows[0].cells[1].text = f"表{i}-B"
        if i % 11 == 0:
            png = src_dir / f"img_{i}.png"
            png.write_bytes(_png_bytes(i))
            doc.add_picture(str(png), width=Cm(5))
        doc.save(path)
        source_paths.append(str(path))
        ir = read_docx(path)
        expected_block_counts.append(len(ir.blocks))
        expected_nonempty_texts.append(
            [b.text for b in ir.blocks if isinstance(b, ParagraphBlock) and b.text]
        )

    out = tmp_path / "assembled_50.docx"
    start = time.perf_counter()
    result = assemble_documents(source_paths, str(out), explicit_format_hint="正式公文")
    elapsed = time.perf_counter() - start

    assert result["status"] == "ok"
    assert result["total"] == 50
    assert result["processed"] == 50
    assert result["failed"] == 0
    assert out.exists()
    assert all(item["per_item_preservation"]["passed"] for item in result["items"])

    assembled_ir = read_docx(out)
    total_imported = sum(expected_block_counts)
    # page_break_between_items=True adds one separator paragraph between items
    assert len(assembled_ir.blocks) == total_imported + 49

    for item, expected_texts in zip(result["items"], expected_nonempty_texts):
        assert item["per_item_preservation"]["passed"] is True
        # per-item preservation already verifies text/sequence; extra visibility check:
        for text in expected_texts:
            assert any(
                isinstance(b, ParagraphBlock) and b.text == text for b in assembled_ir.blocks
            )

    print(f"\n50-doc assembly elapsed: {elapsed:.3f}s")

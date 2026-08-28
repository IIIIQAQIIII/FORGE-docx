"""Mission 04-C — Document Chrome: source-preserving PAGE field tests."""

import re
from zipfile import ZipFile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from document_ir import read_docx
from format_model import FormatProfile, FormatSource
from profiles import registry as profile_registry
from reformat_engine.planner import build_plan
from reformat_engine.renderer import render_reformat
from semantics.annotator import annotate_document


def _profile(profile_id, page_number):
    profile_registry.register_profile(
        FormatProfile(
            profile_id=profile_id,
            name=profile_id,
            source=FormatSource(),
            page_number=page_number,
        )
    )


def _plan(src, profile_id):
    ir = read_docx(src)
    annotate_document(ir)
    return build_plan(ir, profile_id)


def _zip_read(path, name):
    with ZipFile(path) as z:
        return z.read(name).decode("utf-8")


def _footer_parts(out):
    with ZipFile(out) as z:
        return sorted(n for n in z.namelist() if re.match(r"word/footer\d+\.xml", n))


def _add_page_field_runs(paragraph, font=None, size_pt=None):
    def make_rpr():
        rpr = OxmlElement("w:rPr")
        if font:
            rfonts = OxmlElement("w:rFonts")
            rfonts.set(qn("w:ascii"), font)
            rfonts.set(qn("w:hAnsi"), font)
            rfonts.set(qn("w:eastAsia"), font)
            rpr.append(rfonts)
        if size_pt is not None:
            half = str(int(round(size_pt * 2)))
            sz = OxmlElement("w:sz")
            sz.set(qn("w:val"), half)
            rpr.append(sz)
            szcs = OxmlElement("w:szCs")
            szcs.set(qn("w:val"), half)
            rpr.append(szcs)
        return rpr

    def add_run(fld_type=None, instr=None):
        run = paragraph.add_run()
        if fld_type:
            fld = OxmlElement("w:fldChar")
            fld.set(qn("w:fldCharType"), fld_type)
            run._r.append(fld)
        elif instr is not None:
            instr_el = OxmlElement("w:instrText")
            instr_el.set(qn("xml:space"), "preserve")
            instr_el.text = instr
            run._r.append(instr_el)

    add_run(fld_type="begin")
    add_run(instr=" PAGE ")
    add_run(fld_type="end")


def test_1_new_footer_page_field(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    doc = Document()
    doc.add_paragraph("正文")
    doc.save(src)
    _profile("_pn_new", {"enabled": True, "position": "footer", "alignment": "center"})
    result = render_reformat(src, _plan(src, "_pn_new"), out)
    assert result["status"] == "ok"
    xml = _zip_read(out, "word/document.xml")
    assert "w:footerReference" in xml
    footers = _footer_parts(out)
    assert len(footers) >= 1
    footer_xml = _zip_read(out, footers[0])
    assert "PAGE" in footer_xml
    assert '<w:jc w:val="center"/>' in footer_xml


def test_2_alignment_right(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    doc = Document()
    doc.add_paragraph("正文")
    doc.save(src)
    _profile("_pn_right", {"enabled": True, "position": "footer", "alignment": "right"})
    render_reformat(src, _plan(src, "_pn_right"), out)
    footers = _footer_parts(out)
    footer_xml = _zip_read(out, footers[0])
    assert '<w:jc w:val="right"/>' in footer_xml


def test_3_start_at(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    doc = Document()
    doc.add_paragraph("正文")
    doc.save(src)
    _profile("_pn_start", {"enabled": True, "position": "footer", "alignment": "center", "start_at": 5})
    render_reformat(src, _plan(src, "_pn_start"), out)
    xml = _zip_read(out, "word/document.xml")
    assert 'w:pgNumType' in xml
    assert 'w:start="5"' in xml


def test_4_show_on_first_page_false(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    doc = Document()
    doc.add_paragraph("正文")
    doc.save(src)
    _profile(
        "_pn_nofirst",
        {"enabled": True, "position": "footer", "alignment": "center", "show_on_first_page": False},
    )
    render_reformat(src, _plan(src, "_pn_nofirst"), out)
    xml = _zip_read(out, "word/document.xml")
    assert "w:titlePg" in xml
    footers = _footer_parts(out)
    assert len(footers) >= 2
    page_footers = []
    blank_footers = []
    for name in footers:
        content = _zip_read(out, name)
        if "PAGE" in content:
            page_footers.append(name)
        else:
            blank_footers.append(name)
    assert len(page_footers) == 1
    assert len(blank_footers) >= 1


def test_5_existing_footer_text_kept(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    doc = Document()
    doc.add_paragraph("正文")
    doc.sections[0].footer.paragraphs[0].text = "机密文件 第 页"
    doc.save(src)
    _profile("_pn_text", {"enabled": True, "position": "footer", "alignment": "center"})
    render_reformat(src, _plan(src, "_pn_text"), out)
    footers = _footer_parts(out)
    footer_xml = _zip_read(out, footers[0])
    assert "机密文件" in footer_xml
    assert "第" in footer_xml
    assert "PAGE" in footer_xml


def test_6_existing_header_text_kept(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    doc = Document()
    doc.add_paragraph("正文")
    doc.sections[0].header.paragraphs[0].text = "绝密 请勿外传"
    doc.save(src)
    _profile("_pn_header", {"enabled": True, "position": "footer", "alignment": "center"})
    render_reformat(src, _plan(src, "_pn_header"), out)
    with ZipFile(out) as z:
        headers = [n for n in z.namelist() if re.match(r"word/header\d+\.xml", n)]
    assert headers
    header_xml = _zip_read(out, headers[0])
    assert "绝密" in header_xml
    assert "请勿外传" in header_xml


def test_7_replace_existing_page_field_keeps_other_footer_content(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    doc = Document()
    doc.add_paragraph("正文")
    p = doc.sections[0].footer.paragraphs[0]
    p.text = "第 "
    _add_page_field_runs(p, font="宋体", size_pt=9)
    p.add_run(" 页")
    doc.save(src)
    _profile("_pn_replace", {"enabled": True, "position": "footer", "alignment": "center", "font": "黑体", "size_pt": 12})
    render_reformat(src, _plan(src, "_pn_replace"), out)
    footers = _footer_parts(out)
    footer_xml = _zip_read(out, footers[0])
    assert "第" in footer_xml and "页" in footer_xml
    assert footer_xml.count("PAGE") >= 1
    assert 'w:eastAsia="黑体"' in footer_xml
    assert 'w:val="24"' in footer_xml


def test_8_section_page_settings_not_broken(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    doc = Document()
    section = doc.sections[0]
    section.page_width = 210 * 360  # EMU? python-docx uses Emu; Cm easier
    from docx.shared import Cm

    section.page_width = Cm(20)
    section.page_height = Cm(28)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2.8)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.4)
    doc.add_paragraph("正文")
    doc.save(src)
    _profile("_pn_sect", {"enabled": True, "position": "footer", "alignment": "center"})
    render_reformat(src, _plan(src, "_pn_sect"), out)
    before = read_docx(src).sections[0]
    after = read_docx(out).sections[0]
    assert (before.page_width, before.page_height) == (after.page_width, after.page_height)
    assert (before.margin_top, before.margin_bottom) == (after.margin_top, after.margin_bottom)
    assert (before.margin_left, before.margin_right) == (after.margin_left, after.margin_right)


def test_9_content_fingerprints_unaffected_by_page_number(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    doc = Document()
    doc.add_paragraph("标题内容")
    doc.add_paragraph("正文内容")
    doc.save(src)
    _profile("_pn_fp", {"enabled": True, "position": "footer", "alignment": "center"})
    render_reformat(src, _plan(src, "_pn_fp"), out)
    src_ir = read_docx(src)
    out_ir = read_docx(out)
    assert out_ir.content_fingerprint.text_sha256 == src_ir.content_fingerprint.text_sha256
    assert out_ir.content_fingerprint.structure_sha256 == src_ir.content_fingerprint.structure_sha256
    assert out_ir.content_fingerprint.media_sha256 == src_ir.content_fingerprint.media_sha256
    assert out_ir.content_fingerprint.content_sequence_sha256 == src_ir.content_fingerprint.content_sequence_sha256


def test_10_source_file_not_modified(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    doc = Document()
    doc.add_paragraph("正文")
    doc.save(src)
    _profile("_pn_src", {"enabled": True, "position": "footer", "alignment": "center"})
    before = src.read_bytes()
    render_reformat(src, _plan(src, "_pn_src"), out)
    assert src.read_bytes() == before

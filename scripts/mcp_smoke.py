"""Sanitized MCP smoke test (runs in a clean environment, no pytest needed).

Verifies the main MCP entrypoints using generated synthetic fixtures.
"""

from __future__ import annotations

import json
import os
import sys
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document

import server


def main() -> None:
    tmp = Path(tempfile.mkdtemp(prefix="forge_smoke_"))
    os.environ["FORGE_HOME"] = str(tmp / "forge_home")

    checks = []

    def check(name, condition):
        checks.append((name, bool(condition)))
        print(("PASS" if condition else "FAIL"), name)

    # 1. list_document_types
    data = json.loads(server.list_document_types())
    check("list_document_types", "document_types" in data)

    # 2. resolve_document_format
    data = json.loads(server.resolve_document_format("写一个活动方案"))
    check("resolve_document_format", "classification" in data and "resolution" in data)

    # 3. generate a fixture for inspect/edit/reformat/assemble
    src = tmp / "smoke_src.docx"
    doc = Document()
    doc.add_paragraph("冒烟测试标题")
    doc.add_paragraph("冒烟测试正文段落。")
    doc.save(src)

    # inspect_document
    data = json.loads(server.inspect_document(source_path=str(src)))
    check("inspect_document", data["status"] == "ok" and data["total_blocks"] >= 2)
    sha = data["source_file_sha256"]

    # edit_document
    data = json.loads(
        server.edit_document(
            source_path=str(src),
            expected_source_sha256=sha,
            edits=[
                {
                    "op": "replace_text",
                    "source_locator": "body/0",
                    "old_text": "冒烟测试标题",
                    "new_text": "已编辑标题",
                    "expected_occurrences": 1,
                }
            ],
            output_path=str(tmp / "smoke_edited.docx"),
        )
    )
    check("edit_document", data["status"] == "ok" and data["preservation"]["passed"])

    # reformat_document
    data = json.loads(
        server.reformat_document(
            source_path=str(src),
            output_path=str(tmp / "smoke_reformatted.docx"),
            explicit_format_hint="正式公文",
        )
    )
    check("reformat_document", data["status"] == "ok" and data["content_preservation"]["passed"])

    # assemble_documents
    a = tmp / "asm_a.docx"
    b = tmp / "asm_b.docx"
    for path, text in ((a, "汇编材料一"), (b, "汇编材料二")):
        d = Document()
        d.add_paragraph(text)
        d.save(path)
    data = json.loads(
        server.assemble_documents(
            source_paths=[str(a), str(b)],
            output_path=str(tmp / "smoke_assembled.docx"),
            explicit_format_hint="正式公文",
        )
    )
    check("assemble_documents", data["status"] == "ok" and data["processed"] == 2)

    # list_format_profiles
    data = json.loads(server.list_format_profiles())
    check("list_format_profiles", "format_profiles" in data and "origins" in data)

    # list_document_templates
    data = json.loads(server.list_document_templates())
    check("list_document_templates", "templates" in data)

    failed = [name for name, ok in checks if not ok]
    if failed:
        print("FAILED:", failed)
        raise SystemExit(1)
    print(f"smoke ok: {len(checks)} checks")


if __name__ == "__main__":
    main()

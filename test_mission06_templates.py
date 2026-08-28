"""Mission 06 — Dynamic Template Registry tests."""

import json
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document

from open_format.template_registry import (
    ERROR_NO_TEMPLATE_PLACEHOLDERS,
    list_document_templates,
    register_document_template,
    resolve_document_template,
)


def _make_docx_template(path, texts):
    doc = Document()
    for text in texts:
        doc.add_paragraph(text)
    doc.save(path)


def test_valid_docxtpl_registration_and_generation(tmp_path, monkeypatch):
    monkeypatch.setenv("FORGE_HOME", str(tmp_path / "forge_home"))
    tpl = tmp_path / "tpl.docx"
    _make_docx_template(tpl, ["{{title}}", "{{body}}"])
    result = register_document_template(
        template_path=tpl,
        template_id="my_new_template",
        name="我的新模板",
        kind="docxtpl",
        profile_id="generic_document",
        aliases=["我的新模板"],
    )
    assert result["status"] == "ok"
    assert "title" in result["template"]["schema"]["placeholders"]
    assert "body" in result["template"]["schema"]["placeholders"]

    resolved = resolve_document_template("我的新模板")
    assert resolved["origin"] == "user"
    assert resolved["manifest"]["template_id"] == "my_new_template"

    templates = list_document_templates()
    user_entries = [t for t in templates if t["origin"] == "user"]
    assert any(t["template_id"] == "my_new_template" for t in user_entries)

    # dynamic generation through server.generate_by_type (no DOCUMENT_TYPES change)
    import server

    raw = server.generate_by_type("my_new_template", {"title": "动态标题", "body": "动态正文"}, str(tmp_path / "gen.docx"))
    assert raw.startswith("Created:")
    out_doc = Document(tmp_path / "gen.docx")
    texts = [p.text for p in out_doc.paragraphs if p.text.strip()]
    assert texts == ["动态标题", "动态正文"]


def test_docxtpl_without_placeholders_rejected(tmp_path, monkeypatch):
    monkeypatch.setenv("FORGE_HOME", str(tmp_path / "forge_home"))
    tpl = tmp_path / "plain.docx"
    _make_docx_template(tpl, ["没有占位符的普通文档"])
    result = register_document_template(
        template_path=tpl, template_id="plain_template", name="普通", kind="docxtpl", profile_id="generic_document"
    )
    assert result["error"] == ERROR_NO_TEMPLATE_PLACEHOLDERS


def test_reference_kind_accepts_plain_docx(tmp_path, monkeypatch):
    monkeypatch.setenv("FORGE_HOME", str(tmp_path / "forge_home"))
    tpl = tmp_path / "plain.docx"
    _make_docx_template(tpl, ["普通 reference 文档"])
    result = register_document_template(
        template_path=tpl, template_id="ref_template", name="参考", kind="reference", profile_id="generic_document"
    )
    assert result["status"] == "ok"
    assert result["template"]["kind"] == "reference"


def test_template_id_conflict(tmp_path, monkeypatch):
    monkeypatch.setenv("FORGE_HOME", str(tmp_path / "forge_home"))
    tpl = tmp_path / "tpl.docx"
    _make_docx_template(tpl, ["{{title}}"])
    ok = register_document_template(template_path=tpl, template_id="dup_template", name="一", kind="docxtpl", profile_id="generic_document")
    assert ok["status"] == "ok"
    conflict = register_document_template(template_path=tpl, template_id="dup_template", name="二", kind="docxtpl", profile_id="generic_document")
    assert conflict["error"] == "TEMPLATE_ID_CONFLICT"


def test_template_profile_not_found(tmp_path, monkeypatch):
    monkeypatch.setenv("FORGE_HOME", str(tmp_path / "forge_home"))
    tpl = tmp_path / "tpl.docx"
    _make_docx_template(tpl, ["{{title}}"])
    result = register_document_template(template_path=tpl, template_id="bad_profile", name="x", kind="docxtpl", profile_id="no_such_profile")
    assert result["error"] == "PROFILE_NOT_FOUND"


def test_malformed_docx_rejected(tmp_path, monkeypatch):
    monkeypatch.setenv("FORGE_HOME", str(tmp_path / "forge_home"))
    bad = tmp_path / "bad.docx"
    bad.write_bytes(b"this is not a zip")
    result = register_document_template(template_path=bad, template_id="bad_docx", name="坏", kind="docxtpl", profile_id="generic_document")
    assert result["error"] == "MALFORMED_DOCX"


def test_macro_rejected(tmp_path, monkeypatch):
    monkeypatch.setenv("FORGE_HOME", str(tmp_path / "forge_home"))
    tpl = tmp_path / "macro.docx"
    _make_docx_template(tpl, ["{{title}}"])
    with ZipFile(tpl) as zin:
        parts = {info.filename: zin.read(info.filename) for info in zin.infolist()}
    parts["word/vbaProject.bin"] = b"fake macro"
    with ZipFile(tpl, "w", ZIP_DEFLATED) as zout:
        for name, data in parts.items():
            zout.writestr(name, data)
    result = register_document_template(template_path=tpl, template_id="macro_template", name="宏", kind="docxtpl", profile_id="generic_document")
    assert result["error"] == "MACRO_REJECTED"


def test_user_template_survives_registry_reload(tmp_path, monkeypatch):
    monkeypatch.setenv("FORGE_HOME", str(tmp_path / "forge_home"))
    tpl = tmp_path / "tpl.docx"
    _make_docx_template(tpl, ["{{title}}"])
    register_document_template(template_path=tpl, template_id="reload_template", name="重载", kind="docxtpl", profile_id="generic_document")
    # second listing re-reads persisted manifests
    templates = list_document_templates()
    assert any(t["template_id"] == "reload_template" for t in templates)
    resolved = resolve_document_template("reload_template")
    assert resolved is not None


def test_builtin_generation_regression(tmp_path):
    import server

    raw = server.generate_by_type("传统公文", {"organization": "某单位", "title": "测试标题", "opening": "开头", "sections": [], "signature": "某单位", "date": "2026年3月19日"}, str(tmp_path / "builtin.docx"))
    assert raw.startswith("Created:")
    assert (tmp_path / "builtin.docx").exists()

"""Mission 03-A Document IR + Faithful Ordered DOCX Reader 测试。"""

import hashlib
import os
import struct
import zlib

from docx import Document
from docx.shared import Cm

from document_ir import ParagraphBlock, TableBlock, read_docx


def _sha256_file(path):
    with open(path, "rb") as f:
        return hashlib.sha256(f.read()).hexdigest()


def _png_bytes():
    def chunk(tag, data):
        c = tag + data
        return struct.pack(">I", len(data)) + c + struct.pack(">I", zlib.crc32(c))

    ihdr = struct.pack(">IIBBBBB", 8, 8, 8, 2, 0, 0, 0)
    idat = zlib.compress(b"\x00" + b"\xAA\xBB\xCC" * 8)
    return b"\x89PNG\r\n\x1a\n" + chunk(b"IHDR", ihdr) + chunk(b"IDAT", idat) + chunk(b"IEND", b"")


def _make_docx(path, build):
    doc = Document()
    build(doc)
    doc.save(path)


def test_1_pure_paragraphs_order(tmp_path):
    p = tmp_path / "a.docx"
    _make_docx(p, lambda d: [d.add_paragraph(t) for t in ["A", "B", "C"]])
    ir = read_docx(p)
    texts = [b.text for b in ir.blocks if isinstance(b, ParagraphBlock)]
    assert texts == ["A", "B", "C"]


def test_2_paragraph_table_paragraph_order(tmp_path):
    p = tmp_path / "b.docx"
    doc = Document()
    doc.add_paragraph("Before")
    t = doc.add_table(rows=1, cols=1)
    t.rows[0].cells[0].text = "Cell"
    doc.add_paragraph("After")
    doc.save(p)

    ir = read_docx(p)
    kinds = [type(b).__name__ for b in ir.blocks]
    assert kinds == ["ParagraphBlock", "TableBlock", "ParagraphBlock"]
    assert ir.blocks[0].text == "Before"
    assert ir.blocks[2].text == "After"


def test_3_table_cell_content_complete(tmp_path):
    p = tmp_path / "c.docx"
    doc = Document()
    t = doc.add_table(rows=2, cols=2)
    values = [["r0c0", "r0c1"], ["r1c0", "r1c1"]]
    for ri in range(2):
        for ci in range(2):
            t.rows[ri].cells[ci].text = values[ri][ci]
    doc.save(p)

    ir = read_docx(p)
    table = ir.blocks[0]
    assert isinstance(table, TableBlock)
    for ri in range(2):
        for ci in range(2):
            assert table.rows[ri][ci].blocks[0].text == values[ri][ci]


def test_4_empty_paragraph_preserved(tmp_path):
    p = tmp_path / "d.docx"
    doc = Document()
    doc.add_paragraph("A")
    doc.add_paragraph("")
    doc.add_paragraph("B")
    doc.save(p)

    ir = read_docx(p)
    texts = [b.text for b in ir.blocks if isinstance(b, ParagraphBlock)]
    assert texts == ["A", "", "B"]


def test_5_image_recognized(tmp_path):
    p = tmp_path / "e.docx"
    png = tmp_path / "img.png"
    png.write_bytes(_png_bytes())
    doc = Document()
    doc.add_paragraph("Before")
    doc.add_picture(str(png), width=Cm(2))
    doc.add_paragraph("After")
    doc.save(p)

    ir = read_docx(p)
    assert ir.statistics.image_count == 1
    image_inlines = [
        inline
        for b in ir.blocks
        if isinstance(b, ParagraphBlock)
        for inline in b.inline
        if inline.type == "image"
    ]
    assert len(image_inlines) == 1
    assert image_inlines[0].media_id is not None
    assert len(ir.media) == 1


def test_6_media_sha256_stable(tmp_path):
    p = tmp_path / "f.docx"
    png = tmp_path / "img.png"
    png_bytes = _png_bytes()
    png.write_bytes(png_bytes)
    doc = Document()
    doc.add_picture(str(png), width=Cm(2))
    doc.save(p)

    ir = read_docx(p)
    assert ir.media[0].sha256 == hashlib.sha256(png_bytes).hexdigest()


def test_7_section_margins(tmp_path):
    p = tmp_path / "g.docx"
    doc = Document()
    s = doc.sections[0]
    s.top_margin = Cm(3.7)
    s.bottom_margin = Cm(3.5)
    s.left_margin = Cm(2.8)
    s.right_margin = Cm(2.6)
    doc.save(p)

    ir = read_docx(p)
    section = ir.sections[0]
    assert round(section.margin_top, 1) == 3.7
    assert round(section.margin_bottom, 1) == 3.5
    assert round(section.margin_left, 1) == 2.8
    assert round(section.margin_right, 1) == 2.6


def test_8_unicode_not_garbled(tmp_path):
    p = tmp_path / "h.docx"
    doc = Document()
    doc.add_paragraph("中文 English ① —— （一）")
    doc.save(p)

    ir = read_docx(p)
    assert ir.blocks[0].text == "中文 English ① —— （一）"


def test_9_fingerprint_identical_twice(tmp_path):
    p = tmp_path / "i.docx"
    _make_docx(p, lambda d: [d.add_paragraph(t) for t in ["A", "B", "C"]])
    ir1 = read_docx(p)
    ir2 = read_docx(p)
    assert ir1.content_fingerprint.text_sha256 == ir2.content_fingerprint.text_sha256
    assert ir1.content_fingerprint.structure_sha256 == ir2.content_fingerprint.structure_sha256
    assert ir1.content_fingerprint.media_sha256 == ir2.content_fingerprint.media_sha256


def test_10_source_sha256_unchanged(tmp_path):
    p = tmp_path / "j.docx"
    _make_docx(p, lambda d: [d.add_paragraph(t) for t in ["A", "B"]])
    before = _sha256_file(p)
    read_docx(p)
    after = _sha256_file(p)
    assert before == after


def test_11_full_order_with_image(tmp_path):
    p = tmp_path / "k.docx"
    png = tmp_path / "img.png"
    png.write_bytes(_png_bytes())
    doc = Document()
    doc.add_paragraph("P1")
    t = doc.add_table(rows=1, cols=1)
    t.rows[0].cells[0].text = "T1"
    doc.add_paragraph("P2")
    doc.add_picture(str(png), width=Cm(2))
    doc.add_paragraph("P3")
    doc.save(p)

    ir = read_docx(p)
    kinds = [type(b).__name__ for b in ir.blocks]
    assert kinds == ["ParagraphBlock", "TableBlock", "ParagraphBlock", "ParagraphBlock", "ParagraphBlock"]
    assert ir.blocks[0].text == "P1"
    assert ir.blocks[2].text == "P2"
    assert ir.blocks[4].text == "P3"
    image_paragraph = ir.blocks[3]
    assert any(i.type == "image" for i in image_paragraph.inline)

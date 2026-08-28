"""Mission 04-B2 Source-Preserving Table & Image Formatting 测试。"""

import struct
import zlib
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from document_ir import ParagraphBlock, TableBlock, read_docx
from format_model import FormatProfile, FormatSource
from profiles import registry as profile_registry
from reformat_engine.planner import build_plan
from reformat_engine.renderer import CONTENT_PRESERVATION_FAILED, render_reformat
from semantics.annotator import annotate_document


def _png_bytes():
    def chunk(tag, data):
        c = tag + data
        return struct.pack(">I", len(data)) + c + struct.pack(">I", zlib.crc32(c))

    ihdr = struct.pack(">IIBBBBB", 8, 8, 8, 2, 0, 0, 0)
    idat = zlib.compress(b"\x00" + b"\xAA\xBB\xCC" * 8)
    return b"\x89PNG\r\n\x1a\n" + chunk(b"IHDR", ihdr) + chunk(b"IDAT", idat) + chunk(b"IEND", b"")


def _profile(profile_id, inherits="generic_document", **slots):
    profile_registry.register_profile(
        FormatProfile(profile_id=profile_id, source=FormatSource(), inherits=inherits, **slots)
    )


def _plan(path, profile_id="official_standard"):
    ir = read_docx(path)
    annotate_document(ir)
    return build_plan(ir, profile_id)


def _xml(path):
    with ZipFile(path) as z:
        return z.read("word/document.xml").decode("utf-8")


def _table_profile():
    _profile(
        "_b2_table",
        table={
            "alignment": "center",
            "preferred_width_cm": 12,
            "autofit": True,
            "text": {"font": "宋体", "size_pt": 10, "bold": False},
            "cell": {
                "vertical_alignment": "center",
                "margin_top_cm": 0.1,
                "margin_bottom_cm": 0.1,
                "margin_left_cm": 0.2,
                "margin_right_cm": 0.2,
                "shading": "EEEEEE",
            },
            "borders": {
                "top": {"val": "single", "sz": "6", "color": "000000"},
                "bottom": {"val": "single", "sz": "6", "color": "000000"},
                "left": {"val": "single", "sz": "6", "color": "000000"},
                "right": {"val": "single", "sz": "6", "color": "000000"},
                "inside_horizontal": {"val": "single", "sz": "6", "color": "000000"},
                "inside_vertical": {"val": "single", "sz": "6", "color": "000000"},
            },
        },
    )


def _image_profile(max_width_cm=8, allow_upscale=False, alignment="center"):
    _profile(
        "_b2_image",
        image={
            "max_width_cm": max_width_cm,
            "preserve_aspect_ratio": True,
            "allow_upscale": allow_upscale,
            "alignment": alignment,
        },
    )


def test_1_table_font_text_unchanged(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    for ri in range(2):
        for ci in range(2):
            table.rows[ri].cells[ci].text = f"r{ri}c{ci}"
    doc.save(src)
    _table_profile()
    plan = _plan(src, "_b2_table")
    result = render_reformat(src, plan, out)
    assert result["status"] == "ok"
    xml = _xml(out)
    assert "r0c0" in xml and "r1c1" in xml


def test_2_table_alignment(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    doc = Document()
    doc.add_table(rows=1, cols=1).rows[0].cells[0].text = "x"
    doc.save(src)
    _table_profile()
    plan = _plan(src, "_b2_table")
    render_reformat(src, plan, out)
    xml = _xml(out)
    assert '<w:jc w:val="center"/>' in xml


def test_3_table_border(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    doc = Document()
    doc.add_table(rows=1, cols=1).rows[0].cells[0].text = "x"
    doc.save(src)
    _table_profile()
    plan = _plan(src, "_b2_table")
    render_reformat(src, plan, out)
    xml = _xml(out)
    assert "w:tblBorders" in xml
    assert 'w:insideH' in xml and 'w:insideV' in xml


def test_4_cell_vertical_alignment(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    doc = Document()
    doc.add_table(rows=1, cols=1).rows[0].cells[0].text = "x"
    doc.save(src)
    _table_profile()
    plan = _plan(src, "_b2_table")
    render_reformat(src, plan, out)
    xml = _xml(out)
    assert 'w:vAlign' in xml and 'w:val="center"' in xml


def test_5_cell_margins(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    doc = Document()
    doc.add_table(rows=1, cols=1).rows[0].cells[0].text = "x"
    doc.save(src)
    _table_profile()
    plan = _plan(src, "_b2_table")
    render_reformat(src, plan, out)
    xml = _xml(out)
    assert "w:tcMar" in xml


def test_6_grid_span_preserved(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).merge(table.cell(0, 1))
    doc.save(src)
    _table_profile()
    plan = _plan(src, "_b2_table")
    render_reformat(src, plan, out)
    ir = read_docx(out)
    tbl = [b for b in ir.blocks if isinstance(b, TableBlock)][0]
    assert tbl.rows[0][0].metadata.get("grid_span") == "2"


def test_7_v_merge_preserved(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    doc = Document()
    doc.add_table(rows=2, cols=1)
    doc.save(src)
    with ZipFile(src) as zin:
        parts = {info.filename: zin.read(info.filename) for info in zin.infolist()}
    xml = parts["word/document.xml"].decode("utf-8")
    xml = xml.replace(
        "<w:tbl>",
        '<w:tbl><w:tr><w:tc><w:tcPr><w:vMerge w:val="restart"/></w:tcPr><w:p/></w:tc></w:tr><w:tr><w:tc><w:tcPr><w:vMerge/></w:tcPr><w:p/></w:tc></w:tr>',
        1,
    )
    parts["word/document.xml"] = xml.encode("utf-8")
    with ZipFile(src, "w", ZIP_DEFLATED) as zout:
        for name, data in parts.items():
            zout.writestr(name, data)
    _table_profile()
    plan = _plan(src, "_b2_table")
    render_reformat(src, plan, out)
    ir = read_docx(out)
    tbl = [b for b in ir.blocks if isinstance(b, TableBlock)][-1]
    assert tbl.rows[0][0].metadata.get("v_merge") == "restart"
    assert tbl.rows[1][0].metadata.get("v_merge") == "continue"


def test_8_multi_paragraph_cell_order(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    doc = Document()
    cell = doc.add_table(rows=1, cols=1).rows[0].cells[0]
    cell.text = "第一段"
    cell.add_paragraph("第二段")
    doc.save(src)
    _table_profile()
    plan = _plan(src, "_b2_table")
    render_reformat(src, plan, out)
    ir = read_docx(out)
    tbl = [b for b in ir.blocks if isinstance(b, TableBlock)][0]
    assert [p.text for p in tbl.rows[0][0].blocks] == ["第一段", "第二段"]


def test_9_cell_image_preserved(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    png = tmp_path / "img.png"
    png.write_bytes(_png_bytes())
    doc = Document()
    cell = doc.add_table(rows=1, cols=1).rows[0].cells[0]
    cell.paragraphs[0].add_run().add_picture(str(png), width=Cm(3))
    doc.save(src)
    _table_profile()
    plan = _plan(src, "_b2_table")
    render_reformat(src, plan, out)
    ir = read_docx(out)
    assert ir.statistics.image_count == 1


def test_10_undefined_width_kept(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    doc = Document()
    doc.add_table(rows=1, cols=1).rows[0].cells[0].text = "x"
    doc.save(src)
    _profile("_b2_nowidth", table={"text": {"font": "宋体", "size_pt": 10}})
    plan = _plan(src, "_b2_nowidth")
    render_reformat(src, plan, out)
    import re

    def tbl_w(path):
        return re.search(r"<w:tblW [^/]*/>", _xml(path)).group(0)

    assert tbl_w(src) == tbl_w(out)


def test_11_explicit_width_applied(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    doc = Document()
    doc.add_table(rows=1, cols=1).rows[0].cells[0].text = "x"
    doc.save(src)
    _profile("_b2_width", table={"preferred_width_cm": 12})
    plan = _plan(src, "_b2_width")
    render_reformat(src, plan, out)
    xml = _xml(out)
    assert "w:tblW" in xml


def test_12_inline_image_downscale(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    png = tmp_path / "img.png"
    png.write_bytes(_png_bytes())
    doc = Document()
    p = doc.add_paragraph()
    p.add_run().add_picture(str(png), width=Cm(12))
    doc.save(src)
    _image_profile(max_width_cm=8)
    plan = _plan(src, "_b2_image")
    result = render_reformat(src, plan, out)
    assert result["status"] == "ok"
    xml = _xml(out)
    # cx 应缩放到 8cm = 8*360000 = 2880000
    assert 'cx="2880000"' in xml


def test_13_small_image_no_upscale(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    png = tmp_path / "img.png"
    png.write_bytes(_png_bytes())
    doc = Document()
    p = doc.add_paragraph()
    p.add_run().add_picture(str(png), width=Cm(4))
    doc.save(src)
    _image_profile(max_width_cm=8, allow_upscale=False)
    plan = _plan(src, "_b2_image")
    render_reformat(src, plan, out)
    xml = _xml(out)
    # 4cm = 1440000 EMU，保持不变
    assert 'cx="1440000"' in xml


def test_14_media_sha_unchanged(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    png = tmp_path / "img.png"
    png.write_bytes(_png_bytes())
    doc = Document()
    p = doc.add_paragraph()
    p.add_run().add_picture(str(png), width=Cm(12))
    doc.save(src)
    _image_profile(max_width_cm=8)
    plan = _plan(src, "_b2_image")
    render_reformat(src, plan, out)
    before = read_docx(src)
    after = read_docx(out)
    assert before.media[0].sha256 == after.media[0].sha256


def test_15_image_relationship_unchanged(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    png = tmp_path / "img.png"
    png.write_bytes(_png_bytes())
    doc = Document()
    p = doc.add_paragraph()
    p.add_run().add_picture(str(png), width=Cm(12))
    doc.save(src)
    _image_profile(max_width_cm=8)
    plan = _plan(src, "_b2_image")
    render_reformat(src, plan, out)
    before = read_docx(src)
    after = read_docx(out)
    assert before.media[0].relationship_id == after.media[0].relationship_id
    assert before.media[0].part_name == after.media[0].part_name


def test_16_mixed_paragraph_no_alignment_change(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    png = tmp_path / "img.png"
    png.write_bytes(_png_bytes())
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("TEXT A")
    p.add_run().add_picture(str(png), width=Cm(12))
    p.add_run("TEXT B")
    doc.save(src)
    _image_profile(max_width_cm=8, alignment="center")
    plan = _plan(src, "_b2_image")
    result = render_reformat(src, plan, out)
    assert "IMAGE_ALIGNMENT_DEFERRED_FOR_MIXED_PARAGRAPH" in result["warnings"]


def test_17_media_only_alignment_applied(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    png = tmp_path / "img.png"
    png.write_bytes(_png_bytes())
    doc = Document()
    p = doc.add_paragraph()
    p.add_run().add_picture(str(png), width=Cm(12))
    doc.save(src)
    _image_profile(max_width_cm=8, alignment="center")
    plan = _plan(src, "_b2_image")
    result = render_reformat(src, plan, out)
    assert result["status"] == "ok"
    xml = _xml(out)
    assert '<w:jc w:val="center"/>' in xml


def test_18_anchor_deferred(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    png = tmp_path / "img.png"
    png.write_bytes(_png_bytes())
    doc = Document()
    doc.add_picture(str(png), width=Cm(12))
    doc.save(src)
    with ZipFile(src) as zin:
        parts = {info.filename: zin.read(info.filename) for info in zin.infolist()}
    xml = parts["word/document.xml"].decode("utf-8")
    xml = xml.replace(
        "<wp:inline",
        '<wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="0" behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1"',
    ).replace("</wp:inline>", "</wp:anchor>")
    parts["word/document.xml"] = xml.encode("utf-8")
    with ZipFile(src, "w", ZIP_DEFLATED) as zout:
        for name, data in parts.items():
            zout.writestr(name, data)
    _image_profile(max_width_cm=8)
    plan = _plan(src, "_b2_image")
    result = render_reformat(src, plan, out)
    assert "ANCHOR_IMAGE_FORMAT_DEFERRED" in result["warnings"]


def test_19_caption_uses_caption_slot(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    doc = Document()
    doc.add_paragraph("图1 活动现场")
    doc.save(src)
    _profile(
        "_b2_caption",
        caption={"font": "黑体", "size_pt": 10.5, "align": "center"},
        image={"alignment": "right"},
    )
    plan = _plan(src, "_b2_caption")
    render_reformat(src, plan, out)
    ir = read_docx(out)
    cap = [b for b in ir.blocks if isinstance(b, ParagraphBlock)][0]
    assert cap.style.font_name == "黑体"
    assert cap.style.alignment == "center"


def test_20_four_fingerprints_pass(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    png = tmp_path / "img.png"
    png.write_bytes(_png_bytes())
    doc = Document()
    doc.add_paragraph("一、标题")
    doc.add_paragraph("正文")
    p = doc.add_paragraph()
    p.add_run().add_picture(str(png), width=Cm(12))
    table = doc.add_table(rows=1, cols=1)
    table.rows[0].cells[0].text = "表格"
    doc.save(src)
    _table_profile()
    _image_profile(max_width_cm=8)
    plan = _plan(src, "_b2_image")
    result = render_reformat(src, plan, out)
    assert result["content_preservation"]["passed"] is True
    for key in ("text", "structure", "media", "sequence"):
        assert result["content_preservation"][key] is True


def test_21_table_structure_preserved(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).merge(table.cell(0, 1))
    doc.save(src)
    _table_profile()
    plan = _plan(src, "_b2_table")
    result = render_reformat(src, plan, out)
    assert result["content_preservation"]["table_structure"] is True


def test_22_media_relationships_preserved(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    png = tmp_path / "img.png"
    png.write_bytes(_png_bytes())
    doc = Document()
    doc.add_picture(str(png), width=Cm(12))
    doc.save(src)
    _image_profile(max_width_cm=8)
    plan = _plan(src, "_b2_image")
    result = render_reformat(src, plan, out)
    assert result["content_preservation"]["media_relationships"] is True


def test_24_cell_shading_applied(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    doc = Document()
    doc.add_table(rows=1, cols=1).rows[0].cells[0].text = "x"
    doc.save(src)
    _table_profile()
    plan = _plan(src, "_b2_table")
    render_reformat(src, plan, out)
    xml = _xml(out)
    assert 'w:fill="EEEEEE"' in xml


def test_25_image_max_height_downscale(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    png = tmp_path / "img.png"
    png.write_bytes(_png_bytes())
    doc = Document()
    p = doc.add_paragraph()
    # 宽 8cm、高 12cm 的图片：宽不超限但高超过 max_height_cm=6
    p.add_run().add_picture(str(png), width=Cm(8), height=Cm(12))
    doc.save(src)
    _profile(
        "_b2_image_h",
        image={
            "max_width_cm": 20,
            "max_height_cm": 6,
            "preserve_aspect_ratio": True,
            "allow_upscale": False,
            "alignment": "center",
        },
    )
    plan = _plan(src, "_b2_image_h")
    render_reformat(src, plan, out)
    xml = _xml(out)
    # 高缩到 6cm=2160000 EMU，宽按比例 4cm=1440000 EMU
    assert 'cy="2160000"' in xml
    assert 'cx="1440000"' in xml


def test_26_aspect_ratio_preserved_on_width_downscale(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    png = tmp_path / "img.png"
    png.write_bytes(_png_bytes())
    doc = Document()
    p = doc.add_paragraph()
    # 12cm x 6cm，宽缩到 8cm 时高应为 4cm
    p.add_run().add_picture(str(png), width=Cm(12), height=Cm(6))
    doc.save(src)
    _image_profile(max_width_cm=8)
    plan = _plan(src, "_b2_image")
    render_reformat(src, plan, out)
    xml = _xml(out)
    assert 'cx="2880000"' in xml
    assert 'cy="1440000"' in xml


def test_23_preservation_failure_no_output(tmp_path, monkeypatch):
    import reformat_engine.renderer as renderer_mod

    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    doc = Document()
    doc.add_paragraph("正文")
    doc.save(src)
    plan = _plan(src)

    real_read = renderer_mod.read_docx

    def fake_read(path):
        ir = real_read(path)
        ir.content_fingerprint.text_sha256 = "deadbeef"
        return ir

    monkeypatch.setattr(renderer_mod, "read_docx", fake_read)
    result = renderer_mod.render_reformat(src, plan, out)
    assert result["status"] == "error"
    assert CONTENT_PRESERVATION_FAILED in result["errors"]
    assert not out.exists()

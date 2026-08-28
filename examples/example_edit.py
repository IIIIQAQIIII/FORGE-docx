"""Synthetic example: inspect + targeted edit with source version protection."""

from pathlib import Path

from docx import Document

from edit_engine.service import edit_document, inspect_document

OUT = Path(__file__).resolve().parent.parent / "outputs"
OUT.mkdir(exist_ok=True)

source = OUT / "example_edit_source.docx"
target = OUT / "example_edited.docx"

doc = Document()
doc.add_paragraph("旧文字需要替换")
doc.add_paragraph("其他段落保持不变")
doc.save(source)

inspection = inspect_document(source, query="旧文字")
block = inspection["blocks"][0]
sha = inspection["source_file_sha256"]

result = edit_document(
    source,
    expected_source_sha256=sha,
    edits=[
        {
            "op": "replace_text",
            "source_locator": block["source_locator"],
            "old_text": "旧文字",
            "new_text": "新文字",
            "expected_occurrences": 1,
        }
    ],
    output_path=target,
)
print(result["status"])
print(result["preservation"]["passed"])

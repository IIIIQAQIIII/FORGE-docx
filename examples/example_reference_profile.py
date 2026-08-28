"""Synthetic example: learn a profile from a reference DOCX and reuse it."""

import os
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from open_format.profile_store import save_user_profile
from open_format.reference_builder import analyze_reference_docx
from reformat_engine.service import reformat_document

tmp = Path(tempfile.mkdtemp(prefix="forge_reference_example_"))
os.environ.setdefault("FORGE_HOME", str(tmp / "forge_home"))

reference = tmp / "reference.docx"
doc = Document()
section = doc.sections[0]
section.top_margin = Cm(3.7)
title = doc.add_paragraph("参考标题")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.bold = True
title.runs[0].font.size = Pt(22)
for i in range(10):
    doc.add_paragraph(f"参考正文第{i}段，这是用于学习格式的合成内容。")
doc.save(reference)

analysis = analyze_reference_docx(reference, base_profile_id="generic_document", profile_id="learned_profile")
print("analysis status:", analysis["status"])

save_user_profile("learned_profile", "学到的格式", "", "reference", "generic_document", analysis["draft"]["rules"])

source = tmp / "source.docx"
d2 = Document()
p = d2.add_paragraph("全新内容标题")
p.runs[0].font.bold = True
p.runs[0].font.size = Pt(22)
d2.add_paragraph("全新内容的正文段落。")
d2.save(source)

result = reformat_document(source, output_path=tmp / "reformatted.docx", saved_profile_id="learned_profile")
print(result["status"])
print(result["content_preservation"]["passed"])

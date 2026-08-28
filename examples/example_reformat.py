"""Synthetic example: reformat a generated DOCX with a built-in profile."""

from pathlib import Path

from docx import Document

from reformat_engine.service import reformat_document

OUT = Path(__file__).resolve().parent.parent / "outputs"
OUT.mkdir(exist_ok=True)

source = OUT / "example_source.docx"
target = OUT / "example_reformatted.docx"

doc = Document()
p = doc.add_paragraph("示例标题")
p.runs[0].font.bold = True
doc.add_paragraph("一、示例章节")
doc.add_paragraph("这是用于演示的正文段落，内容完全由脚本生成。")
doc.save(source)

result = reformat_document(source, output_path=target, explicit_format_hint="正式公文")
print(result["status"])
print(result["output"])
print(result["content_preservation"]["passed"])
